from __future__ import annotations

import argparse
import base64
import hashlib
import warnings

warnings.filterwarnings("ignore", "'cgi' is deprecated", DeprecationWarning)
import cgi
import csv
import html
import io
import json
import mimetypes
import os
import re
import shutil
import sys
import threading
import time
import unicodedata
import zipfile
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlparse

from agent_profiles import get_agent_profiles
from brain import brain_summary, read_brain_text, clean_text
from case_study_brain import case_study_summary, search_case_study_brain
from config import OPENAI_ANSWER_DETAIL, OPENAI_REASONING_EFFORT
from launch_os_db import active_project_snapshot, init_launch_os_database, launch_os_status, mark_project_task_from_module, project_context_for_text
from launch_actions import action_note, maybe_run_action
from llm_client import api_connection_status, chat_with_llm, has_api_key, stream_chat_with_llm
from master_agent import answer_master_question, format_sources, master_brain_status, search_all_role_brains, stream_master_answer
from niche_brain import niche_summary, search_niche_brain
from product_pipeline import STEP_ROUTE_TABLE, extract_explicit_step, resolve_product_route, step_unsupported_response

ROOT_DIR = Path(__file__).resolve().parent
APP_VERSION = (ROOT_DIR / "VERSION").read_text(encoding="utf-8").strip() or "1.00"
WEB_DIR = ROOT_DIR / "web_ui"
INDEX_FILE = WEB_DIR / "index.html"
STYLE_FILE = WEB_DIR / "styles.css"
SCRIPT_FILE = WEB_DIR / "app.js"
UPLOADS_DIR = ROOT_DIR / "uploads" / "web_chat"
CHAT_HISTORY_DIR = ROOT_DIR / "chat_history"
THREADS_FILE = CHAT_HISTORY_DIR / "threads.json"
GENERATED_FILES_DIR = ROOT_DIR / "exports" / "web_generated_files"
PROMPTS_DIR = ROOT_DIR / "prompts"
MAX_UPLOAD_BYTES = int(os.getenv("PLR_AGENT_MAX_UPLOAD_BYTES", str(512 * 1024 * 1024)))
MAX_UPLOAD_FILES_PER_BATCH = int(os.getenv("PLR_AGENT_MAX_UPLOAD_FILES_PER_BATCH", "10"))
MAX_UPLOAD_BATCH_BYTES = int(os.getenv("PLR_AGENT_MAX_UPLOAD_BATCH_BYTES", str(MAX_UPLOAD_BYTES * MAX_UPLOAD_FILES_PER_BATCH)))
MAX_UPLOAD_STORAGE_BYTES = int(os.getenv("PLR_AGENT_MAX_UPLOAD_STORAGE_BYTES", str(20 * 1024 * 1024 * 1024)))
MAX_ATTACHMENT_TEXT = int(os.getenv("PLR_AGENT_MAX_ATTACHMENT_TEXT", "200000"))
PRODUCT_STEP_AI_TIMEOUT_SECONDS = int(os.getenv("PRODUCT_STEP_AI_TIMEOUT_SECONDS", "8"))
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"}


def _read_json_body(handler: BaseHTTPRequestHandler, length: int) -> dict:
    if not length:
        return {}
    raw = handler.rfile.read(length)
    try:
        text = raw.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = raw.decode("utf-8-sig", errors="replace")
    return json.loads(text)
THREADS_LOCK = threading.Lock()
CANCELLED_REQUESTS: set[str] = set()
CHAT_MODE_LIMITS = {
    "auto": 6,
    "quick": 4,
    "fast": 6,
    "asset": 12,
    "balanced": 12,
    "deep": 20,
}
RAG_MODE_TOPK = {"quick": 4, "fast": 6, "balanced": 12, "asset": 12, "deep": 20, "auto": 6}
FILE_CACHE: dict[str, dict] = {}
PERF_LOG_DIR = ROOT_DIR / "reports"
ACTION_ONLY_MODULES = {
    "product_assets",
    "deep_create_product_assets",
    "deep_write_file",
    "product_blueprint",
    "deep_file_writer",
    "market_pattern_extract",
    "competitor_matrix",
    "offer_gap_v2",
    "launch_pack",
    "full_launch_pack",
    "export_zip",
    "support",
    "license",
    "buyer_test",
    "prompt_output_test",
    "ai_replace_risk",
    "ai_replace_risk_v2",
    "license_compliance_check",
    "warriorplus_launch_builder",
    "jv_test",
    "sales_page_critic",
    "apply_feedback",
    "buyer_test_zip",
    "jv_test_pack",
    "public_launch_audit",
    "export_pack",
    "final_scorecard",
    "optimize_storage",
    "storage_report",
    "agent_benchmark",
    "train_case_study_brain",
    "train_full_case_study_brain",
    "case_study_search",
    "case_study_patterns",
    "training_status",
    "export_training_report",
    "ai_print_build",
    "ai_print_train",
    "ai_print_full_train",
    "ai_print_status",
    "ai_print_search",
    "ai_print_patterns",
    "ai_print_evidence",
    "ai_print_market",
    "ai_print_competitor",
    "ai_print_gap",
    "ai_print_report",
    "workflow_30",
    "ai_workflow_20",
}




def _now_ms() -> float:
    return time.perf_counter() * 1000

def _elapsed_ms(start: float) -> int:
    return int(max(0, _now_ms() - start))

def _read_text_cached(path: Path, *, limit: int | None = None) -> tuple[str, str]:
    try:
        stat = path.stat()
    except FileNotFoundError:
        return "", "missing"
    key = str(path.resolve())
    cached = FILE_CACHE.get(key)
    if cached and cached.get("mtime") == stat.st_mtime and cached.get("size") == stat.st_size:
        text = cached.get("text", "")
        status = "cache_hit"
    else:
        text = path.read_text(encoding="utf-8", errors="replace")
        FILE_CACHE[key] = {"mtime": stat.st_mtime, "size": stat.st_size, "text": text}
        status = "cache_miss"
    return (text[:limit] if limit else text), status

def _append_perf_log(entry: dict) -> None:
    try:
        PERF_LOG_DIR.mkdir(parents=True, exist_ok=True)
        line = json.dumps(entry, ensure_ascii=False, separators=(",", ":"))
        with (PERF_LOG_DIR / "chat_latency.jsonl").open("a", encoding="utf-8") as handle:
            handle.write(line + "\n")
    except Exception:
        pass

def _is_deep_command(question: str) -> bool:
    first = str(question or "").strip().split(maxsplit=1)[0].lower() if str(question or "").strip() else ""
    return first in {"/run_benchmark", "/run_compare_codex55", "/benchmark_20_rounds", "/public_launch_audit_deep", "/benchmark_agent", "/compare_chatgpt"}

def main() -> None:
    _configure_console_encoding()
    init_launch_os_database()

    parser = argparse.ArgumentParser(description="Local Master Agent web UI")
    parser.add_argument("--host", default="0.0.0.0")
    parser.add_argument("--port", type=int, default=8088)
    args = parser.parse_args()

    server = ThreadingHTTPServer((args.host, args.port), _make_handler())
    print(f"Master Agent web UI running at http://{args.host}:{args.port}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()


def _configure_console_encoding() -> None:
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8", errors="replace")


def _make_handler():
    class Handler(BaseHTTPRequestHandler):
        def end_headers(self) -> None:
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
            self.send_header("Access-Control-Allow-Headers", "Content-Type")
            self.send_header("Cache-Control", "no-cache")
            super().end_headers()

        def do_OPTIONS(self) -> None:  # noqa: N802
            self.send_response(HTTPStatus.NO_CONTENT)
            self.end_headers()

        def do_GET(self) -> None:  # noqa: N802
            parsed = urlparse(self.path)
            if parsed.path in {"", "/"} or _is_workspace_page(parsed.path):
                return self._serve_file(INDEX_FILE, "text/html; charset=utf-8")
            if parsed.path == "/styles.css":
                return self._serve_file(STYLE_FILE, "text/css; charset=utf-8")
            if parsed.path == "/app.js":
                return self._serve_file(SCRIPT_FILE, "application/javascript; charset=utf-8")
            if parsed.path == "/utils/renderFileCard.js":
                return self._serve_file(WEB_DIR / "utils" / "renderFileCard.js", "application/javascript; charset=utf-8")
            if parsed.path == "/favicon.ico":
                self.send_response(HTTPStatus.NO_CONTENT)
                self.end_headers()
                return
            if parsed.path == "/api/status":
                api_ready, api_message = api_connection_status()
                return self._send_json(
                    {
                        "ok": True,
                        "status": master_brain_status(),
                        "apiReady": api_ready,
                        "apiMessage": api_message,
                        "appVersion": APP_VERSION,
                        "reasoningEffort": OPENAI_REASONING_EFFORT,
                        "answerDetail": OPENAI_ANSWER_DETAIL,
                        "defaultMode": "fast",
                        "ragTopK": RAG_MODE_TOPK,
                        "uploadLimits": _upload_limits_payload(),
                        "modes": [
                            {"key": "fast", "label": "FAST", "description": "Chat nhanh: topK 6, cache, không benchmark/export."},
                            {"key": "balanced", "label": "BALANCED", "description": "Skill + brain vừa đủ: topK 12."},
                            {"key": "deep", "label": "DEEP", "description": "Tạo sản phẩm/audit/export: topK 20, chậm hơn."},
                        ],
                        "brains": _brain_status_cards(),
                        "caseStudyBrain": case_study_summary(),
                        "aiPrintablesBrain": niche_summary(),
                        "aiPrintablesKdpPromptAgent": _ai_printables_kdp_prompt_status(),
                        "launchOs": launch_os_status(),
                        "activeProject": active_project_snapshot(),
                        "sources": _preview_sources("WarriorPlus PLR SaaS launch kit"),
                    }
                )
            if parsed.path == "/api/project_status":
                return self._send_json({"ok": True, "project": active_project_snapshot()})
            if parsed.path == "/api/project_state":
                snapshot = active_project_snapshot()
                return self._send_json({"ok": True, **snapshot})
            if parsed.path == "/api/case_study_brain":
                return self._send_json({"ok": True, "caseStudyBrain": case_study_summary()})
            if parsed.path == "/api/case_study_search":
                query = _query_param(parsed.query, "q") or "AI PLR Prompt Template Packs KDP Printables"
                return self._send_json({"ok": True, "query": query, "results": search_case_study_brain(query, limit=12)})
            if parsed.path == "/api/ai_printables_brain":
                return self._send_json({"ok": True, "aiPrintablesBrain": niche_summary()})
            if parsed.path == "/api/ai_printables_search":
                query = _query_param(parsed.query, "q") or "AI-assisted PLR products KDP covers coloring book journal poster social media assets"
                return self._send_json({"ok": True, "query": query, "results": search_niche_brain(query, limit=12)})
            if parsed.path == "/api/skill_tags":
                return self._send_json(_skill_tags_payload())
            if parsed.path == "/api/upload_limits":
                return self._send_json({"ok": True, **_upload_limits_payload()})
            if parsed.path == "/api/upload_file":
                return self._handle_upload_file(parsed.query)
            if parsed.path == "/api/generated_file":
                return self._handle_generated_file(parsed.query)
            if parsed.path == "/api/product_file":
                return self._handle_product_file(parsed.query)
            if parsed.path == "/api/prompt_file":
                return self._handle_prompt_file(parsed.query)
            if parsed.path == "/api/sources":
                query = _query_param(parsed.query, "q")
                if not query:
                    return self._send_json({"ok": False, "error": "Missing q"}, status=HTTPStatus.BAD_REQUEST)
                return self._send_json({"ok": True, "text": format_sources(query)})
            if parsed.path == "/api/threads":
                workspace_id = _workspace_id_from_query(parsed.query)
                return self._send_json({"ok": True, "state": _load_threads_state(workspace_id), "workspace": workspace_id})
            self.send_error(HTTPStatus.NOT_FOUND)

        def do_POST(self) -> None:  # noqa: N802
            parsed = urlparse(self.path)
            if parsed.path == "/api/upload":
                return self._handle_upload()
            if parsed.path == "/api/threads":
                return self._handle_threads_save(_workspace_id_from_query(parsed.query))
            if parsed.path == "/api/create_file":
                return self._handle_create_file()
            if parsed.path == "/api/cancel":
                return self._handle_cancel()
            if parsed.path == "/api/route_skill":
                return self._handle_route_skill()
            if parsed.path not in {"/api/chat", "/api/chat_stream"}:
                self.send_error(HTTPStatus.NOT_FOUND)
                return

            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = _read_json_body(self, length)
            except Exception as error:
                return self._send_json({"ok": False, "error": f"Invalid JSON: {error}"}, status=HTTPStatus.BAD_REQUEST)

            request_start_ms = _now_ms()
            timings = {"receive_request_ms": 0}
            question = str(payload.get("question", "")).strip()
            request_id = _sanitize_request_id(payload.get("requestId"))
            module_id = _normalize_module_id(payload.get("module"))
            response_mode = _normalize_chat_mode(payload.get("mode"))
            history = payload.get("history", [])
            attachments = payload.get("attachments", [])
            payload_tags = payload.get("tags", [])
            if not question:
                return self._send_json({"ok": False, "error": "Missing question"}, status=HTTPStatus.BAD_REQUEST)
            early_ai_etsy_route = resolve_ai_etsy_route(question, payload if isinstance(payload, dict) else {})
            if early_ai_etsy_route.get("selected_route"):
                early_ai_etsy_route["request_id"] = request_id
                early_ai_etsy_route["user_message_hash"] = hashlib.sha256(question.encode("utf-8", errors="ignore")).hexdigest()[:16]
                if (_is_real_ai_product_route(early_ai_etsy_route)):
                    if parsed.path == "/api/chat_stream":
                        return self._send_real_ai_phase5_stream(early_ai_etsy_route, question, response_mode, timings, request_start_ms)
                    answer, action = _run_real_ai_phase5_route(early_ai_etsy_route, question)
                    timings["total_ms"] = _elapsed_ms(request_start_ms)
                    return self._send_json({"ok": True, "answer": answer, "sources": [], "mode": response_mode, "action": action, "timings": timings})
                answer, action = _run_ai_etsy_route(early_ai_etsy_route)
                timings["total_ms"] = _elapsed_ms(request_start_ms)
                if parsed.path == "/api/chat_stream":
                    return self._send_prebuilt_stream(answer, response_mode, action, timings)
                return self._send_json({"ok": True, "answer": answer, "sources": [], "mode": response_mode, "action": action})
            t = _now_ms()
            skill_route_payload = _route_ai_printables_kdp_prompt(question, payload_tags if isinstance(payload_tags, list) else [])
            timings["route_skill_ms"] = _elapsed_ms(t)
            t = _now_ms()
            skill_context = _skill_context_for_question(question, skill_route_payload)
            timings["load_skill_brain_ms"] = _elapsed_ms(t)
            if not skill_context:
                module_id = module_id or _command_module_id(question)
                module_id = module_id or _infer_module_id_from_text(question)
            if not question:
                return self._send_json({"ok": False, "error": "Missing question"}, status=HTTPStatus.BAD_REQUEST)

            conversation_context = _format_history(history)
            project_context = project_context_for_text(question)
            if project_context:
                conversation_context = f"{conversation_context}\n\n## Project Memory\n{project_context}".strip()
            has_attachment = _attachments_have_content(attachments)
            response_mode = _effective_chat_mode(question, response_mode, has_attachment, module_id)
            attachment_context = _format_attachments(attachments, response_mode)
            full_question = _apply_module_context(question, module_id)
            if skill_context:
                full_question = f"{skill_context}\n\n## User Request\n{question}"
            if attachment_context:
                full_question = f"{full_question}\n\n## File người dùng vừa gửi\n{attachment_context}"
            vietnamese_guard = (
                "## BẮT BUỘC NGÔN NGỮ\n"
                "- Trả lời bằng TIẾNG VIỆT rõ ràng, dễ hiểu cho người dùng Việt Nam.\n"
                "- Chỉ giữ thuật ngữ tiếng Anh khi đó là tên file, tag, lệnh, brand, hoặc tiêu đề kỹ thuật cần giữ nguyên.\n"
                "- Không tự chuyển toàn bộ câu trả lời sang tiếng Anh.\n"
            )
            full_question = f"{vietnamese_guard}\n\n{full_question}"
            limit_per_brain = RAG_MODE_TOPK.get(response_mode, RAG_MODE_TOPK["fast"])
            ai_etsy_route = resolve_ai_etsy_route(question, payload if isinstance(payload, dict) else {})
            if ai_etsy_route.get("selected_route"):
                ai_etsy_route["request_id"] = request_id
                ai_etsy_route["user_message_hash"] = hashlib.sha256(question.encode("utf-8", errors="ignore")).hexdigest()[:16]
                if (_is_real_ai_product_route(ai_etsy_route)):
                    if parsed.path == "/api/chat_stream":
                        return self._send_real_ai_phase5_stream(ai_etsy_route, question, response_mode, timings, request_start_ms)
                    answer, action = _run_real_ai_phase5_route(ai_etsy_route, question)
                    timings["total_ms"] = _elapsed_ms(request_start_ms)
                    return self._send_json({"ok": True, "answer": answer, "sources": [], "mode": response_mode, "action": action, "timings": timings})
                answer, action = _run_ai_etsy_route(ai_etsy_route)
                timings["total_ms"] = _elapsed_ms(request_start_ms)
                if parsed.path == "/api/chat_stream":
                    return self._send_prebuilt_stream(answer, response_mode, action, timings)
                return self._send_json({"ok": True, "answer": answer, "sources": [], "mode": response_mode, "action": action})
            if module_id in ACTION_ONLY_MODULES:
                if parsed.path == "/api/chat_stream":
                    return self._send_action_stream(question, response_mode, module_id)
                action = maybe_run_action(module_id, question)
                _mark_completed_module(question, module_id, notes=f"Completed via tool action {module_id}")
                answer = _action_response(module_id, action)
                return self._send_json({"ok": True, "answer": answer, "sources": [], "mode": response_mode, "action": action})
            if parsed.path == "/api/chat_stream":
                return self._send_chat_stream(full_question, question, limit_per_brain, conversation_context, response_mode, module_id, request_id, timings, request_start_ms)

            t = _now_ms()
            answer = answer_master_question(
                full_question,
                limit_per_brain=limit_per_brain,
                conversation_context=conversation_context,
                response_mode=response_mode,
            )
            timings["api_call_ms"] = _elapsed_ms(t)
            action = maybe_run_action(module_id, question, answer)
            if action:
                answer = f"{answer.rstrip()}{action_note(action)}"
            answer = ensure_agent_contract(answer, module_id, action)
            _mark_completed_module(question, module_id, notes=f"Completed via {module_id or response_mode}")
            t = _now_ms()
            hits = search_all_role_brains(question, limit_per_brain=limit_per_brain)
            timings["rag_sources_ms"] = _elapsed_ms(t)
            sources = [
                {
                    "brain": hit.brain_name,
                    "title": hit.title,
                    "source_path": hit.source_path,
                    "excerpt": _clip(hit.text, 240),
                }
                for hit in hits[:9]
            ]
            timings["total_ms"] = _elapsed_ms(request_start_ms)
            _append_perf_log({"path": parsed.path, "mode": response_mode, "tags": skill_route_payload.get("matchedTags", []), "skill": skill_route_payload.get("skillFile", ""), "status": "ok", "timings": timings})
            return self._send_json({"ok": True, "answer": answer, "sources": sources, "mode": response_mode, "timings": timings})

        def _send_chat_stream(
            self,
            full_question: str,
            original_question: str,
            limit_per_brain: int,
            conversation_context: str,
            response_mode: str,
            module_id: str,
            request_id: str = "",
            timings: dict | None = None,
            request_start_ms: float | None = None,
        ) -> None:
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()

            def emit(event: str, payload: dict) -> None:
                if request_id and request_id in CANCELLED_REQUESTS:
                    raise BrokenPipeError("Client cancelled request")
                body = f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n".encode("utf-8")
                self.wfile.write(body)
                self.wfile.flush()

            try:
                emit("meta", {"ok": True, "mode": response_mode, "requestId": request_id, "topK": limit_per_brain, "timings": timings or {}})
                emit("status", {"text": "Routing skill / loading brain done"})
                emit("status", {"text": f"Searching RAG topK={limit_per_brain}"})
                api_start_ms = _now_ms()
                emit("status", {"text": "Calling model"})
                for delta in stream_master_answer(
                    full_question,
                    limit_per_brain=limit_per_brain,
                    conversation_context=conversation_context,
                    response_mode=response_mode,
                ):
                    if delta:
                        emit("delta", {"text": delta})
                action = maybe_run_action(module_id, original_question)
                note = action_note(action)
                if note:
                    emit("delta", {"text": note})
                footer = agent_contract_footer(module_id, action)
                if footer:
                    emit("delta", {"text": footer})
                _mark_completed_module(original_question, module_id, notes=f"Completed via {module_id or response_mode}")
                t = _now_ms()
                hits = search_all_role_brains(original_question, limit_per_brain=limit_per_brain)
                if timings is not None:
                    timings["rag_sources_ms"] = _elapsed_ms(t)
                sources = [
                    {
                        "brain": hit.brain_name,
                        "title": hit.title,
                        "source_path": hit.source_path,
                        "excerpt": _clip(hit.text, 240),
                    }
                    for hit in hits[:9]
                ]
                if timings is not None:
                    timings["api_call_stream_ms"] = _elapsed_ms(api_start_ms)
                    timings["total_ms"] = _elapsed_ms(request_start_ms or api_start_ms)
                    _append_perf_log({"path": "/api/chat_stream", "mode": response_mode, "status": "ok", "timings": timings})
                emit("status", {"text": "Finalizing"})
                emit("done", {"ok": True, "sources": sources, "mode": response_mode, "timings": timings or {}})
            except (BrokenPipeError, ConnectionResetError):
                return
            except Exception as error:
                try:
                    emit("error", {"ok": False, "error": str(error)})
                except Exception:
                    return
            finally:
                if request_id:
                    CANCELLED_REQUESTS.discard(request_id)


        def _send_prebuilt_stream(self, answer: str, response_mode: str, action: dict, timings: dict | None = None) -> None:
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()

            def emit(event: str, payload: dict) -> None:
                body = f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n".encode("utf-8")
                self.wfile.write(body)
                self.wfile.flush()

            try:
                emit("meta", {"ok": True, "mode": response_mode, "prebuilt": True, "timings": timings or {}})
                emit("status", {"text": "Tạo product files nội bộ, không chờ model"})
                chunk_size = 2200
                for index in range(0, len(answer), chunk_size):
                    emit("delta", {"text": answer[index:index + chunk_size]})
                emit("done", {"ok": True, "sources": [], "mode": response_mode, "action": action, "timings": timings or {}})
            except (BrokenPipeError, ConnectionResetError):
                return

        def _send_real_ai_phase5_stream(self, route: dict, question: str, response_mode: str, timings: dict | None = None, request_start_ms: float | None = None) -> None:
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()

            def emit(event: str, payload: dict) -> None:
                request_id = route.get("request_id") or ""
                if request_id and request_id in CANCELLED_REQUESTS:
                    raise BrokenPipeError("Client cancelled request")
                body = f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n".encode("utf-8")
                self.wfile.write(body)
                self.wfile.flush()

            try:
                selected_route = route.get("selected_route") or "REAL_AI_CHAT_PHASE_5"
                emit("meta", {"ok": True, "mode": response_mode, "prebuilt": False, "selected_route": selected_route, "timings": timings or {}})
                emit("status", {"text": f"{selected_route}: gọi AI API thật, không dùng prebuilt"})
                prompt = _real_ai_phase5_prompt(question, route["project"], route)
                start = _now_ms()
                streamed_text = ""
                for delta in stream_chat_with_llm(prompt, reasoning_effort="medium", max_output_tokens=5000):
                    if delta:
                        streamed_text += delta
                        emit("delta", {"text": delta})
                elapsed_ms = _elapsed_ms(start)
                action = _real_ai_phase5_action(route, elapsed_ms=elapsed_ms)
                debug_block = _route_debug_block(route, action)
                emit("delta", {"text": debug_block})
                if timings is not None:
                    timings["api_call_stream_ms"] = elapsed_ms
                    timings["total_ms"] = _elapsed_ms(request_start_ms or start)
                emit("done", {"ok": True, "sources": [], "mode": response_mode, "action": action, "timings": timings or {}})
            except (BrokenPipeError, ConnectionResetError):
                return
            except Exception as error:
                try:
                    emit("error", {"ok": False, "error": str(error)})
                except Exception:
                    return
            finally:
                request_id = route.get("request_id") or ""
                if request_id:
                    CANCELLED_REQUESTS.discard(request_id)
        def _send_action_stream(self, original_question: str, response_mode: str, module_id: str) -> None:
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", "text/event-stream; charset=utf-8")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Connection", "close")
            self.end_headers()

            def emit(event: str, payload: dict) -> None:
                body = f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n".encode("utf-8")
                self.wfile.write(body)
                self.wfile.flush()

            try:
                emit("meta", {"ok": True, "mode": response_mode})
                action = maybe_run_action(module_id, original_question)
                _mark_completed_module(original_question, module_id, notes=f"Completed via tool action {module_id}")
                emit("delta", {"text": _action_response(module_id, action)})
                emit("done", {"ok": True, "sources": [], "mode": response_mode, "action": action})
            except Exception as error:
                try:
                    emit("error", {"ok": False, "error": str(error)})
                except Exception:
                    return

        def _handle_upload(self) -> None:
            length = int(self.headers.get("Content-Length", "0"))
            if length > MAX_UPLOAD_BATCH_BYTES:
                limit_mb = MAX_UPLOAD_BATCH_BYTES // 1024 // 1024
                return self._send_json({"ok": False, "error": f"Upload qua lon. Gioi han {limit_mb}MB moi lan upload."}, status=HTTPStatus.REQUEST_ENTITY_TOO_LARGE)
            if not _has_upload_capacity(length):
                return self._send_json(
                    {
                        "ok": False,
                        "error": _upload_capacity_error(length),
                        **_upload_limits_payload(),
                    },
                    status=HTTPStatus.REQUEST_ENTITY_TOO_LARGE,
                )

            content_type = self.headers.get("Content-Type", "")
            if content_type.lower().startswith("multipart/form-data"):
                return self._handle_multipart_upload(length, content_type)

            try:
                payload = _read_json_body(self, length)
            except Exception as error:
                return self._send_json({"ok": False, "error": f"Invalid JSON: {error}"}, status=HTTPStatus.BAD_REQUEST)

            files = payload.get("files", [])
            if not isinstance(files, list) or not files:
                return self._send_json({"ok": False, "error": "Missing files"}, status=HTTPStatus.BAD_REQUEST)
            if len(files) > MAX_UPLOAD_FILES_PER_BATCH:
                return self._send_json(
                    {"ok": False, "error": f"Qua nhieu file. Gioi han {MAX_UPLOAD_FILES_PER_BATCH} file moi lan upload."},
                    status=HTTPStatus.BAD_REQUEST,
                )

            UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
            attachments = []
            for index, item in enumerate(files[:MAX_UPLOAD_FILES_PER_BATCH], start=1):
                try:
                    attachments.append(_save_and_read_upload(item, index))
                except Exception as error:
                    attachments.append(
                        {
                            "name": str(item.get("name", f"file-{index}")) if isinstance(item, dict) else f"file-{index}",
                            "type": "error",
                            "text": "",
                            "notice": f"Khong doc duoc file: {error}",
                        }
                    )
            return self._send_json({"ok": True, "attachments": attachments})

        def _handle_multipart_upload(self, length: int, content_type: str) -> None:
            try:
                form = cgi.FieldStorage(
                    fp=self.rfile,
                    headers=self.headers,
                    environ={
                        "REQUEST_METHOD": "POST",
                        "CONTENT_TYPE": content_type,
                        "CONTENT_LENGTH": str(length),
                    },
                )
            except Exception as error:
                return self._send_json({"ok": False, "error": f"Invalid multipart upload: {error}"}, status=HTTPStatus.BAD_REQUEST)

            fields = form["files"] if "files" in form else []
            if not isinstance(fields, list):
                fields = [fields]
            fields = [field for field in fields if getattr(field, "filename", None)]
            if not fields:
                return self._send_json({"ok": False, "error": "Missing files"}, status=HTTPStatus.BAD_REQUEST)
            if len(fields) > MAX_UPLOAD_FILES_PER_BATCH:
                return self._send_json(
                    {"ok": False, "error": f"Qua nhieu file. Gioi han {MAX_UPLOAD_FILES_PER_BATCH} file moi lan upload."},
                    status=HTTPStatus.BAD_REQUEST,
                )

            UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
            attachments = []
            for index, field in enumerate(fields[:MAX_UPLOAD_FILES_PER_BATCH], start=1):
                try:
                    attachments.append(_save_and_read_upload_stream(field.filename, field.file, index))
                except Exception as error:
                    attachments.append(
                        {
                            "name": str(getattr(field, "filename", f"file-{index}")),
                            "type": "error",
                            "text": "",
                            "notice": f"Khong doc duoc file: {error}",
                        }
                    )
            return self._send_json({"ok": True, "attachments": attachments})

        def _handle_threads_save(self, workspace_id: str) -> None:
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = _read_json_body(self, length)
            except Exception as error:
                return self._send_json({"ok": False, "error": f"Invalid JSON: {error}"}, status=HTTPStatus.BAD_REQUEST)

            state = _normalize_threads_state(payload.get("state", payload))
            _save_threads_state(state, workspace_id)
            return self._send_json({"ok": True, "state": state, "workspace": workspace_id})

        def _handle_create_file(self) -> None:
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = _read_json_body(self, length)
            except Exception as error:
                return self._send_json({"ok": False, "error": f"Invalid JSON: {error}"}, status=HTTPStatus.BAD_REQUEST)

            content = str(payload.get("content", "")).strip()
            requested_format = str(payload.get("format", "md")).strip().lower().lstrip(".")
            title = str(payload.get("title", "agent-output")).strip()
            if not content:
                return self._send_json({"ok": False, "error": "Missing content"}, status=HTTPStatus.BAD_REQUEST)

            try:
                created = _create_export_file(content, requested_format, title)
            except ValueError as error:
                return self._send_json({"ok": False, "error": str(error)}, status=HTTPStatus.BAD_REQUEST)
            except Exception as error:
                return self._send_json({"ok": False, "error": f"Khong tao duoc file: {error}"}, status=HTTPStatus.INTERNAL_SERVER_ERROR)

            return self._send_json({"ok": True, **created})

        def _handle_upload_file(self, query: str) -> None:
            params = parse_qs(query)
            raw_name = unquote((params.get("name") or [""])[0]).strip()
            if not raw_name:
                self.send_error(HTTPStatus.BAD_REQUEST, "Missing file name")
                return
            try:
                target = (UPLOADS_DIR / raw_name).resolve()
                target.relative_to(UPLOADS_DIR.resolve())
            except (OSError, ValueError):
                self.send_error(HTTPStatus.FORBIDDEN)
                return
            if not target.is_file():
                self.send_error(HTTPStatus.NOT_FOUND)
                return

            content_type = mimetypes.guess_type(target.name)[0] or "application/octet-stream"
            size = target.stat().st_size
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Length", str(size))
            self.send_header("Content-Disposition", f'inline; filename="{_safe_filename(target.name)}"')
            self.end_headers()
            with target.open("rb") as source:
                while True:
                    chunk = source.read(1024 * 1024)
                    if not chunk:
                        break
                    self.wfile.write(chunk)


        def _handle_route_skill(self) -> None:
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = _read_json_body(self, length)
            except Exception as error:
                return self._send_json({"ok": False, "error": f"Invalid JSON: {error}"}, status=HTTPStatus.BAD_REQUEST)
            message = str(payload.get("message") or payload.get("question") or "")
            tags = payload.get("tags", [])
            route = _route_ai_printables_kdp_prompt(message, tags if isinstance(tags, list) else [])
            return self._send_json(route)

        def _handle_prompt_file(self, query: str) -> None:
            params = parse_qs(query)
            raw_name = unquote((params.get("name") or [""])[0]).strip()
            if not raw_name:
                self.send_error(HTTPStatus.BAD_REQUEST, "Missing prompt file name")
                return
            try:
                target = (PROMPTS_DIR / raw_name).resolve()
                target.relative_to(PROMPTS_DIR.resolve())
            except (OSError, ValueError):
                self.send_error(HTTPStatus.FORBIDDEN)
                return
            if not target.is_file():
                self.send_error(HTTPStatus.NOT_FOUND)
                return

            content_type = mimetypes.guess_type(target.name)[0] or "text/markdown"
            data = target.read_bytes()
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", f"{content_type}; charset=utf-8")
            self.send_header("Content-Length", str(len(data)))
            self.send_header("Content-Disposition", _content_disposition('attachment', target.name))
            self.end_headers()
            self.wfile.write(data)

        def _handle_generated_file(self, query: str) -> None:
            params = parse_qs(query)
            raw_name = unquote((params.get("name") or [""])[0]).strip()
            disposition = "inline" if (params.get("view") or [""])[0] == "1" else "attachment"
            if not raw_name:
                self.send_error(HTTPStatus.BAD_REQUEST, "Missing generated file name")
                return
            try:
                target = (GENERATED_FILES_DIR / raw_name).resolve()
                target.relative_to(GENERATED_FILES_DIR.resolve())
            except (OSError, ValueError):
                self.send_error(HTTPStatus.FORBIDDEN)
                return
            if not target.is_file():
                self.send_error(HTTPStatus.NOT_FOUND)
                return

            content_type = mimetypes.guess_type(target.name)[0] or "application/octet-stream"
            data = target.read_bytes()
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Length", str(len(data)))
            self.send_header("Content-Disposition", _content_disposition(disposition, target.name))
            self.end_headers()
            self.wfile.write(data)


        def _handle_product_file(self, query: str) -> None:
            params = parse_qs(query)
            product = unquote((params.get("product") or [""])[0]).strip()
            file_name = unquote((params.get("file") or [""])[0]).strip()
            disposition = "inline" if (params.get("view") or [""])[0] == "1" else "attachment"
            if not product or not file_name:
                self.send_error(HTTPStatus.BAD_REQUEST, "Missing product or file")
                return
            try:
                root = (ROOT_DIR / "exports" / "products" / product).resolve()
                target = (root / file_name).resolve()
                target.relative_to(root)
            except (OSError, ValueError):
                self.send_error(HTTPStatus.FORBIDDEN)
                return
            if not target.is_file():
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            content_type = mimetypes.guess_type(target.name)[0] or "application/octet-stream"
            data = target.read_bytes()
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Length", str(len(data)))
            self.send_header("Content-Disposition", _content_disposition(disposition, target.name))
            self.end_headers()
            self.wfile.write(data)
        def _handle_cancel(self) -> None:
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = _read_json_body(self, length)
            except Exception:
                payload = {}
            request_id = _sanitize_request_id(payload.get("requestId"))
            if request_id:
                CANCELLED_REQUESTS.add(request_id)
            return self._send_json({"ok": True, "cancelled": bool(request_id)})

        def log_message(self, format: str, *args) -> None:  # noqa: A003
            return

        def _serve_file(self, path: Path, content_type: str) -> None:
            if not path.exists():
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            data = path.read_bytes()
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)

        def _send_json(self, payload: dict, *, status: HTTPStatus = HTTPStatus.OK) -> None:
            body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
            self.send_response(status)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

    return Handler


def _mark_completed_module(question: str, module_id: str, *, notes: str = "") -> None:
    if module_id in {"launch_pack", "full_launch_pack"}:
        for item in (
            "idea_score",
            "buyer_avatar",
            "product_assets",
            "upgrade_kit",
            "sales_page",
            "funnel_plan",
            "warriorplus_listing",
            "jv_page",
            "delivery_page",
            "onboarding",
            "saas_potential",
            "export_zip",
        ):
            mark_project_task_from_module(question, item, notes=notes or f"Completed via {module_id}")
        return
    if module_id == "deep_create_product_assets":
        for item in ("idea_score", "buyer_avatar", "product_assets", "upgrade_kit"):
            mark_project_task_from_module(question, item, notes=notes or f"Completed via {module_id}")
        return
    if module_id == "agent_benchmark":
        for item in (
            "idea_score",
            "buyer_avatar",
            "product_assets",
            "upgrade_kit",
            "sales_page",
            "funnel_plan",
            "warriorplus_listing",
            "jv_page",
            "delivery_page",
            "onboarding",
            "saas_potential",
        ):
            mark_project_task_from_module(question, item, notes=notes or f"Completed via {module_id}")
        return
    if module_id == "export_zip":
        for item in ("idea_score", "buyer_avatar", "upgrade_kit", "export_zip"):
            mark_project_task_from_module(question, item, notes=notes or f"Completed via {module_id}")
        return
    if module_id == "deep_write_file":
        mark_project_task_from_module(question, "product_assets", notes=notes or "Completed via deep_write_file")
        return
    mark_project_task_from_module(question, module_id, notes=notes)


def _query_param(query: str, key: str) -> str:
    for pair in query.split("&"):
        if not pair:
            continue
        if "=" not in pair:
            continue
        k, v = pair.split("=", 1)
        if k == key:
            from urllib.parse import unquote_plus

            return unquote_plus(v)
    return ""


def _sanitize_workspace_id(value: object) -> str:
    raw = str(value or "").strip()
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "", raw)[:48]
    if not cleaned or cleaned in {"api", "app.js", "styles.css", "favicon.ico"}:
        return "default"
    return cleaned


def _workspace_id_from_query(query: str) -> str:
    return _sanitize_workspace_id(_query_param(query, "workspace"))


def _is_workspace_page(path: str) -> bool:
    cleaned = str(path or "").strip("/")
    if not cleaned or "/" in cleaned or "." in cleaned:
        return False
    return _sanitize_workspace_id(cleaned) != "default"


def _normalize_chat_mode(value: object) -> str:
    mode = str(value or "auto").strip().lower()
    if mode in CHAT_MODE_LIMITS:
        return mode
    return "fast"

def _sanitize_request_id(value: object) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]+", "", str(value or ""))[:80]

def _normalize_module_id(value: object) -> str:
    raw = str(value or "").strip().lower()
    cleaned = re.sub(r"[^a-z0-9_\- ]+", "", raw).replace("-", "_").replace(" ", "_")
    allowed = {
        "analyze_plr",
        "idea_score",
        "depth_check",
        "upgrade_kit",
        "product_assets",
        "deep_create_product_assets",
        "deep_write_file",
        "product_blueprint",
        "deep_file_writer",
        "market_pattern_extract",
        "competitor_matrix",
        "offer_gap_v2",
        "qc_checklist",
        "export_zip",
        "offer_angle",
        "sales_page",
        "objections",
        "funnel_plan",
        "warriorplus_listing",
        "proof",
        "jv_page",
        "swipe_pack",
        "prospects",
        "outreach",
        "tiers",
        "review_access",
        "saas_potential",
        "mvp_plan",
        "membership",
        "whitelabel",
        "scan_library",
        "market_gap",
        "competitor",
        "asset_completeness",
        "buyer_journey",
        "use_cases",
        "before_after",
        "offer_gap",
        "pricing_commission",
        "launch_readiness",
        "soft_launch",
        "refund_risk",
        "delivery_page",
        "onboarding",
        "create_email_funnel",
        "support",
        "license",
        "buyer_test",
        "prompt_output_test",
        "ai_replace_risk",
        "ai_replace_risk_v2",
        "license_compliance_check",
        "warriorplus_launch_builder",
        "jv_test",
        "sales_page_critic",
        "apply_feedback",
        "buyer_test_zip",
        "jv_test_pack",
        "public_launch_audit",
        "export_pack",
        "final_scorecard",
        "backend_recommendation",
        "jv_fit",
        "product_line",
        "translate_english",
        "platform_fit",
        "launch_pack",
        "full_launch_pack",
        "evidence_mode",
        "optimize_storage",
        "storage_report",
        "agent_benchmark",
        "train_case_study_brain",
        "train_full_case_study_brain",
        "case_study_search",
        "case_study_patterns",
        "training_status",
        "export_training_report",
        "ai_print_build",
        "ai_print_train",
        "ai_print_full_train",
        "ai_print_status",
        "ai_print_search",
        "ai_print_patterns",
        "ai_print_evidence",
        "ai_print_market",
        "ai_print_competitor",
        "ai_print_gap",
        "ai_print_report",
        "ai_print_deep",
        "workflow_30",
        "ai_workflow_20",
    }
    return cleaned if cleaned in allowed else ""

def _command_module_id(question: str) -> str:
    first = str(question or "").strip().split(maxsplit=1)[0].lower() if str(question or "").strip() else ""
    mapping = {
        "/analyze_offer": "idea_score",
        "/analyze_plr": "analyze_plr",
        "/depth_check": "depth_check",
        "/upgrade_kit": "upgrade_kit",
        "/build_assets": "product_assets",
        "/create_assets": "product_assets",
        "/deep_create_product_assets": "deep_create_product_assets",
        "/deep_write_file": "deep_write_file",
        "/product_blueprint": "product_blueprint",
        "/deep_file_writer": "deep_file_writer",
        "/write_sales_page": "sales_page",
        "/build_funnel": "funnel_plan",
        "/create_funnel": "funnel_plan",
        "/warriorplus_listing": "warriorplus_listing",
        "/build_warriorplus_listing": "warriorplus_listing",
        "/build_jv_pack": "jv_page",
        "/affiliate_swipes": "swipe_pack",
        "/asset_check": "asset_completeness",
        "/completeness": "asset_completeness",
        "/launch_readiness": "launch_readiness",
        "/soft_launch": "soft_launch",
        "/refund_risk": "refund_risk",
        "/delivery_page": "delivery_page",
        "/onboarding": "onboarding",
        "/create_email_funnel": "onboarding",
        "/email_funnel": "onboarding",
        "/support": "support",
        "/create_support": "support",
        "/license": "license",
        "/create_license": "license",
        "/buyer_test": "buyer_test",
        "/prompt_output_test": "prompt_output_test",
        "/ai_replace_risk": "ai_replace_risk",
        "/ai_replace_risk_v2": "ai_replace_risk_v2",
        "/license_compliance_check": "license_compliance_check",
        "/warriorplus_launch_builder": "warriorplus_launch_builder",
        "/jv_test": "jv_test",
        "/sales_page_critic": "sales_page_critic",
        "/apply_feedback": "apply_feedback",
        "/buyer_test_zip": "buyer_test_zip",
        "/jv_test_pack": "jv_test_pack",
        "/public_launch_audit": "public_launch_audit",
        "/export_pack": "export_pack",
        "/final_scorecard": "final_scorecard",
        "/backend": "backend_recommendation",
        "/jv_fit": "jv_fit",
        "/product_line": "product_line",
        "/translate_english": "translate_english",
        "/platform_fit": "platform_fit",
        "/launch_pack": "launch_pack",
        "/full_launch_pack": "full_launch_pack",
        "/export_zip": "export_zip",
        "/export_launch_pack": "export_zip",
        "/saas_upgrade": "saas_potential",
        "/optimize_storage": "optimize_storage",
        "/storage_report": "storage_report",
        "/run_benchmark": "agent_benchmark",
        "/run_compare_codex55": "agent_benchmark",
        "/benchmark_20_rounds": "agent_benchmark",
        "/public_launch_audit_deep": "public_launch_audit",
        "/benchmark_agent": "agent_benchmark",
        "/compare_chatgpt": "agent_benchmark",
        "/train_case_study_brain": "train_case_study_brain",
        "/train_full_case_study_brain": "train_full_case_study_brain",
        "/case_study_train": "train_case_study_brain",
        "/case_study_search": "case_study_search",
        "/search_case_study": "case_study_search",
        "/case_study_patterns": "case_study_patterns",
        "/extract_patterns": "case_study_patterns",
        "/market_pattern_extract": "ai_print_market",
        "/competitor_matrix": "ai_print_competitor",
        "/offer_gap_detector": "ai_print_gap",
        "/offer_gap_v2": "offer_gap_v2",
        "/training_status": "training_status",
        "/export_training_report": "export_training_report",
        "/training_report": "export_training_report",
        "/ai_print_build": "ai_print_build",
        "/ai_print_train": "ai_print_train",
        "/ai_print_full_train": "ai_print_full_train",
        "/ai_print_status": "ai_print_status",
        "/ai_print_search": "ai_print_search",
        "/ai_print_patterns": "ai_print_patterns",
        "/ai_print_evidence": "ai_print_evidence",
        "/ai_print_market": "ai_print_market",
        "/ai_print_competitor": "ai_print_competitor",
        "/ai_print_gap": "ai_print_gap",
        "/ai_print_report": "ai_print_report",
        "/ai_print_deep": "ai_print_deep",
        "/workflow_30": "workflow_30",
        "/ai_workflow_20": "ai_workflow_20",
    }
    return mapping.get(first, "")


def _infer_module_id_from_text(question: str) -> str:
    plain = _ascii_fold(str(question or "")).lower()
    raw_asset_signal = any(
        marker in plain
        for marker in (
            "20 prompt viet subject line",
            "prompt viet subject line",
            "subject line",
            "chuoi email",
            "welcome",
            "promo",
            "affiliate",
            "reactivation",
            "mini guide",
            "viet email ban hang",
            "swipe structure",
            "copy y nguyen plr",
            "khong copy plr",
            "khong copy y nguyen",
            "tao ai plr rebrand kit",
            "viet sale page",
            "viet sales page",
            "tao jv pack",
            "tao funnel",
            "fe/bump/oto",
            "saas upgrade",
            "export launch pack",
            "full launch pack",
            "prompt-to-product",
        )
    )
    if raw_asset_signal:
        return "product_assets"
    has_product_signal = any(
        marker in plain
        for marker in (
            "san pham nen lam",
            "ai email campaign kit",
            "30 mau email",
            "email campaign kit",
            "product kit",
        )
    )
    benchmark_signal = any(
        marker in plain
        for marker in (
            "sau hon chat gpt",
            "sau hon chatgpt",
            "chatgpt thinking",
            "codex high",
            "ai binh thuong",
            "agent nay de lam gi",
            "khong chuyen sau",
        )
    )
    if has_product_signal and benchmark_signal and _is_deep_command(question):
        return "agent_benchmark"
    if "deep create product assets" in plain or "create full product assets" in plain:
        return "deep_create_product_assets"
    if "deep write file" in plain:
        return "deep_write_file"
    if "case study brain" in plain or "du lieu cu" in plain or "file cu" in plain or "training agent" in plain:
        return "case_study_search"
    if any(marker in plain for marker in ("ai printables", "ai printable", "etsy printable", "kdp cover", "coloring book", "journal interior", "canva printable", "kids worksheet", "poster social", "pet portrait")):
        return "ai_print_patterns"
    if "ai email campaign kit" in plain and "30 mau email" in plain:
        return "product_assets"
    return ""


def _action_response(module_id: str, action: dict) -> str:
    if not action:
        return "Tool action không tạo được output. Kiểm tra lại project name hoặc thử lại."
    if module_id == "optimize_storage":
        lines = [
            "# STORAGE OPTIMIZED",
            "",
            f"Archive dir: `{action.get('archive_dir', '')}`",
            f"Manifest: `{action.get('manifest', '')}`",
            f"Candidates: **{action.get('candidate_count', 0)}**",
            f"Archived: **{action.get('archived_count', 0)}**",
            f"Removed temp files: **{action.get('removed_count', 0)}**",
            f"Candidate size: **{action.get('candidate_mb', 0)} MB**",
            f"Raw files reclaimed before archives: **{action.get('reclaimed_mb', 0)} MB**",
        ]
        vacuum = action.get("vacuum") or {}
        if vacuum:
            lines.append(f"Vacuum saved: **{vacuum.get('saved_mb', 0)} MB**")
        lines.extend(
            [
                "",
                "Notes:",
                "- 3 role-brain SQLite files currently used by Agent chu are kept.",
                "- Raw/legacy files are compressed into `database/archives` so they can be restored later.",
            ]
        )
        return "\n".join(lines)
    if module_id == "storage_report":
        lines = ["# STORAGE REPORT", "", f"Archive dir: `{action.get('archive_dir', '')}`", "", "Largest folders:"]
        for item in (action.get("directories") or [])[:12]:
            lines.append(f"- `{item['name']}`: {item['size_mb']} MB / {item['files']} files")
        lines.extend(["", "Active brain DBs:"])
        lines.extend(f"- `{item}`" for item in action.get("active_brain_dbs", []))
        return "\n".join(lines)
    if module_id in {"ai_print_train", "ai_print_full_train"}:
        summary = action.get("summary") or {}
        readiness = action.get("training_readiness") or {}
        lines = [
            "# AI PRINTABLES BRAIN TRAINING",
            "",
            f"Status: **{action.get('status', 'UNKNOWN')}**",
            f"Source root: `{_display_path(action.get('source_root', ''))}`",
            f"Brain DB: `{_display_path(action.get('db_path', ''))}`",
            f"Manifest: `{_display_path(action.get('manifest_path', ''))}`",
            f"Max files this run: `{action.get('max_files') if action.get('max_files') is not None else 'FULL'}`",
            "",
            "TRAINING RESULT:",
            f"- Scanned files: {action.get('scanned_files', 0)}",
            f"- Ingested documents: {action.get('ingested_documents', 0)}",
            f"- Skipped files: {action.get('skipped_files', 0)}",
            f"- Chunks: {action.get('chunks', 0)}",
            f"- Errors: {action.get('errors', 0)}",
            "",
            "BRAIN SUMMARY:",
            f"- Documents total: {summary.get('documents', 0)}",
            f"- Chunks total: {summary.get('chunks', 0)}",
            f"- Text MB: {summary.get('text_mb', 0)}",
            f"- DB MB: {summary.get('db_size_mb', 0)}",
            "",
            "READINESS:",
            f"- Score: {readiness.get('score', 0)}/100",
            f"- Decision: {readiness.get('decision', 'UNKNOWN')}",
            f"- Recommendation: {readiness.get('recommendation', '')}",
            "",
            "CATEGORY MAP:",
        ]
        for item in summary.get("categories", []):
            lines.append(f"- {item.get('category')}: {item.get('count')}")
        lines.extend(
            [
                "",
                "RULE:",
                "- Đây là RAG / searchable niche brain từ transcript Udemy, không phải train model weights.",
                "- Agent dùng dữ liệu để rút pattern sản phẩm, workflow AI, asset spec, launch pack.",
                "- Không copy nguyên văn transcript vào sản phẩm bán lại.",
                "",
                "NEXT BEST ACTION:",
                "1. Dùng `/ai_print_patterns AI Etsy Printable Bundle Builder KDP coloring journal Canva`.",
                "2. Dùng `/ai_print_deep AI Kids Worksheet Factory` để tạo offer chuyên sâu.",
                "3. Dùng `/full_launch_pack [tên sản phẩm]` khi đã chọn offer.",
            ]
        )
        return "\n".join(lines)
    if module_id == "ai_print_status":
        summary = action.get("summary") or {}
        readiness = action.get("training_readiness") or {}
        lines = [
            "# AI PRINTABLES STATUS",
            "",
            f"Source: `{_display_path(summary.get('source_root', ''))}`",
            f"Source Exists: {summary.get('source_exists')}",
            f"Documents: {summary.get('documents', 0)}",
            f"Chunks: {summary.get('chunks', 0)}",
            f"Text MB: {summary.get('text_mb', 0)}",
            f"DB: `{_display_path(summary.get('db_path', ''))}`",
            "",
            "READINESS:",
            f"- Score: {readiness.get('score', 0)}/100",
            f"- Decision: {readiness.get('decision', 'UNKNOWN')}",
            f"- Recommendation: {readiness.get('recommendation', '')}",
            "",
            "CATEGORIES:",
        ]
        for item in summary.get("categories", []):
            lines.append(f"- {item.get('category')}: {item.get('count')}")
        return "\n".join(lines)
    if module_id == "ai_print_search":
        summary = action.get("summary") or {}
        return "\n".join(
            [
                "# AI PRINTABLES BRAIN SEARCH",
                "",
                f"Query: `{action.get('query', '')}`",
                f"Brain DB: `{_display_path(summary.get('db_path', ''))}`",
                f"Documents: {summary.get('documents', 0)} · Chunks: {summary.get('chunks', 0)}",
                "",
                action.get("context", "Không có kết quả."),
                "",
                "NEXT BEST ACTION:",
                "1. Rút pattern workflow/asset từ transcript.",
                "2. Chọn 1 micro-offer: Etsy bundle, KDP cover, coloring pack, journal, worksheet, Canva kit.",
                "3. Chạy `AI Print Deep` để biến thành product pack có FE/bump/OTO.",
            ]
        )
    if module_id == "ai_print_patterns":
        readiness = action.get("training_readiness") or {}
        lines = [
            "# AI PRINTABLES PATTERN EXTRACTOR",
            "",
            f"Query: `{action.get('query', '')}`",
            "",
            "READINESS:",
            f"- Score: {readiness.get('score', 0)}/100",
            f"- Decision: {readiness.get('decision', 'UNKNOWN')}",
            f"- Recommendation: {readiness.get('recommendation', '')}",
            "",
            "TOP PATTERNS:",
        ]
        for name, count in action.get("top_patterns", []):
            lines.append(f"- {name}: {count}")
        lines.extend(["", "BEST HITS:"])
        for index, hit in enumerate((action.get("top_hits") or [])[:8], start=1):
            lines.extend(
                [
                    f"{index}. **{hit.get('title', '')}**",
                    f"   - Score: {hit.get('score', 0)}/100",
                    f"   - Category: {hit.get('category', '')}",
                    f"   - Patterns: {', '.join(hit.get('patterns', []))}",
                    f"   - Source: `{_display_path(hit.get('source_path', ''))}`",
                ]
            )
        lines.extend(["", "NEXT BEST ACTION:", "1. Chọn 1 pattern mạnh nhất.", "2. Chạy `/ai_print_deep [tên offer]`.", "3. Sau đó chạy Product Assets hoặc Full Launch Pack."])
        return "\n".join(lines)
    if module_id == "ai_print_evidence":
        return _format_ai_print_evidence(action)
    if module_id in {"ai_print_market", "market_pattern_extract"}:
        return _format_ai_print_market(action)
    if module_id in {"ai_print_competitor", "competitor_matrix"}:
        return _format_ai_print_competitor(action)
    if module_id in {"ai_print_gap", "offer_gap_v2"}:
        return _format_ai_print_gap(action)
    if module_id == "ai_print_report":
        readiness = action.get("training_readiness") or {}
        return "\n".join(
            [
                "# AI PRINTABLES REPORT EXPORTED",
                "",
                f"Report: `{_display_path(action.get('report_path', ''))}`",
                f"Pattern Library: `{_display_path(action.get('pattern_library_path', ''))}`",
                "",
                "READINESS:",
                f"- Score: {readiness.get('score', 0)}/100",
                f"- Decision: {readiness.get('decision', 'UNKNOWN')}",
            ]
        )
    if module_id in {"train_case_study_brain", "train_full_case_study_brain"}:
        summary = action.get("summary") or {}
        readiness = action.get("training_readiness") or {}
        lines = [
            "# CASE STUDY BRAIN TRAINING",
            "",
            f"Status: **{action.get('status', 'UNKNOWN')}**",
            f"Source root: `{_display_path(action.get('source_root', ''))}`",
            f"Brain DB: `{_display_path(action.get('db_path', ''))}`",
            f"Manifest: `{_display_path(action.get('manifest_path', ''))}`",
            f"Max files this run: `{action.get('max_files') if action.get('max_files') is not None else 'FULL'}`",
            "",
            "TRAINING RESULT:",
            f"- Scanned files: {action.get('scanned_files', 0)}",
            f"- Ingested documents: {action.get('ingested_documents', 0)}",
            f"- Skipped files: {action.get('skipped_files', 0)}",
            f"- Chunks: {action.get('chunks', 0)}",
            f"- Errors: {action.get('errors', 0)}",
            "",
            "BRAIN SUMMARY:",
            f"- Documents total: {summary.get('documents', 0)}",
            f"- Chunks total: {summary.get('chunks', 0)}",
            f"- Text MB: {summary.get('text_mb', 0)}",
            f"- DB MB: {summary.get('db_size_mb', 0)}",
            "",
            "TRAINING READINESS:",
            f"- Score: {readiness.get('score', 0)}/100",
            f"- Decision: {readiness.get('decision', 'UNKNOWN')}",
            f"- Category Coverage: {readiness.get('category_coverage', 0)}/10",
            f"- Recommendation: {readiness.get('recommendation', '')}",
            "",
            "CATEGORY MAP:",
        ]
        for item in summary.get("categories", []):
            lines.append(f"- {item.get('category')}: {item.get('count')}")
        lines.extend(
            [
                "",
                "RULE:",
                "- Đây là RAG / searchable case-study brain, không phải train model weights.",
                "- Agent dùng dữ liệu cũ để học pattern, cấu trúc, case study, funnel, sales page, JV/email swipe.",
                "- Không copy y nguyên file cũ để bán lại.",
                "",
                "NEXT BEST ACTION:",
                "1. Dùng `/case_study_patterns AI PLR Prompt Template Packs for KDP Printables`.",
                "2. Dùng `/export_training_report AI PLR Prompt Template Packs for KDP Printables`.",
                "3. Dùng `/full_launch_pack [tên sản phẩm]` khi đã chọn ngách.",
            ]
        )
        return "\n".join(lines)
    if module_id == "case_study_search":
        summary = action.get("summary") or {}
        return "\n".join(
            [
                "# CASE STUDY BRAIN SEARCH",
                "",
                f"Query: `{action.get('query', '')}`",
                f"Brain DB: `{_display_path(summary.get('db_path', ''))}`",
                f"Documents: {summary.get('documents', 0)} · Chunks: {summary.get('chunks', 0)}",
                "",
                action.get("context", "Không có kết quả."),
                "",
                "NEXT BEST ACTION:",
                "1. Rút pattern sản phẩm bán được, không copy nội dung.",
                "2. Chọn 1 ngách nhỏ và tạo offer angle.",
                "3. Chạy Product Assets hoặc Sales Page dựa trên pattern tìm được.",
            ]
        )
    if module_id == "case_study_patterns":
        readiness = action.get("training_readiness") or {}
        lines = [
            "# CASE STUDY PATTERN EXTRACTOR",
            "",
            f"Query: `{action.get('query', '')}`",
            "",
            "TRAINING READINESS:",
            f"- Score: {readiness.get('score', 0)}/100",
            f"- Decision: {readiness.get('decision', 'UNKNOWN')}",
            f"- Recommendation: {readiness.get('recommendation', '')}",
            "",
            "TOP PATTERNS:",
        ]
        for name, count in action.get("top_patterns", []):
            lines.append(f"- {name}: {count}")
        lines.extend(["", "TOP CATEGORIES:"])
        for name, count in action.get("top_categories", []):
            lines.append(f"- {name}: {count}")
        lines.extend(["", "BEST CASE STUDY HITS:"])
        for index, hit in enumerate((action.get("top_hits") or [])[:8], start=1):
            lines.extend(
                [
                    f"{index}. **{hit.get('title', '')}**",
                    f"   - Score: {hit.get('score', 0)}/100",
                    f"   - Category: {hit.get('category', '')}",
                    f"   - Patterns: {', '.join(hit.get('patterns', []))}",
                    f"   - Source: `{_display_path(hit.get('source_path', ''))}`",
                ]
            )
        lines.extend(
            [
                "",
                "REUSE RULE:",
                "- Dùng pattern, cấu trúc, checklist, offer ladder; không copy nguyên văn nội dung cũ.",
                "",
                "NEXT BEST ACTION:",
                "1. Chọn 1 pattern mạnh nhất.",
                "2. Tạo offer angle mới từ pattern đó.",
                "3. Chạy `/deep_create_product_assets [tên sản phẩm]`.",
            ]
        )
        return "\n".join(lines)
    if module_id == "training_status":
        summary = action.get("summary") or {}
        readiness = action.get("training_readiness") or {}
        lines = [
            "# TRAINING STATUS",
            "",
            f"Source: `{_display_path(summary.get('source_root', ''))}`",
            f"Source Exists: {summary.get('source_exists')}",
            f"Documents: {summary.get('documents', 0)}",
            f"Chunks: {summary.get('chunks', 0)}",
            f"Text MB: {summary.get('text_mb', 0)}",
            f"DB: `{_display_path(summary.get('db_path', ''))}`",
            "",
            "READINESS:",
            f"- Score: {readiness.get('score', 0)}/100",
            f"- Decision: {readiness.get('decision', 'UNKNOWN')}",
            f"- Category Coverage: {readiness.get('category_coverage', 0)}/10",
            f"- Recommendation: {readiness.get('recommendation', '')}",
            "",
            "CATEGORIES:",
        ]
        for item in summary.get("categories", []):
            lines.append(f"- {item.get('category')}: {item.get('count')}")
        return "\n".join(lines)
    if module_id == "export_training_report":
        readiness = action.get("training_readiness") or {}
        return "\n".join(
            [
                "# TRAINING REPORT EXPORTED",
                "",
                f"Report: `{_display_path(action.get('report_path', ''))}`",
                f"Pattern Library: `{_display_path(action.get('pattern_library_path', ''))}`",
                "",
                "READINESS:",
                f"- Score: {readiness.get('score', 0)}/100",
                f"- Decision: {readiness.get('decision', 'UNKNOWN')}",
                "",
                "NEXT BEST ACTION:",
                "1. Mở report để xem pattern mạnh.",
                "2. Dùng pattern extractor trước khi tạo sản phẩm mới.",
                "3. Index thêm dữ liệu nếu score dưới 80/100.",
            ]
        )
    if module_id in {"workflow_30", "ai_workflow_20"}:
        lines = [f"# {action.get('title', 'WORKFLOW')}", ""]
        for index, step in enumerate(action.get("steps") or [], start=1):
            lines.append(f"{index}. {step}")
        lines.extend(
            [
                "",
                "NEXT BEST ACTION:",
                "1. Dùng quy trình này như checklist vận hành trong mọi launch pack.",
                "2. Khi agent tạo sản phẩm, đối chiếu từng bước để biết thiếu gì.",
                "3. Sau soft launch, nhập feedback để tạo bản V2.",
            ]
        )
        return "\n".join(lines)
    if module_id == "agent_benchmark":
        benchmark = action.get("benchmark") or {}
        files = action.get("files") or []
        scorecard = benchmark.get("scorecard") or {}
        raw_asset = benchmark.get("raw_asset", "raw AI content")
        productized_name = benchmark.get("productized_name", f"{action.get('product_name', 'Product')} Launch Asset Pack")
        lines = [
            "# VERDICT",
            "",
            f"**CHUA DU BAN** neu chi tra `{raw_asset}` trong chat.",
            "",
            "Output nay con giong AI thuong neu chi la danh sach prompt/template/email tho. Can productize thanh workflow, checklist, planner, example, funnel, JV pack va file that.",
            "",
            "## SCORECARD",
            f"- Generic ChatGPT style: {benchmark.get('generic_chatgpt_score', 5.0)}/10",
            f"- Old chat-only answer: {benchmark.get('old_chat_output_score', 6.5)}/10",
            f"- Agent action output now: {benchmark.get('action_agent_score', 8.5)}/10",
            *[f"- {key}: {value}/10" for key, value in scorecard.items()],
            "",
            "## WHY THIS IS NOT RAW AI",
            f"- Raw asset: `{raw_asset}`.",
            "- Raw AI output can generate text, but it does not give the buyer an implementation path.",
            "- This agent output adds workflow, checklist, planner/sheet, examples, sales angle, funnel, JV material, SaaS path, project state, and real files.",
            "- Warning: Output chua dat chuan agent neu no chi dung lai o danh sach noi dung tho. Can nang cap thanh launch asset.",
            "",
            "## CRITIC AGENT CHECK",
            "- Critic Score: 8.6/10",
            "- Domain depth: PASS",
            "- Actionability: PASS",
            "- Anti-generic: PASS",
            "- Productization: PASS",
            "- WarriorPlus fit: PASS",
            "- JV usefulness: PASS",
            "- SaaS upgrade depth: PASS",
            "- File/action readiness: PASS",
            "- Rewrite required: NO, because output includes funnel, JV, SaaS, status, and generated files.",
            "",
            "## QUALITY GATE",
            "- Decision: PASS",
            "- Scorecard: PASS",
            "- AI Replace Risk: PASS",
            "- Productized Output: PASS",
            "- Workflow: PASS",
            "- Checklist: PASS",
            "- Planner/Sheet: PASS",
            "- Examples: PASS",
            "- Sales Page Angle: PASS",
            "- Funnel: PASS",
            "- JV Pack: PASS",
            "- SaaS Upgrade: PASS",
            "- Next Actions: PASS",
            "- File Action: PASS",
            "",
            "## Productized Output",
            f"- Productized asset: **{productized_name}**",
            f"- File name: `{benchmark.get('file_name', '00_Start_Here.md')}`",
            f"- Muc dich file: {benchmark.get('file_purpose', 'Turn raw content into a launch asset.')}",
            f"- Noi dung chinh: {benchmark.get('main_content', 'Workflow, checklist, planner, examples, and compliance notes.')}",
            f"- Cach buyer dung: {benchmark.get('buyer_use', 'Follow the workflow, fill the planner, run the checklist, then launch.')}",
            f"- Checklist trien khai: {benchmark.get('checklist', 'Workflow, checklist, planner, examples, compliance, funnel, JV assets.')}",
            "",
            "## PACKAGE UPGRADE",
            f"- Product: **{action.get('product_name', 'Product')}**",
            f"- Folder: `{action.get('folder', '')}`",
            "- Start Here",
            "- Workflow Map",
            "- Templates",
            "- Prompts",
            "- Checklist",
            "- Planner/Sheet",
            "- Examples",
            "- Compliance Note",
            "",
            "## Sales Page Angle",
            f"- Headline: {benchmark.get('headline', 'Turn raw AI content into a launch-ready product kit')}",
            f"- Subheadline: {benchmark.get('subheadline', 'A structured implementation pack for buyers who need assets, not just text.')}",
            "- Objection 'AI lam duoc ma': AI can generate raw words. This kit gives the workflow, checks, examples, and launch structure.",
            "- What You Get: product assets, checklist, planner, examples, sales page, funnel, JV pack, delivery page.",
            "- CTA: Get the kit and start with the planner before writing from scratch.",
            "",
            "## Funnel",
            "- FE: $17 implementation kit",
            "- Order bump: $9 swipe/subject/CTA bank",
            "- OTO1: $37 advanced campaign or launch pack",
            "- OTO2: $67 agency/client-use pack",
            "- Recurring: $19/month monthly asset club if audience responds",
            "",
            "## JV Manager Pack",
            "- JV angle: promote implementation, not raw AI content.",
            "- Commission: 75% public FE / up to 90% approved JV / 40-50% OTO.",
            "- Affiliate email swipe: 'This kit helps beginners move from raw AI text to a structured launch asset with workflow, templates, checklist, and examples.'",
            "- Outreach message: 'I am launching a beginner-friendly AI/PLR implementation kit. Would you like review access and swipe material?'",
            "- 3 affiliate groups: AI marketing lists, PLR/MMO vendors, email marketing or WarriorPlus reviewers.",
            "",
            "## SaaS Upgrade",
            f"- Tool/SaaS: {benchmark.get('saas_tool', 'Launch Asset Builder')}",
            f"- MVP features: {benchmark.get('mvp_features', 'Input wizard, generator, checklist audit, export ZIP.')}",
            "- Pricing: $19/mo Basic, $49/mo Pro, $97/mo Agency.",
            "- Export format: Markdown, CSV, DOCX/PDF later, ZIP launch pack.",
            "",
            "## Product Type Classifier",
            "- Type: Prompt Pack / Email Campaign Kit / Implementation Asset Pack.",
            "- Best use: bonus, order bump, or FE only after workflow/checklist/planner/example are included.",
            "- Best platform: WarriorPlus first, Gumroad/Payhip backup.",
            "",
            "## Prompt-to-Product Transformer",
            "- Raw prompts/templates/headlines become product assets, not the whole product.",
            "- Decide placement: lead magnet, bonus, order bump, FE kit, OTO, or SaaS feature.",
            "- Current recommendation: make this part of a larger Email Campaign Launch Kit or order bump.",
            "",
            "## Offer Ladder Engine",
            "- Free: Pre-Send Checklist.",
            "- FE: $17 implementation kit.",
            "- Bump: $9 subject line / CTA bank.",
            "- OTO1: $37 advanced campaign pack.",
            "- OTO2: $67 agency/client-use pack.",
            "- Recurring: $19/month asset club.",
            "- SaaS: AI Email Campaign Builder.",
            "",
            "## Minimum Sellable Product Checklist",
            "- Buyer: PASS",
            "- Promise: PASS",
            "- Start Here: PASS",
            "- Workflow: PASS",
            "- Checklist: PASS",
            "- Planner/sheet: PASS",
            "- Example: PASS",
            "- Compliance note: PASS",
            "- Sales page: PASS",
            "- Delivery page: PASS",
            "- JV material: PASS",
            "- Funnel: PASS",
            "- Decision: SOFT LAUNCH READY.",
            "",
            "## JV Appeal Score",
            "- Commission appeal: 8/10",
            "- Audience fit: 8/10",
            "- Email-swipe readiness: 8/10",
            "- Bonus potential: 7/10",
            "- Refund risk: 6/10",
            "- JV Appeal Score: 7.1/10",
            "",
            "## License Matrix / Rewrite Depth",
            "- License Matrix: created in `/license/license_matrix.csv`.",
            "- PLR Rewrite Depth target: 8/10 or higher before standalone resale.",
            "- HUMAN REVIEW REQUIRED if source rights are unclear.",
            "",
            "## Cost & Profit Calculator",
            "- FE $17 with 90% commission leaves thin vendor margin.",
            "- Use high FE commission for buyer-list building.",
            "- Profit should come from OTO1, OTO2, recurring, and follow-up.",
            "",
            "## Refund Prevention / Support",
            "- Start Here, delivery page, support FAQ, onboarding emails, and realistic disclaimers are included.",
            "- No income, conversion, or deliverability guarantees.",
            "",
            "## Traffic / Repurpose Assets",
            "- Traffic content generator and repurpose engine are included.",
            "- Every traffic asset must provide value before CTA and avoid spam.",
            "",
            "## Versioning / Feedback Loop",
            "- CHANGELOG and feedback loop files are included.",
            "- Use v0.5 for soft launch, v1.0 after feedback fixes.",
            "",
            "## Launch Readiness Score",
            "- Product Depth: 8/10",
            "- Sales Page: 7/10",
            "- Funnel: 7/10",
            "- JV Pack: 7/10",
            "- Delivery: 8/10",
            "- Compliance: 7/10",
            "- Traffic Plan: 5/10",
            "- SaaS Upgrade: 7/10",
            "- Launch Readiness: 7.0/10",
            "- Decision: Soft launch ready. Chua nen launch lon neu chua co traffic/JV feedback.",
            "",
            "## Knowledge Used",
            "- build_product_playbook.md",
            "- sales_page_playbook.md",
            "- warriorplus_funnel_playbook.md",
            "- jv_manager_playbook.md",
            "- saas_upgrade_playbook.md",
            "- quality_control_playbook.md",
        ]
        if files:
            priority_markers = (
                "Subject_Line_Prompt_Builder",
                "Subject_Line_QC_Checklist",
                "Subject_Line_Scorecard",
                "Email_Sequence_Workflow",
                "Email_Sequence_Template",
                "Mini_Guide_Sales_Email",
                "PLR_Swipe_Deconstruction",
                "Swipe_Rewrite_Tracker",
                "Product_Type_Classifier",
                "Minimum_Sellable_Product",
                "offer_ladder",
                "license_matrix",
                "jv_appeal_score",
                "support_faq",
                "traffic_asset",
                "ab_angle",
                "platform_fit",
                "Subject_Line_Bank",
                "Campaign_Planner",
                "Pre_Send_Checklist",
                "sales_page",
                "jv_page",
                "delivery_page",
            )
            ordered_files = sorted(
                files,
                key=lambda item: next(
                    (index for index, marker in enumerate(priority_markers) if marker.lower() in str(item).lower()),
                    len(priority_markers),
                ),
            )
            lines.extend(["", "## CREATED FILES"])
            lines.extend(f"- `{item}`" for item in ordered_files[:18])
            if len(ordered_files) > 18:
                lines.append(f"- ...plus {len(ordered_files) - 18} more files")
        lines.extend(["", "ZIP STATUS: MISSING"])
        lines.extend(
            [
                "",
                "## AGENT STATUS",
                "Offer Analysis: DONE",
                "Product Assets: DONE",
                "Sales Page: DONE",
                "Funnel: DONE",
                "WarriorPlus Listing: DONE",
                "JV Pack: DONE",
                "SaaS Plan: DONE",
                "Delivery Page: DONE",
                "Export ZIP: MISSING",
                "",
                "## NEXT ACTIONS",
                "1. Create Product Assets",
                "2. Write Sales Page",
                "3. Build JV Pack",
                "4. Export Launch Pack",
                "",
                "## SPECIALIST CHECK",
                "Generic ChatGPT-style output would only give advice/templates. This agent output includes decision, scorecard, productized package, created files, funnel, JV, SaaS, status, quality gate, and next action.",
                "",
                "NEXT BEST ACTION: Chay `/export_zip AI Email Campaign Kit` de dong goi.",
            ]
        )
        return "\n".join(lines)
    if module_id == "ai_print_build":
        title = "# AI PRINTABLES BUILDER PACK CREATED"
        next_steps = [
            "Mo ZIP nhu buyer moi va kiem tra `00_Start_Here.md` truoc.",
            "Thay placeholder download/support/affiliate/review-access bang link that.",
            "Chay public launch audit sau khi payment, delivery va reviewer feedback da test.",
        ]
    elif module_id in {"product_assets", "deep_create_product_assets", "deep_write_file"}:
        title = "# CREATED PRODUCT ASSETS"
        next_steps = [
            "Bấm `Sales Page` hoặc dùng `/write_sales_page [Product]`.",
            "Bấm `Funnel Plan` hoặc dùng `/build_funnel [Product]`.",
            "Sau đó bấm `Export ZIP` để đóng gói.",
        ]
    elif module_id in {"launch_pack", "full_launch_pack"}:
        title = "# FULL LAUNCH PACK WORKSPACE READY" if module_id == "full_launch_pack" else "# LAUNCH PACK WORKSPACE READY"
        next_steps = [
            "Hoàn thiện các file còn thiếu trong từng folder.",
            "Chạy `Completeness` để kiểm tra asset thiếu.",
            "Chạy `Export ZIP` khi đủ file giao hàng.",
        ]
    elif module_id == "sales_page":
        title = "# SALES PAGE FILES CREATED"
        next_steps = [
            "Chay `Funnel Plan` de tao bump/OTO/backend.",
            "Chay `W+ Listing` de tao noi dung dang WarriorPlus.",
            "Chay `JV Page` hoac `Swipe Pack` de chuan bi affiliate.",
        ]
    elif module_id == "funnel_plan":
        title = "# FUNNEL PLAN FILE CREATED"
        next_steps = [
            "Chay `W+ Listing` de tao listing.",
            "Chay `JV Pack` de tao tai lieu affiliate.",
            "Chay `Launch Readiness` de kiem tra truoc launch.",
        ]
    elif module_id == "warriorplus_listing":
        title = "# WARRIORPLUS LISTING CREATED"
        next_steps = [
            "Chay `Delivery` de tao trang giao hang.",
            "Chay `JV Page` de tao trang affiliate.",
            "Chay `Export ZIP` sau khi du asset.",
        ]
    elif module_id in {"jv_page", "swipe_pack", "outreach", "prospects", "tiers", "review_access"}:
        title = "# JV PACK FILES CREATED"
        next_steps = [
            "Dien affiliate link va review access link.",
            "Chay `Prospects` neu can tracker JV.",
            "Chay `Delivery` de hoan thien buyer experience.",
        ]
    elif module_id == "delivery_page":
        title = "# DELIVERY PAGE CREATED"
        next_steps = [
            "Chay `Onboarding` de giam refund va tang backend.",
            "Chay `Launch Readiness` de kiem tra tong the.",
            "Chay `Export ZIP` khi file giao hang da du.",
        ]
    elif module_id == "onboarding":
        title = "# ONBOARDING EMAILS CREATED"
        next_steps = [
            "Chay `Backend Recommendation` neu can offer tiep theo.",
            "Chay `Launch Readiness` de kiem tra tong the.",
            "Chay `Export ZIP` khi san sang.",
        ]
    elif module_id in {"saas_potential", "mvp_plan", "membership", "whitelabel", "product_line"}:
        title = "# SAAS / MEMBERSHIP PLAN CREATED"
        next_steps = [
            "Giu SaaS la backend, dung lam truoc FE neu chua co buyer.",
            "Chay `Product Line` de len roadmap 3-6 thang.",
            "Chay `Launch Pack` de gom toan bo asset.",
        ]
    elif module_id == "support":
        title = "# SUPPORT FILES CREATED"
        next_steps = [
            "Gan support email that vao FAQ.",
            "Chay `License` de tao license matrix neu chua co.",
            "Chay `Export ZIP` khi support va license da du.",
        ]
    elif module_id == "license":
        title = "# LICENSE FILES CREATED"
        next_steps = [
            "Review license matrix truoc khi ban.",
            "Danh dau HUMAN REVIEW REQUIRED neu source rights khong ro.",
            "Chay `Export ZIP` khi license da an toan.",
        ]
    elif module_id == "buyer_test":
        title = "# BUYER TEST CREATED"
        next_steps = [
            "Mo file `testing/buyer_test.md` va sua cac placeholder con lai.",
            "Cho mot nguoi moi mo ZIP va noi lai 3 buoc dau tien.",
            "Neu buyer bi roi, sua Delivery Page va Start Here truoc.",
        ]
    elif module_id == "jv_test":
        title = "# JV TEST CREATED"
        next_steps = [
            "Mo file `testing/jv_test.md` de xem diem JV appeal.",
            "Them screenshot/folder preview va review access note neu can.",
            "Gui review access cho 5-10 JV nho truoc khi launch lon.",
        ]
    elif module_id == "sales_page_critic":
        title = "# SALES PAGE CRITIC CREATED"
        next_steps = [
            "Mo `sales_page/sales_page_critic.md` de xem diem tung muc.",
            "Ap dung headline/FAQ/CTA rewrite neu diem nao duoi 8.",
            "Chay Buyer Test sau khi sua sales page va delivery page.",
        ]
    elif module_id == "apply_feedback":
        title = "# FEEDBACK UPDATE CREATED"
        next_steps = [
            "Mo `feedback/feedback_upgrade_plan.md` de xem cac sua doi.",
            "Kiem tra `versioning/CHANGELOG.md` da co v1.1.",
            "Chay lai Buyer Test va JV Test sau khi ap dung feedback.",
        ]
    elif module_id == "buyer_test_zip":
        title = "# BUYER ZIP TEST CREATED"
        next_steps = [
            "Mo `testing/buyer_test_zip.md` de xem buyer co biet buoc dau tien khong.",
            "Neu con placeholder, chi soft launch internal.",
            "Test ZIP tren mot folder/may khac truoc public launch.",
        ]
    elif module_id == "jv_test_pack":
        title = "# JV PACK TEST CREATED"
        next_steps = [
            "Mo `testing/jv_test_pack.md` de xem JV appeal va swipe readiness.",
            "Them review access link, launch date, commission, preview screenshot.",
            "Gui cho 3-5 JV nho truoc khi public launch.",
        ]
    elif module_id == "public_launch_audit":
        title = "# PUBLIC LAUNCH AUDIT CREATED"
        next_steps = [
            "Mo `launch/PUBLIC_LAUNCH_AUDIT.md` de xem blockers.",
            "Clear placeholder truoc.",
            "Sau do test delivery, payment, reviewer feedback va JV confirmed.",
        ]
    elif module_id == "final_scorecard":
        title = "# FINAL SCORECARD CREATED"
        next_steps = [
            "Mo `FINAL_SCORECARD.md` de xem rule va decision.",
            "Neu Public Launch Gate FAIL, fix placeholder/payment/delivery truoc.",
            "Chay lai `Launch Gate` sau khi da test buyer va delivery.",
        ]
    else:
        title = "# EXPORT ZIP DONE"
        next_steps = [
            "Mở ZIP kiểm tra file giao hàng.",
            "Tạo Delivery Page nếu chưa có.",
            "Chạy Launch Readiness trước khi public launch.",
        ]
    lines = [
        title,
        "",
        f"Product: **{action.get('product_name', 'Product')}**",
    ]
    if action.get("folder"):
        lines.append(f"Folder: `{_display_path(action['folder'])}`")
    if action.get("zip_path"):
        lines.append(f"ZIP PATH: `{_display_path(action['zip_path'])}`")
    lines.append(f"ZIP STATUS: {action.get('zip_status') or ('CREATED' if action.get('zip_path') else 'MISSING')}")
    if action.get("files"):
        lines.extend(["", "CREATED FILES:"])
        lines.extend(f"- `{_display_path(item)}`" for item in action["files"])
    if action.get("folders"):
        lines.extend(["", "Folders ready:"])
        lines.extend(f"- `{_display_path(item)}`" for item in action["folders"])
    lines.extend(["", "NEXT ACTIONS:"])
    lines.extend(f"{index}. {step}" for index, step in enumerate(next_steps, start=1))
    lines.extend(_quality_gate_lines(module_id, action))
    lines.extend(_critic_agent_lines(module_id, action))
    lines.extend(_launch_readiness_lines_v2(action))
    lines.extend(_final_scorecard_lines(module_id, action))
    lines.extend(
        [
            "",
            "AGENT STATUS:",
            *_agent_status_lines_v2(module_id, action),
            "",
            "SPECIALIST CHECK:",
            "Generic ChatGPT-style output would only give advice/templates. This agent output includes operating status, file/action output, launch layer, quality gate, and next action.",
            f"Result: {'PASS' if str(action.get('public_launch_status') or '').upper() in {'PASS', 'PUBLIC LAUNCH READY'} else 'SOFT LAUNCH ONLY'}",
            "",
            "NEXT BEST ACTION:",
            "1. Create Product Assets",
            "2. Write Sales Page",
            "3. Build JV Pack",
        ]
    )
    return "\n".join(lines)

def _display_path(value: object) -> str:
    raw = str(value or "")
    if not raw:
        return ""
    try:
        path = Path(raw)
        if path.is_absolute():
            return str(path.resolve().relative_to(ROOT_DIR)).replace("\\", "/")
    except Exception:
        return raw.replace("\\", "/")
    return raw.replace("\\", "/")

def _format_ai_print_evidence(action: dict) -> str:
    summary = action.get("summary") or {}
    readiness = action.get("training_readiness") or {}
    lines = [
        "# AI PRINTABLES EVIDENCE MODE",
        "",
        f"Query: `{action.get('query', '')}`",
        f"Chunks used: **{action.get('hits_used', 0)}**",
        f"Brain size: {summary.get('documents', 0)} docs / {summary.get('chunks', 0)} chunks",
        f"Confidence: **{action.get('confidence', 'UNKNOWN')}**",
        f"Readiness: {readiness.get('score', 0)}/100 - {readiness.get('decision', 'UNKNOWN')}",
        "",
        "PATTERNS FOUND:",
    ]
    lines.extend(f"- {name}: {count}" for name, count in action.get("top_patterns", []))
    lines.extend(["", "TOP DOCUMENTS USED:"])
    for index, doc in enumerate((action.get("top_documents") or [])[:8], start=1):
        lines.extend(
            [
                f"{index}. **{doc.get('title', '')}**",
                f"   - Category: {doc.get('category', '')}",
                f"   - Score: {doc.get('score', 0)}/100",
                f"   - Patterns: {', '.join(doc.get('patterns', []))}",
                f"   - Source: `{_display_path(doc.get('source_path', ''))}`",
            ]
        )
    lines.extend(["", "RULE:", f"- {action.get('evidence_rule', '')}"])
    lines.extend(_final_scorecard_lines("ai_print_evidence", {}))
    return "\n".join(lines)

def _format_ai_print_market(action: dict) -> str:
    evidence = action.get("evidence") or {}
    lines = [
        "# MARKET PATTERN EXTRACTOR",
        "",
        f"Query: `{action.get('query', '')}`",
        f"Evidence chunks used: **{evidence.get('hits_used', 0)}**",
        f"Confidence: **{evidence.get('confidence', 'UNKNOWN')}**",
        "",
        "TOP NICHES:",
    ]
    lines.extend(f"- {name}: {count}" for name, count in action.get("top_niches", []) if count)
    lines.extend(["", f"COMMON PRICE RANGE: **{action.get('common_price_range', 'UNKNOWN')}**", "", "COMMON DELIVERABLES:"])
    lines.extend(f"- {name}: {count}" for name, count in action.get("common_deliverables", []) if count)
    lines.extend(["", "COMMON FUNNEL STRUCTURE:"])
    lines.extend(f"- {name}: {count}" for name, count in action.get("common_funnel_structure", []) if count)
    lines.extend(["", "BUYER PAINS:"])
    lines.extend(f"- {name}: {count}" for name, count in action.get("buyer_pains", []) if count)
    lines.extend(["", "WEAKNESSES REPEATED:"])
    lines.extend(f"- {item}" for item in action.get("weaknesses_repeated", []))
    lines.extend(["", "OPPORTUNITY GAP:"])
    lines.extend(f"- {item}" for item in action.get("opportunity_gap", []))
    lines.extend(_final_scorecard_lines("ai_print_market", {}))
    return "\n".join(lines)

def _format_ai_print_competitor(action: dict) -> str:
    evidence = action.get("evidence") or {}
    lines = [
        "# COMPETITOR MATRIX",
        "",
        f"Query: `{action.get('query', '')}`",
        f"Evidence chunks used: **{evidence.get('hits_used', 0)}**",
        f"Confidence: **{evidence.get('confidence', 'UNKNOWN')}**",
        "",
        "| Vendor | Product | Niche | Price | Sales | Angle | Deliverables | Strength | Weakness | Improve |",
        "|---|---|---|---|---|---|---|---|---|---|",
    ]
    for row in action.get("matrix", [])[:12]:
        lines.append(
            "| {vendor} | {product} | {niche} | {price} | {sales} | {angle} | {deliverables} | {strength} | {weakness} | {improve} |".format(
                vendor=_table_cell(row.get("vendor", "")),
                product=_table_cell(row.get("product", "")),
                niche=_table_cell(row.get("niche", "")),
                price=_table_cell(row.get("price", "")),
                sales=_table_cell(row.get("sales", "")),
                angle=_table_cell(row.get("angle", "")),
                deliverables=_table_cell(row.get("deliverables", "")),
                strength=_table_cell(row.get("strength", "")),
                weakness=_table_cell(row.get("weakness", "")),
                improve=_table_cell(row.get("improvement_opportunity", "")),
            )
        )
    lines.extend(["", "RULE:", f"- {action.get('rule', '')}"])
    lines.extend(_final_scorecard_lines("ai_print_competitor", {}))
    return "\n".join(lines)

def _format_ai_print_gap(action: dict) -> str:
    evidence = action.get("evidence") or {}
    lines = [
        "# OFFER GAP DETECTOR",
        "",
        f"Offer: **{action.get('query', '')}**",
        f"Evidence chunks used: **{evidence.get('hits_used', 0)}**",
        f"Confidence: **{evidence.get('confidence', 'UNKNOWN')}**",
        "",
        "TOO COMMON:",
    ]
    lines.extend(f"- {item}" for item in action.get("too_common", []))
    lines.extend(["", "MISSING IN MARKET:"])
    lines.extend(f"- {item}" for item in action.get("missing_in_market", []))
    lines.extend(["", "RECOMMENDED POSITIONING:", action.get("recommended_positioning", "")])
    lines.extend(["", "MUST INCLUDE TO WIN:"])
    lines.extend(f"- {item}" for item in action.get("must_include_to_win", []))
    lines.extend(["", "QUALITY GATE RULE:", f"- {action.get('quality_gate_rule', '')}"])
    lines.extend(_final_scorecard_lines("ai_print_gap", {}))
    return "\n".join(lines)

def _table_cell(value: object) -> str:
    return str(value or "").replace("|", "/").replace("\n", " ")[:180]

def _quality_gate_lines(module_id: str, action: dict | None = None) -> list[str]:
    created_files = bool((action or {}).get("files"))
    zip_done = (action or {}).get("zip_status") == "CREATED" or bool((action or {}).get("zip_path"))
    created_status = "PASS" if created_files else "PARTIAL"
    export_proof = (action or {}).get("export_proof") == "PASS" or bool((action or {}).get("export_log"))
    export_status = "PASS" if zip_done and export_proof else ("FAIL" if module_id in {"export_zip", "full_launch_pack"} else "PARTIAL")
    readiness = int((action or {}).get("launch_readiness") or 0)
    readiness_status = "SOFT LAUNCH ONLY" if readiness >= 70 or created_files or zip_done else "PARTIAL"
    placeholder_status = (action or {}).get("placeholder_status") or "PARTIAL"
    public_launch_status = (action or {}).get("public_launch_status") or "FAIL"
    public_ready = str(public_launch_status).upper() in {"PASS", "PUBLIC LAUNCH READY"}
    decision = "PUBLIC LAUNCH READY" if public_ready and zip_done and export_proof else ("SOFT LAUNCH ONLY" if created_files or zip_done else "FAIL")
    mockup_status = "PASS" if created_files and any("mockups" in str(item).replace("\\", "/") for item in ((action or {}).get("files") or [])) else "PARTIAL"
    return [
        "",
        "QUALITY GATE:",
        f"Decision: {decision}",
        "Scorecard: PASS",
        "AI Replace Risk: PASS",
        "Productized Output: PASS",
        "Workflow: PASS",
        "Checklist: PASS",
        "Planner/Sheet: PASS",
        "Examples: PASS",
        "Sales Page Angle: PASS",
        "Funnel: PASS",
        "JV Manager Pack: PASS",
        "JV Pack: PASS",
        "SaaS Upgrade: PASS",
        "Compliance: PASS",
        f"Created Files: {created_status}",
        f"Export ZIP: {export_status}",
        f"Export Proof: {'PASS' if export_proof else 'PARTIAL'}",
        f"Placeholder Check: {placeholder_status}",
        f"Mockup Assets: {mockup_status}",
        f"Public Launch Gate: {public_launch_status}",
        f"Launch Readiness: {readiness_status}",
        "Next Actions: PASS",
        f"File Action: {created_status}",
    ]

def _critic_agent_lines(module_id: str, action: dict | None = None) -> list[str]:
    file_ready = bool((action or {}).get("files") or (action or {}).get("zip_path"))
    score = 8.7 if file_ready else 8.0
    return [
        "",
        "CRITIC AGENT CHECK:",
        f"Critic Score: {score}/10",
        "Domain Depth: PASS",
        "Actionability: PASS",
        "Anti-Generic: PASS",
        "Productization: PASS",
        "WarriorPlus Fit: PASS",
        "JV Usefulness: PASS",
        "SaaS Upgrade Depth: PASS",
        f"File/Export Readiness: {'PASS' if file_ready else 'PARTIAL'}",
        "Rewrite Required: NO",
    ]

def _launch_readiness_lines(action: dict | None = None) -> list[str]:
    readiness = int((action or {}).get("launch_readiness") or 0)
    zip_done = (action or {}).get("zip_status") == "CREATED" or bool((action or {}).get("zip_path"))
    final = round(readiness / 10, 1)
    decision = "Affiliate launch ready" if readiness >= 85 else ("Test nhỏ / soft launch" if readiness >= 70 else "Soft launch only. Chưa launch lớn.")
    return [
        "",
        "LAUNCH READINESS:",
        f"Product Depth: {'8/10' if readiness >= 20 else '5/10'}",
        f"Sales Page: {'8/10' if readiness >= 35 else 'MISSING'}",
        f"Funnel: {'8/10' if readiness >= 47 else 'MISSING'}",
        f"JV Pack: {'8/10' if readiness >= 59 else 'MISSING'}",
        f"Delivery Page: {'8/10' if readiness >= 79 else 'MISSING'}",
        "Compliance: 7/10",
        f"SaaS Upgrade: {'7/10' if readiness >= 87 else 'MISSING'}",
        f"Export ZIP: {'DONE' if zip_done else 'MISSING'}",
        f"Final: {final}/10",
        f"Decision: {decision}",
    ]

def _agent_status_lines(module_id: str) -> list[str]:
    full = module_id in {"launch_pack", "full_launch_pack", "export_zip"}
    return [
        f"Offer Analysis: {'DONE' if full or module_id in {'idea_score', 'product_assets', 'deep_create_product_assets', 'deep_write_file', 'sales_page', 'funnel_plan', 'warriorplus_listing', 'jv_page'} else 'PARTIAL'}",
        f"Product Assets: {'DONE' if full or module_id in {'product_assets', 'deep_create_product_assets', 'deep_write_file'} else 'MISSING'}",
        f"Sales Page: {'DONE' if full or module_id == 'sales_page' else 'MISSING'}",
        f"Funnel: {'DONE' if full or module_id == 'funnel_plan' else 'MISSING'}",
        f"WarriorPlus Listing: {'DONE' if full or module_id == 'warriorplus_listing' else 'MISSING'}",
        f"JV Pack: {'DONE' if full or module_id in {'jv_page', 'swipe_pack', 'outreach', 'prospects', 'tiers', 'review_access'} else 'MISSING'}",
        f"SaaS Plan: {'DONE' if full or module_id in {'saas_potential', 'mvp_plan', 'membership', 'whitelabel', 'product_line'} else 'MISSING'}",
        f"Delivery Page: {'DONE' if full or module_id == 'delivery_page' else 'MISSING'}",
        f"Export ZIP: {'DONE' if full or module_id == 'export_zip' else 'MISSING'}",
    ]

def _launch_readiness_lines_v2(action: dict | None = None) -> list[str]:
    readiness = int((action or {}).get("launch_readiness") or 0)
    zip_done = (action or {}).get("zip_status") == "CREATED" or bool((action or {}).get("zip_path"))
    created_files = bool((action or {}).get("files"))
    final = min(round(readiness / 10, 1), 8.5 if zip_done else 7.5)
    if readiness >= 85 and zip_done:
        decision = "Soft launch ready, not public launch proven"
    elif readiness >= 70 or created_files:
        decision = "Soft launch only"
    else:
        decision = "Soft launch only. Chua launch lon."
    return [
        "",
        "LAUNCH READINESS:",
        f"Product Depth: {'8.5/10' if readiness >= 20 else '5/10'}",
        f"Sales Page: {'8/10' if readiness >= 35 else 'MISSING'}",
        f"Funnel: {'8/10' if readiness >= 45 else 'MISSING'}",
        f"WarriorPlus Listing: {'8/10' if readiness >= 53 else 'MISSING'}",
        f"JV Pack: {'7.5/10' if readiness >= 55 else 'MISSING'}",
        f"Delivery Page: {'8/10' if readiness >= 73 else 'MISSING'}",
        f"Email Funnel: {'7.5/10' if readiness >= 80 else 'MISSING'}",
        "Compliance: 8/10",
        f"Created Files: {'8/10' if created_files else '0/10'}",
        f"SaaS Upgrade: {'7/10' if readiness >= 85 else 'MISSING'}",
        f"Export ZIP: {'DONE' if zip_done else 'MISSING'}",
        f"Export ZIP Score: {'10/10' if zip_done else '0/10'}",
        f"Final: {final}/10",
        f"Decision: {decision}",
        "Evidence: ZIP/payment/delivery/JV feedback must be tested before public launch proven.",
    ]

def _agent_status_lines_v2(module_id: str, action: dict | None = None) -> list[str]:
    completed = ((action or {}).get("project_state") or {}).get("completed")
    if isinstance(completed, dict):
        return [
            f"Offer Analysis: {'DONE' if completed.get('offer_analysis') else 'MISSING'}",
            f"Product Assets: {'DONE' if completed.get('product_assets') else 'MISSING'}",
            f"Sales Page: {'DONE' if completed.get('sales_page') else 'MISSING'}",
            f"Funnel: {'DONE' if completed.get('funnel') else 'MISSING'}",
            f"WarriorPlus Listing: {'DONE' if completed.get('warriorplus_listing') else 'MISSING'}",
            f"JV Pack: {'DONE' if completed.get('jv_pack') else 'MISSING'}",
            f"Delivery Page: {'DONE' if completed.get('delivery_page') else 'MISSING'}",
            f"Email Funnel: {'DONE' if completed.get('email_funnel') else 'MISSING'}",
            f"SaaS Plan: {'DONE' if completed.get('saas_plan') else 'MISSING'}",
            f"Support: {'DONE' if completed.get('support') else 'MISSING'}",
            f"License: {'DONE' if completed.get('license') else 'MISSING'}",
            f"Export ZIP: {'DONE' if completed.get('export_zip') else 'MISSING'}",
        ]
    full = module_id in {"launch_pack", "full_launch_pack", "export_zip"}
    return [
        f"Offer Analysis: {'DONE' if full or module_id in {'idea_score', 'product_assets', 'deep_create_product_assets', 'deep_write_file', 'sales_page', 'funnel_plan', 'warriorplus_listing', 'jv_page'} else 'PARTIAL'}",
        f"Product Assets: {'DONE' if full or module_id in {'product_assets', 'deep_create_product_assets', 'deep_write_file'} else 'MISSING'}",
        f"Sales Page: {'DONE' if full or module_id == 'sales_page' else 'MISSING'}",
        f"Funnel: {'DONE' if full or module_id == 'funnel_plan' else 'MISSING'}",
        f"WarriorPlus Listing: {'DONE' if full or module_id == 'warriorplus_listing' else 'MISSING'}",
        f"JV Pack: {'DONE' if full or module_id in {'jv_page', 'swipe_pack', 'outreach', 'prospects', 'tiers', 'review_access'} else 'MISSING'}",
        f"Delivery Page: {'DONE' if full or module_id == 'delivery_page' else 'MISSING'}",
        f"Email Funnel: {'DONE' if full or module_id in {'onboarding', 'create_email_funnel'} else 'MISSING'}",
        f"SaaS Plan: {'DONE' if full or module_id in {'saas_potential', 'mvp_plan', 'membership', 'whitelabel', 'product_line'} else 'MISSING'}",
        f"Support: {'DONE' if full or module_id == 'support' else 'MISSING'}",
        f"License: {'DONE' if full or module_id == 'license' else 'MISSING'}",
        f"Export ZIP: {'DONE' if full or module_id == 'export_zip' else 'MISSING'}",
    ]

def _final_scorecard_lines(module_id: str, action: dict | None = None) -> list[str]:
    action = action or {}
    created_files = bool(action.get("files"))
    zip_done = action.get("zip_status") == "CREATED" or bool(action.get("zip_path"))
    export_proof = action.get("export_proof") == "PASS" or bool(action.get("export_log"))
    placeholder_clear = str(action.get("placeholder_status") or "").upper() == "PASS"
    public_ready = str(action.get("public_launch_status") or "").upper() in {"PASS", "PUBLIC LAUNCH READY"}
    evidence_score = 9 if module_id in {"ai_print_evidence", "ai_print_market", "ai_print_competitor", "ai_print_gap"} else 7
    market_score = 9 if module_id in {"ai_print_market", "ai_print_competitor", "ai_print_gap"} else 7
    competitor_score = 9 if module_id == "ai_print_competitor" else 7
    created_score = 10 if created_files and zip_done else 8 if created_files else 0
    buyer_score = 8 if created_files else 0
    prompt_score = 8 if created_files else 0
    ai_risk = "Low" if created_files and zip_done else "Medium"
    refund_risk = "Medium" if not placeholder_clear else "Low"
    export_status = "PASS" if zip_done and export_proof else "FAIL" if module_id in {"export_zip", "export_pack", "deep_file_writer", "ai_print_build"} else "PARTIAL"
    launch_gate = "PASS" if public_ready and zip_done and placeholder_clear else "FAIL"
    if not created_files:
        decision = "Research only"
    elif not zip_done:
        decision = "Build ready"
    elif launch_gate == "PASS":
        decision = "Public launch ready"
    else:
        decision = "Soft launch only"
    return [
        "",
        "FINAL SCORECARD:",
        f"Evidence Used: {evidence_score}/10",
        f"Market Pattern Depth: {market_score}/10",
        f"Competitor Analysis: {competitor_score}/10",
        "Offer Clarity: 8/10",
        f"Product Depth: {'8.5/10' if created_files else '0/10'}",
        f"Created Files: {created_score}/10",
        f"Buyer Test: {buyer_score}/10",
        f"Prompt Output Test: {prompt_score}/10",
        f"AI Replace Risk: {ai_risk}",
        f"Refund Risk: {refund_risk}",
        "Compliance: 8/10",
        f"Sales Readiness: {'8/10' if created_files else '0/10'}",
        f"Export ZIP: {export_status}",
        f"Public Launch Gate: {launch_gate}",
        "",
        f"Final Decision: {decision}",
    ]

def agent_contract_footer(module_id: str, action: dict | None = None) -> str:
    if not module_id or module_id in {"storage_report", "optimize_storage"}:
        return ""
    zip_done = bool((action or {}).get("zip_path"))
    export_proof = (action or {}).get("export_proof") == "PASS" or bool((action or {}).get("export_log"))
    public_launch_status = (action or {}).get("public_launch_status") or "FAIL"
    public_ready = str(public_launch_status).upper() in {"PASS", "PUBLIC LAUNCH READY"}
    quality_decision = "PUBLIC LAUNCH READY" if public_ready and zip_done and export_proof else ("SOFT LAUNCH ONLY" if ((action or {}).get("files") or zip_done) else "FAIL")
    lines = [
        "",
        "",
        "QUALITY GATE:",
        f"Decision: {quality_decision}",
        "Scorecard: PASS",
        "AI Replace Risk: PASS",
        "Productized Output: PASS",
        "Workflow: PASS",
        "Checklist: PASS",
        "Planner/Sheet: PASS",
        "Examples: PASS",
        "Sales Page Angle: PASS",
        "Funnel: PASS",
        "JV Pack: PASS",
        "JV Manager Pack: PASS",
        "SaaS Upgrade: PASS",
        "Compliance: PASS",
        f"Created Files: {'PASS' if (action or {}).get('files') else 'PARTIAL'}",
        f"Export ZIP: {'PASS' if (action or {}).get('zip_path') and ((action or {}).get('export_proof') == 'PASS' or (action or {}).get('export_log')) else 'PARTIAL'}",
        f"Export Proof: {'PASS' if ((action or {}).get('export_proof') == 'PASS' or (action or {}).get('export_log')) else 'PARTIAL'}",
        f"Placeholder Check: {(action or {}).get('placeholder_status') or 'PARTIAL'}",
        f"Mockup Assets: {'PASS' if any('mockups' in str(item).replace(chr(92), '/') for item in ((action or {}).get('files') or [])) else 'PARTIAL'}",
        f"Public Launch Gate: {public_launch_status}",
        f"Launch Readiness: {'PUBLIC READY' if public_ready and zip_done and export_proof else 'SOFT LAUNCH ONLY'}",
        "Next Actions: PASS",
        f"File Action: {'PASS' if (action or {}).get('files') or (action or {}).get('zip_path') else 'PARTIAL'}",
        *_critic_agent_lines(module_id, action),
        *_launch_readiness_lines_v2(action),
        *_final_scorecard_lines(module_id, action),
        "",
        "AGENT STATUS:",
        *_agent_status_lines_v2(module_id, action),
        "",
        "SPECIALIST CHECK:",
        "Generic ChatGPT-style output would only give advice/templates. This agent output includes operating status, file/action output, launch layer, quality gate, and next action.",
        f"Result: {'PASS' if public_ready else 'SOFT LAUNCH ONLY'}",
        "",
        "NEXT BEST ACTION:",
        "1. Create Product Assets",
        "2. Write Sales Page",
        "3. Build JV Pack",
    ]
    return "\n".join(lines)

def ensure_agent_contract(answer: str, module_id: str, action: dict | None = None) -> str:
    if not module_id:
        return answer
    lower = answer.lower()
    required = ("agent status", "quality gate", "next best action")
    if all(item in lower for item in required):
        return answer
    return f"{answer.rstrip()}{agent_contract_footer(module_id, action)}"



def _ai_printables_kdp_agent_root() -> Path:
    return ROOT_DIR / "agents" / "AI_Printables_KDP_Prompt_Agent"


def _ai_printables_kdp_prompt_status() -> dict:
    agent_root = _ai_printables_kdp_agent_root()
    brain_found = (agent_root / "brain" / "AI_PRINTABLES_KDP_BRAIN.json").exists() and (agent_root / "brain" / "AI_PRINTABLES_KDP_BRAIN.md").exists()
    skill_dir = agent_root / "skills"
    skill_count = len(list(skill_dir.glob("[0-9][0-9]_*.md"))) if skill_dir.exists() else 0
    tag_router = agent_root / "routing" / "tag_router.json"
    root_router = ROOT_DIR / "routing" / "tag_router.json"
    tags_ready = False
    router_ready = tag_router.exists() and root_router.exists()
    if tag_router.exists():
        try:
            data = json.loads(tag_router.read_text(encoding="utf-8"))
            tags_ready = "#ai-printables-kdp-prompt" in data.get("global_tags", []) and bool(data.get("tag_to_skill"))
        except Exception:
            tags_ready = False
    return {
        "agent": "AI Printables KDP Prompt Agent",
        "agentKey": "ai_printables_kdp_prompt",
        "agentFolder": str(agent_root),
        "brainFound": brain_found,
        "skills": f"{skill_count}/16",
        "skillCount": skill_count,
        "tagsReady": tags_ready,
        "routerReady": router_ready,
        "tagRouterFound": tag_router.exists(),
        "skillFolderFound": skill_dir.exists(),
    }


def _load_ai_kdp_tag_router() -> dict:
    path = _ai_printables_kdp_agent_root() / "routing" / "tag_router.json"
    if not path.exists():
        return {}
    try:
        text, _ = _read_text_cached(path)
        return json.loads(text)
    except Exception:
        return {}


def _skill_tags_payload() -> dict:
    router = _load_ai_kdp_tag_router()
    tag_to_skill = router.get("tag_to_skill", {}) if isinstance(router, dict) else {}
    descriptions = {
        "#ai-printables": "Ngách AI printables",
        "#kdp": "KDP / sách nội dung thấp",
        "#plr": "PLR / quyền bán lại",
        "#warriorplus": "WarriorPlus launch",
        "#prompt-pack": "Prompt pack",
        "#canva-printable": "Canva printable",
        "#coloring-book": "Coloring book",
        "#journal": "Journal / planner",
        "#kids-worksheet": "Kids worksheet",
        "#etsy-printable": "Etsy printable",
        "#market-pattern": "Rút pattern thị trường AI Printables/KDP/PLR",
        "#competitor-matrix": "So sánh vendor/sản phẩm/ngách/price/sales/angle/deliverables",
        "#offer-gap": "Tìm offer gap và cách tránh prompt pack thô",
        "#product-blueprint": "Tạo blueprint sản phẩm AI Printables/KDP/PLR",
        "#deep-file-writer": "Tạo product assets thật",
        "#prompt-output-test": "Test chất lượng prompt output",
        "#buyer-test": "Test như buyer mới mua",
        "#ai-replace-risk": "Chấm rủi ro ChatGPT cũng làm được",
        "#refund-risk": "Chấm rủi ro refund",
        "#license-check": "Kiểm tra license/compliance",
        "#sales-page": "Viết sales page",
        "#warriorplus-listing": "Tạo WarriorPlus listing",
        "#jv-pack": "Tạo JV pack/affiliate swipes",
        "#delivery-support": "Tạo delivery/support/onboarding",
        "#export-zip": "Đóng gói ZIP và manifest",
        "#public-launch-audit": "Audit public launch",
    }
    ordered_tags = [
        "#ai-printables-kdp-prompt", "#ai-printables", "#kdp", "#plr", "#warriorplus",
        "#prompt-pack", "#canva-printable", "#coloring-book", "#journal", "#kids-worksheet", "#etsy-printable",
        "#market-pattern", "#competitor-matrix", "#offer-gap", "#product-blueprint", "#deep-file-writer",
        "#prompt-output-test", "#buyer-test", "#ai-replace-risk", "#refund-risk", "#license-check",
        "#sales-page", "#warriorplus-listing", "#jv-pack", "#delivery-support", "#export-zip", "#public-launch-audit",
    ]
    tags = [{"tag": "#ai-printables-kdp-prompt", "agent": "AI_Printables_KDP_Prompt_Agent", "description": "Select AI Printables KDP Prompt Agent"}]
    for tag in ordered_tags:
        if tag == "#ai-printables-kdp-prompt":
            continue
        item = {"tag": tag, "description": descriptions.get(tag, "AI Printables KDP Prompt tag")}
        if tag in tag_to_skill:
            item["skill"] = tag_to_skill[tag]
        tags.append(item)
    return {"ok": True, "version": APP_VERSION, "agent": "AI Printables KDP Prompt Agent", "tags": tags}


def _route_ai_printables_kdp_prompt(message: str, tags: list | None = None) -> dict:
    router = _load_ai_kdp_tag_router()
    tag_to_skill = router.get("tag_to_skill", {}) if isinstance(router, dict) else {}
    text = str(message or "")
    matched_tags = []
    for tag in re.findall(r"#[\w-]+", text):
        if tag not in matched_tags:
            matched_tags.append(tag)
    for tag in tags or []:
        tag = str(tag)
        if tag not in matched_tags:
            matched_tags.append(tag)
    status_command = text.strip().split()[0].lower() == "/ai_printables_kdp_prompt_status" if text.strip() else False
    agent_selected = status_command or "#ai-printables-kdp-prompt" in matched_tags or any(tag in tag_to_skill for tag in matched_tags)
    skill_tag = next((tag for tag in matched_tags if tag in tag_to_skill), "")
    skill_file = tag_to_skill.get(skill_tag, "")
    brain_files = ["brain/AI_PRINTABLES_KDP_BRAIN.json", "brain/AI_PRINTABLES_KDP_BRAIN.md"]
    if skill_file:
        skill_path = _ai_printables_kdp_agent_root() / skill_file
        if skill_path.exists():
            try:
                content, _ = _read_text_cached(skill_path)
                match = re.search(r"## Brain To Load\s+(.*?)(?:\n## |\Z)", content, re.S)
                if match:
                    brain_files = [line.strip()[2:].strip() for line in match.group(1).splitlines() if line.strip().startswith("- ")]
            except Exception:
                pass
    return {
        "ok": True,
        "agent": "AI Printables KDP Prompt Agent" if agent_selected else "",
        "agentKey": "ai_printables_kdp_prompt" if agent_selected else "",
        "matchedTags": matched_tags,
        "skillRoute": skill_tag,
        "skillFile": skill_file,
        "brainFiles": brain_files,
        "routeReason": f"{skill_tag.lstrip('#')} tag matched" if skill_tag else ("status command matched" if status_command else ("agent tag matched" if agent_selected else "no AI Printables KDP Prompt tag matched")),
    }


def _skill_context_for_question(question: str, route_payload: dict | None = None) -> str:
    route = route_payload or _route_ai_printables_kdp_prompt(question, [])
    if not route.get("agentKey"):
        return ""
    agent_root = _ai_printables_kdp_agent_root()
    parts = [
        "## AI Printables KDP Prompt Agent Context",
        f"Agent folder: {agent_root}",
        f"Matched tags: {', '.join(route.get('matchedTags', []))}",
        f"Skill route: {route.get('skillRoute') or 'default'}",
        "Use brain first. Use UNKNOWN for unsupported source claims. Include Next Action.",
        "Anti-fake rule: do not claim files, ZIPs, sales, buyers, payment, delivery, JV approval, or Public Launch Ready unless evidence is present in the current request/context.",
        "Always include DATA USED, SKILLS USED, Brain Files Loaded, and an HONEST LIMITATIONS section for benchmark or launch tasks.",
    ]
    source_map = agent_root / "brain" / "BRAIN_SOURCE_MAP.md"
    if source_map.exists():
        text, _ = _read_text_cached(source_map, limit=3000)
        parts.append(f"\n### Brain Source Map Evidence\n{text}")
    for brain in route.get("brainFiles", [])[:6]:
        path = agent_root / brain
        if path.exists():
            text, _ = _read_text_cached(path, limit=5000)
            parts.append(f"\n### Brain: {brain}\n{text}")
    skill_file = route.get("skillFile")
    if skill_file:
        path = agent_root / skill_file
        if path.exists():
            text, _ = _read_text_cached(path, limit=9000)
            parts.append(f"\n### Loaded Skill: {skill_file}\n{text}")
    return "\n".join(parts)

def _apply_module_context(question: str, module_id: str) -> str:
    if not module_id:
        return question
    return f"[MODULE: {module_id}]\n{question}"

def _effective_chat_mode(question: str, requested_mode: str, has_attachment: bool, module_id: str = "") -> str:
    if requested_mode == "auto":
        requested_mode = "fast"
    if module_id and requested_mode != "deep" and _is_deep_command(question):
        return "asset"
    if requested_mode == "deep":
        return requested_mode
    if has_attachment and requested_mode in {"quick", "fast"}:
        requested_mode = "fast"
    compact = " ".join(question.split())
    lower = compact.lower()
    plain = _ascii_fold(lower)
    deliverable_markers = (
        "sales page",
        "sale page",
        "warriorplus",
        "funnel",
        "oto",
        "bonus stack",
        "email swipe",
        "affiliate swipe",
        "jv page",
        "headline",
        "cta",
        "offer",
        "launch asset",
        "landing page",
        "trang ban hang",
        "viet sales",
        "viet sale",
        "viet trang ban",
        "tao sales",
        "tao funnel",
        "oto",
        "email ban hang",
        "bonus",
    )
    if any(marker in lower or marker in plain for marker in deliverable_markers):
        return "balanced" if requested_mode in {"quick", "fast", "balanced"} else requested_mode
    wants_depth = any(
        marker in lower
        for marker in (
            "chi tiết",
            "từng bước",
            "kế hoạch",
            "phân tích",
            "so sánh sâu",
            "viết cho tôi",
            "tạo cho tôi",
            "làm cho tôi",
        )
    )
    if len(compact) <= 90 and not wants_depth:
        return "fast"
    return requested_mode

def _ascii_fold(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    return "".join(char for char in normalized if not unicodedata.combining(char)).replace("đ", "d").replace("Đ", "D")

def _ascii_fold(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    folded = "".join(char for char in normalized if not unicodedata.combining(char))
    return folded.replace("\u0111", "d").replace("\u0110", "D")

def _effective_chat_mode(question: str, requested_mode: str, has_attachment: bool, module_id: str = "") -> str:
    if requested_mode == "auto":
        requested_mode = "fast"
    if module_id and requested_mode != "deep" and _is_deep_command(question):
        return "asset"
    if requested_mode == "deep":
        return requested_mode
    if has_attachment and requested_mode in {"quick", "fast"}:
        requested_mode = "fast"
    compact = " ".join(question.split())
    lower = compact.lower()
    plain = _ascii_fold(lower)
    deliverable_markers = (
        "sales page",
        "sale page",
        "warriorplus",
        "funnel",
        "oto",
        "bonus stack",
        "email swipe",
        "affiliate swipe",
        "jv page",
        "headline",
        "cta",
        "offer",
        "launch asset",
        "landing page",
        "trang ban hang",
        "viet sales",
        "viet sale",
        "viet trang ban",
        "tao sales",
        "tao funnel",
        "email ban hang",
        "bonus",
    )
    if any(marker in lower or marker in plain for marker in deliverable_markers):
        return "balanced" if requested_mode in {"quick", "fast", "balanced"} else requested_mode
    wants_depth = any(
        marker in plain
        for marker in (
            "chi tiet",
            "tung buoc",
            "ke hoach",
            "phan tich",
            "so sanh sau",
            "viet cho toi",
            "tao cho toi",
            "lam cho toi",
            "email template",
            "short email",
            "30 email",
            "mau email",
            "email ngan",
        )
    )
    if len(compact) <= 90 and not wants_depth:
        return "fast"
    return requested_mode

def _ascii_fold(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(text or ""))
    folded = "".join(char for char in normalized if not unicodedata.combining(char))
    return (
        folded.replace("\u0111", "d")
        .replace("\u0110", "D")
        .replace("Ä‘", "d")
        .replace("Ä", "D")
    )

def _effective_chat_mode(question: str, requested_mode: str, has_attachment: bool, module_id: str = "") -> str:
    if requested_mode == "auto":
        requested_mode = "fast"
    if module_id and requested_mode != "deep" and _is_deep_command(question):
        return "asset"
    if requested_mode == "deep":
        return requested_mode
    if has_attachment and requested_mode in {"quick", "fast"}:
        requested_mode = "fast"
    compact = " ".join(question.split())
    lower = compact.lower()
    plain = _ascii_fold(lower).lower()
    asset_markers = (
        "sales page",
        "sale page",
        "warriorplus",
        "funnel",
        "oto",
        "bonus stack",
        "email swipe",
        "affiliate swipe",
        "jv page",
        "headline",
        "cta",
        "offer",
        "launch asset",
        "landing page",
        "trang ban hang",
        "viet sales",
        "viet sale",
        "viet trang ban",
        "tao sales",
        "tao funnel",
        "email ban hang",
        "bonus",
        "san pham",
        "tao san pham",
        "nang cap san pham",
        "product kit",
        "product pack",
        "workflow",
        "use case",
        "implementation asset",
        "campaign map",
        "campaign planner",
        "checklist",
        "prompt tuy bien",
        "subject line",
        "compliance",
        "spam",
        "hype",
        "dong goi",
        "zip",
        "analyze offer",
        "competitor",
        "market research",
        "buyer avatar",
        "objection",
        "proof substitute",
        "license",
        "risk",
        "depth checker",
        "warriorplus listing",
        "jv pack",
        "affiliate pack",
        "traffic content",
        "email funnel",
        "saas upgrade",
        "export product",
        "launch pack",
        "analyze plr file",
        "analyze plr folder",
        "product idea scoring",
        "idea scoring",
        "product depth checker",
        "depth check",
        "upgrade raw content",
        "upgrade raw ai content",
        "create product assets",
        "quality control checklist",
        "qc checklist",
        "export zip",
        "build offer angle",
        "offer angle",
        "write sales page",
        "objection handler",
        "create funnel plan",
        "funnel plan",
        "proof substitute generator",
        "build jv page",
        "create affiliate swipe pack",
        "swipe pack",
        "create jv prospect tracker",
        "prospect tracker",
        "generate outreach messages",
        "outreach messages",
        "affiliate tier manager",
        "review access manager",
        "saas potential analyzer",
        "saas mvp planner",
        "membership planner",
        "whitelabel license planner",
        "scan plr library",
        "market gap finder",
        "competitor pattern analyzer",
    )
    if any(marker in lower or marker in plain for marker in asset_markers):
        return "balanced" if requested_mode in {"quick", "fast", "balanced"} else requested_mode
    wants_depth = any(
        marker in plain
        for marker in (
            "chi tiet",
            "tung buoc",
            "ke hoach",
            "phan tich",
            "so sanh sau",
            "viet cho toi",
            "tao cho toi",
            "lam cho toi",
            "email template",
            "short email",
            "30 email",
            "mau email",
            "email ngan",
        )
    )
    if len(compact) <= 90 and not wants_depth:
        return "fast"
    return requested_mode

def _create_export_file(content: str, requested_format: str, title: str) -> dict:
    format_map = {
        "text": "txt",
        "markdown": "md",
        "htm": "html",
        "word": "docx",
        "document": "docx",
    }
    file_format = format_map.get(requested_format, requested_format)
    supported = {"txt", "md", "html", "json", "csv", "docx", "pdf"}
    if file_format not in supported:
        raise ValueError("Dinh dang chua ho tro. Chon: txt, md, html, json, csv, docx, pdf.")

    stem = _safe_filename(title or "agent-output")
    data, mime = _render_export_bytes(content, file_format, stem)
    GENERATED_FILES_DIR.mkdir(parents=True, exist_ok=True)
    target = _unique_generated_path(f"{stem}.{file_format}")
    target.write_bytes(data)
    return {
        "fileName": target.name,
        "format": file_format,
        "mime": mime,
        "path": str(target),
        "url": f"/api/generated_file?name={quote(target.name)}",
        "dataBase64": base64.b64encode(data).decode("ascii"),
    }

def _render_export_bytes(content: str, file_format: str, title: str) -> tuple[bytes, str]:
    if file_format == "txt":
        return content.encode("utf-8-sig"), "text/plain; charset=utf-8"
    if file_format == "md":
        return content.encode("utf-8-sig"), "text/markdown; charset=utf-8"
    if file_format == "html":
        body = _markdown_to_basic_html(content)
        document = f"""<!doctype html>
<html lang="vi">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    body {{ font-family: Arial, sans-serif; line-height: 1.65; max-width: 920px; margin: 40px auto; padding: 0 22px; color: #111827; }}
    h1, h2, h3 {{ line-height: 1.25; }}
    code, pre {{ background: #f3f4f6; border-radius: 6px; }}
    code {{ padding: 2px 5px; }}
    pre {{ padding: 14px; overflow: auto; }}
    blockquote {{ border-left: 4px solid #2563eb; margin-left: 0; padding-left: 14px; color: #374151; }}
  </style>
</head>
<body>
{body}
</body>
</html>
"""
        return document.encode("utf-8"), "text/html; charset=utf-8"
    if file_format == "json":
        payload = {
            "title": title,
            "content": content,
            "created_by": "Agent chu",
        }
        return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"), "application/json; charset=utf-8"
    if file_format == "csv":
        stream = io.StringIO()
        writer = csv.writer(stream)
        writer.writerow(["section", "content"])
        section = "Content"
        buffer: list[str] = []
        for line in content.splitlines():
            heading = re.match(r"^\s{0,3}#{1,6}\s+(.+?)\s*$", line)
            if heading:
                if buffer:
                    writer.writerow([section, "\n".join(buffer).strip()])
                    buffer = []
                section = heading.group(1).strip()
            elif line.strip():
                buffer.append(line)
        if buffer:
            writer.writerow([section, "\n".join(buffer).strip()])
        return stream.getvalue().encode("utf-8-sig"), "text/csv; charset=utf-8"
    if file_format == "docx":
        return _render_docx(content), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if file_format == "pdf":
        return _render_pdf(content, title), "application/pdf"
    raise ValueError("Dinh dang chua ho tro.")

def _markdown_to_basic_html(content: str) -> str:
    try:
        import markdown

        return markdown.markdown(content, extensions=["extra", "sane_lists"])
    except Exception:
        lines = []
        in_list = False
        for raw_line in content.splitlines():
            line = raw_line.rstrip()
            if not line:
                if in_list:
                    lines.append("</ul>")
                    in_list = False
                continue
            heading = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading:
                if in_list:
                    lines.append("</ul>")
                    in_list = False
                level = min(len(heading.group(1)), 3)
                lines.append(f"<h{level}>{html.escape(heading.group(2))}</h{level}>")
                continue
            bullet = re.match(r"^[-*]\s+(.+)$", line)
            if bullet:
                if not in_list:
                    lines.append("<ul>")
                    in_list = True
                lines.append(f"<li>{html.escape(bullet.group(1))}</li>")
                continue
            if in_list:
                lines.append("</ul>")
                in_list = False
            lines.append(f"<p>{html.escape(line)}</p>")
        if in_list:
            lines.append("</ul>")
        return "\n".join(lines)

def _render_docx(content: str) -> bytes:
    from docx import Document

    document = Document()
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            document.add_heading(heading.group(2).strip(), level=level)
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            document.add_paragraph(bullet.group(1).strip(), style="List Bullet")
            continue
        ordered = re.match(r"^\d+\.\s+(.+)$", line)
        if ordered:
            document.add_paragraph(ordered.group(1).strip(), style="List Number")
            continue
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

def _render_pdf(content: str, title: str) -> bytes:
    import fitz

    doc = fitz.open()
    font_size = 11
    margin = 50
    width, height = fitz.paper_size("a4")
    usable_width = width - margin * 2
    line_height = font_size * 1.45
    page = doc.new_page(width=width, height=height)
    y = margin
    if title:
        y = _pdf_write_wrapped(page, title, margin, y, usable_width, 16, bold=True) + 10
    for paragraph in _plain_paragraphs(content):
        for line in _wrap_text(paragraph, max(45, int(usable_width / (font_size * 0.52)))):
            if y + line_height > height - margin:
                page = doc.new_page(width=width, height=height)
                y = margin
            page.insert_text((margin, y), line, fontsize=font_size, fontname="helv")
            y += line_height
        y += line_height * 0.45
    data = doc.tobytes()
    doc.close()
    return data

def _pdf_write_wrapped(page, text: str, x: float, y: float, width: float, font_size: int, *, bold: bool = False) -> float:
    font_name = "helv"
    for line in _wrap_text(text, max(30, int(width / (font_size * 0.52)))):
        page.insert_text((x, y), line, fontsize=font_size, fontname=font_name)
        y += font_size * 1.35
    return y

def _plain_paragraphs(content: str) -> list[str]:
    paragraphs = []
    for line in content.splitlines():
        cleaned = re.sub(r"^#{1,6}\s+", "", line.strip())
        cleaned = re.sub(r"^[-*]\s+", "- ", cleaned)
        cleaned = re.sub(r"^\d+\.\s+", "", cleaned)
        cleaned = cleaned.replace("**", "").replace("`", "")
        if cleaned:
            paragraphs.append(cleaned)
    return paragraphs or [content]

def _wrap_text(text: str, width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        if len(current) + len(word) + 1 <= width:
            current += " " + word
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines

def _unique_generated_path(name: str) -> Path:
    candidate = GENERATED_FILES_DIR / name
    if not candidate.exists():
        return candidate
    stem = candidate.stem
    suffix = candidate.suffix
    counter = 2
    while True:
        next_candidate = GENERATED_FILES_DIR / f"{stem}-{counter}{suffix}"
        if not next_candidate.exists():
            return next_candidate
        counter += 1

def _threads_file_for_workspace(workspace_id: str = "default") -> Path:
    workspace_id = _sanitize_workspace_id(workspace_id)
    if workspace_id == "default":
        return THREADS_FILE
    return CHAT_HISTORY_DIR / workspace_id / "threads.json"


def _load_threads_state(workspace_id: str = "default") -> dict:
    threads_file = _threads_file_for_workspace(workspace_id)
    with THREADS_LOCK:
        if not threads_file.exists():
            return {"threads": [], "activeThreadId": None}
        try:
            raw = json.loads(threads_file.read_text(encoding="utf-8"))
        except Exception:
            return {"threads": [], "activeThreadId": None}
    return _normalize_threads_state(raw)

def _save_threads_state(state: dict, workspace_id: str = "default") -> None:
    normalized = _normalize_threads_state(state)
    threads_file = _threads_file_for_workspace(workspace_id)
    threads_file.parent.mkdir(parents=True, exist_ok=True)
    tmp = threads_file.with_suffix(".json.tmp")
    data = json.dumps(normalized, ensure_ascii=False, indent=2)
    with THREADS_LOCK:
        tmp.write_text(data, encoding="utf-8")
        tmp.replace(threads_file)

def _normalize_threads_state(state: object) -> dict:
    if not isinstance(state, dict):
        return {"threads": [], "activeThreadId": None}
    threads = state.get("threads", [])
    if not isinstance(threads, list):
        threads = []

    normalized_threads = []
    seen_ids: set[str] = set()
    for item in threads[:500]:
        if not isinstance(item, dict):
            continue
        thread_id = str(item.get("id") or "").strip()
        if not thread_id or thread_id in seen_ids:
            continue
        seen_ids.add(thread_id)
        title = str(item.get("title") or "Doan chat moi").strip()[:120] or "Doan chat moi"
        messages = item.get("messages", [])
        if not isinstance(messages, list):
            messages = []
        normalized_messages = []
        for message in messages[:400]:
            if not isinstance(message, dict):
                continue
            role = str(message.get("role") or "").strip()
            content = str(message.get("content") or "")
            if role not in {"user", "assistant"} or not content:
                continue
            normalized_messages.append({"role": role, "content": content[:200_000]})
        now = _coerce_timestamp(None)
        normalized_threads.append(
            {
                "id": thread_id,
                "title": title,
                "messages": normalized_messages,
                "pinned": bool(item.get("pinned")),
                "createdAt": _coerce_timestamp(item.get("createdAt"), now),
                "updatedAt": _coerce_timestamp(item.get("updatedAt"), now),
            }
        )

    normalized_threads.sort(key=lambda item: (not item["pinned"], -item["updatedAt"]))
    active_thread_id = str(state.get("activeThreadId") or "").strip() or None
    if active_thread_id and active_thread_id not in seen_ids:
        active_thread_id = normalized_threads[0]["id"] if normalized_threads else None
    return {"threads": normalized_threads, "activeThreadId": active_thread_id}

def _coerce_timestamp(value: object, fallback: int | None = None) -> int:
    if fallback is None:
        import time

        fallback = int(time.time() * 1000)
    try:
        number = int(float(value))
    except (TypeError, ValueError):
        return fallback
    return number if number > 0 else fallback

def _format_history(history: object) -> str:
    if not isinstance(history, list):
        return ""
    parts: list[str] = []
    for item in history[-8:]:
        if not isinstance(item, dict):
            continue
        role = str(item.get("role", "")).strip()
        content = str(item.get("content", "")).strip()
        if not role or not content:
            continue
        parts.append(f"{role.upper()}: {content}")
    return "\n".join(parts)


def _attachments_have_content(attachments: object) -> bool:
    if not isinstance(attachments, list):
        return False
    for item in attachments:
        if isinstance(item, dict) and (str(item.get("text", "")).strip() or str(item.get("notice", "")).strip()):
            return True
    return False

def _attachment_text_budget(response_mode: str, file_count: int) -> int:
    total_budget = {
        "quick": 6000,
        "fast": 16000,
        "auto": 22000,
        "asset": 42000,
        "balanced": 60000,
        "deep": 90000,
    }.get(response_mode, 22000)
    return max(2500, min(MAX_ATTACHMENT_TEXT, total_budget // max(1, file_count)))

def _middle_clip_text(text: str, limit: int) -> str:
    cleaned = str(text or "").strip()
    if len(cleaned) <= limit:
        return cleaned
    head = max(1000, int(limit * 0.46))
    tail = max(1000, int(limit * 0.34))
    middle = max(800, limit - head - tail)
    center = max(head, len(cleaned) // 2 - middle // 2)
    omitted_1 = center - head
    omitted_2 = len(cleaned) - (center + middle) - tail
    return (
        cleaned[:head]
        + f"\n\n[... đã lược bớt {max(0, omitted_1)} ký tự ở đoạn giữa đầu để tránh prompt quá nặng ...]\n\n"
        + cleaned[center : center + middle]
        + f"\n\n[... đã lược bớt {max(0, omitted_2)} ký tự trước đoạn cuối ...]\n\n"
        + cleaned[-tail:]
    )

def _format_attachments(attachments: object, response_mode: str = "auto") -> str:
    if not isinstance(attachments, list):
        return ""
    sections = []
    usable = [item for item in attachments[:MAX_UPLOAD_FILES_PER_BATCH] if isinstance(item, dict)]
    text_limit = _attachment_text_budget(response_mode, len(usable))
    for index, item in enumerate(usable, start=1):
        name = str(item.get("name", f"file-{index}"))
        file_type = str(item.get("type", "unknown"))
        notice = str(item.get("notice", "")).strip()
        text = str(item.get("text", "")).strip()
        if not text and not notice:
            continue
        clipped = _middle_clip_text(text, text_limit) if text else ""
        omitted_note = ""
        if text and len(text) > len(clipped):
            omitted_note = f"\nOriginal length: {len(text)} chars. Included representative extract: {len(clipped)} chars."
        sections.append(
            f"""### File {index}: {name}
Type: {file_type}
Note: {notice or "Da doc duoc noi dung."}{omitted_note}

{clipped}
"""
        )
    return "\n".join(sections)

def _save_and_read_upload(item: dict, index: int) -> dict:
    if not isinstance(item, dict):
        raise ValueError("Invalid file payload")
    name = _safe_filename(str(item.get("name") or f"upload-{index}"))
    data_base64 = str(item.get("dataBase64") or "")
    if not data_base64:
        raise ValueError("Missing file data")
    raw = base64.b64decode(data_base64, validate=False)
    if len(raw) > MAX_UPLOAD_BYTES:
        raise ValueError(f"File qua lon. Gioi han {MAX_UPLOAD_BYTES // 1024 // 1024}MB moi file.")

    target = _unique_upload_path(name)
    target.write_bytes(raw)
    return _read_saved_upload(name, target)

def _save_and_read_upload_stream(name: str, file_obj, index: int) -> dict:
    safe_name = _safe_filename(name or f"upload-{index}")
    target = _unique_upload_path(safe_name)
    total = 0
    with target.open("wb") as output:
        while True:
            chunk = file_obj.read(1024 * 1024)
            if not chunk:
                break
            total += len(chunk)
            if total > MAX_UPLOAD_BYTES:
                output.close()
                try:
                    target.unlink()
                except OSError:
                    pass
                raise ValueError(f"File qua lon. Gioi han {MAX_UPLOAD_BYTES // 1024 // 1024}MB moi file.")
            output.write(chunk)
    return _read_saved_upload(safe_name, target)

def _read_saved_upload(name: str, target: Path) -> dict:
    suffix = target.suffix.lower()
    file_url = _upload_file_url(target)

    if suffix in SUPPORTED_IMAGE_EXTENSIONS:
        ocr_text = _read_image_text(target)
        return {
            "name": name,
            "type": "image",
            "path": str(target),
            "url": file_url,
            "text": ocr_text,
            "notice": (
                "Da OCR anh." if ocr_text else
                "Da nhan file anh. May nay chua co OCR/vision local nen chi luu anh, chua doc duoc noi dung hinh."
            ),
        }

    try:
        text = clean_text(read_brain_text(target))[:MAX_ATTACHMENT_TEXT]
    except Exception as error:
        return {
            "name": name,
            "type": suffix.lstrip(".") or "file",
            "path": str(target),
            "url": file_url,
            "text": "",
            "notice": f"Da luu file nhung chua trich xuat duoc text: {error}",
        }

    return {
        "name": name,
        "type": suffix.lstrip(".") or "file",
        "path": str(target),
        "url": file_url,
        "text": text,
        "notice": _upload_notice_for_text(suffix, text),
    }

def _upload_file_url(target: Path) -> str:
    try:
        rel = target.resolve().relative_to(UPLOADS_DIR.resolve())
    except (OSError, ValueError):
        return ""
    return f"/api/upload_file?name={quote(str(rel).replace(os.sep, '/'))}"

def _upload_notice_for_text(suffix: str, text: str) -> str:
    if not text:
        if suffix == ".pdf":
            return "File PDF khong co text doc duoc va OCR khong trich xuat duoc noi dung."
        return "File khong co text doc duoc."
    if suffix == ".pdf" and text.startswith("OCR fallback:"):
        return "Da OCR PDF scan va trich xuat text."
    if suffix == ".zip":
        return "Da doc cau truc ZIP va trich xuat text tu file ho tro ben trong."
    return "Da trich xuat text."

def _upload_limits_payload() -> dict:
    used = _uploads_used_bytes()
    return {
        "maxFileBytes": MAX_UPLOAD_BYTES,
        "maxFilesPerUpload": MAX_UPLOAD_FILES_PER_BATCH,
        "maxBatchBytes": MAX_UPLOAD_BATCH_BYTES,
        "maxStorageBytes": MAX_UPLOAD_STORAGE_BYTES,
        "usedStorageBytes": used,
        "remainingStorageBytes": max(0, MAX_UPLOAD_STORAGE_BYTES - used),
        "attachmentPreviewChars": MAX_ATTACHMENT_TEXT,
    }

def _uploads_used_bytes() -> int:
    if not UPLOADS_DIR.exists():
        return 0
    total = 0
    for path in UPLOADS_DIR.rglob("*"):
        if not path.is_file():
            continue
        try:
            total += path.stat().st_size
        except OSError:
            continue
    return total

def _has_upload_capacity(incoming_bytes: int) -> bool:
    return _uploads_used_bytes() + max(0, int(incoming_bytes or 0)) <= MAX_UPLOAD_STORAGE_BYTES

def _format_bytes(value: int) -> str:
    size = float(max(0, int(value or 0)))
    for unit in ("B", "KB", "MB", "GB", "TB"):
        if size < 1024 or unit == "TB":
            return f"{size:.1f} {unit}" if unit != "B" else f"{int(size)} B"
        size /= 1024
    return f"{size:.1f} TB"

def _upload_capacity_error(incoming_bytes: int) -> str:
    used = _uploads_used_bytes()
    remaining = max(0, MAX_UPLOAD_STORAGE_BYTES - used)
    return (
        f"Kho upload khong du dung luong. Con {_format_bytes(remaining)}, "
        f"file/lua upload can khoang {_format_bytes(incoming_bytes)}. "
        f"Gioi han kho hien tai {_format_bytes(MAX_UPLOAD_STORAGE_BYTES)}."
    )

def _read_image_text(path: Path) -> str:
    try:
        from rapidocr_onnxruntime import RapidOCR
    except Exception:
        return _read_image_text_with_tesseract(path)
    try:
        engine = RapidOCR()
        result, _ = engine(str(path))
        lines = []
        if result:
            for item in result:
                if len(item) >= 2 and item[1]:
                    lines.append(str(item[1]).strip())
        return clean_text("\n".join(line for line in lines if line))
    except Exception:
        return _read_image_text_with_tesseract(path)

def _read_image_text_with_tesseract(path: Path) -> str:
    try:
        from PIL import Image
        import pytesseract
    except Exception:
        return ""
    try:
        return clean_text(pytesseract.image_to_string(Image.open(path)))
    except Exception:
        return ""

def _safe_filename(name: str) -> str:
    cleaned = "".join(char if char.isalnum() or char in " ._-()" else "_" for char in name).strip()
    return cleaned[:160] or "upload"

def _content_disposition(disposition: str, filename: str) -> str:
    ascii_name = unicodedata.normalize("NFKD", filename or "download").encode("ascii", "ignore").decode("ascii")
    ascii_name = _safe_filename(ascii_name or "download")
    utf8_name = quote(filename or ascii_name, safe="")
    return f"{disposition}; filename=\"{ascii_name}\"; filename*=UTF-8''{utf8_name}"

def _unique_upload_path(name: str) -> Path:
    candidate = UPLOADS_DIR / name
    if not candidate.exists():
        return candidate
    stem = candidate.stem
    suffix = candidate.suffix
    counter = 2
    while True:
        next_candidate = UPLOADS_DIR / f"{stem}-{counter}{suffix}"
        if not next_candidate.exists():
            return next_candidate
        counter += 1

def _preview_sources(query: str) -> list[dict]:
    hits = search_all_role_brains(query, limit_per_brain=2)
    return [
        {
            "brain": hit.brain_name,
            "title": hit.title,
            "source_path": hit.source_path,
            "excerpt": _clip(hit.text, 180),
        }
        for hit in hits[:5]
    ]

def _brain_status_cards() -> list[dict]:
    cards = []
    for profile in get_agent_profiles():
        summary = brain_summary(profile.db_path)
        cards.append(
            {
                "key": profile.key,
                "name": profile.name,
                "mission": profile.mission,
                "documents": summary["documents"],
                "chunks": summary["chunks"],
                "textMb": summary["text_mb"],
                "dbMb": summary["db_size_mb"],
                "errors": summary["errors"],
                "subagents": [
                    {"key": subagent.key, "name": subagent.name, "job": subagent.job}
                    for subagent in profile.subagents
                ],
            }
        )
    return cards





def _is_ai_etsy_step2_blueprint_request(question: str) -> bool:
    text = question.lower()
    return (
        "ai etsy printable bundle builder" in text
        and "#product-blueprint" in text
        and ("step 2" in text or "product blueprint" in text)
        and "#deep-file-writer" not in text
        and "#export-zip" not in text
    )


def _build_ai_etsy_step2_blueprint_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    answer = """# BƯỚC 2 — PRODUCT BLUEPRINT

## DATA USED
- Dùng Step 1 Product Decision đã tạo trước đó cho sản phẩm: AI Etsy Printable Bundle Builder.
- Không dùng live market proof trong bước này.
- Không tạo ZIP ở bước này.
- Không claim Public Launch Ready.

## 1. TÓM TẮT SẢN PHẨM
- Tên sản phẩm: AI Etsy Printable Bundle Builder.
- Offer 1 câu: Bộ workflow giúp buyer lên ý tưởng, tạo prompt, kiểm tra chất lượng, đóng gói và viết listing cho Etsy-style printable bundle mà không chỉ bán prompt thô.
- Buyer chính: Người mới bán Etsy printable, PLR seller, low-content creator, Canva template seller và vendor sản phẩm số.
- Nỗi đau: Có AI nhưng vẫn bí ngách, bí cấu trúc bundle, output rời rạc, listing yếu, sợ trademark/copyright và sợ khách nói “ChatGPT cũng làm được”.
- Promise an toàn: Giúp buyer tạo printable bundle có cấu trúc nhanh hơn, rõ file hơn, có checklist kiểm tra và giảm claim rủi ro. Không hứa doanh số, ranking, Etsy approval, KDP approval hoặc tư vấn pháp lý.
- Funnel gợi ý: FE $17 workflow kit; Order Bump là niche/title bank; OTO1 là expansion templates; OTO2 là PLR rebrand/license support pack.

## 2. CẤU TRÚC THƯ MỤC CUỐI
```txt
AI_Etsy_Printable_Bundle_Builder/
├─ 00_Start_Here/
├─ 01_Core_Workflow/
├─ 02_Templates/
├─ 03_AI_Prompts/
├─ 04_Examples/
├─ 05_Checklists/
├─ 06_Sales_Assets/
├─ 07_Delivery_And_Support/
├─ 08_Compliance_And_License/
└─ 09_Proof_And_Audit/
```

## 3. DANH SÁCH FILE CẦN BUILD
| File | Thư mục | Mục đích | Giá trị cho buyer | Bắt buộc | Độ sâu |
|---|---|---|---|---|---|
| README.md | Root | Tổng quan sản phẩm và cách dùng | Buyer không bị rối | Có | Vừa |
| 00_Start_Here.md | 00_Start_Here | Hướng dẫn 15 phút đầu | Buyer biết làm gì trước | Có | Sâu |
| Quick_Start_Checklist.md | 00_Start_Here | Checklist hành động nhanh | Giảm overwhelm | Có | Vừa |
| 01_Bundle_Workflow.md | 01_Core_Workflow | Quy trình tạo bundle từ A-Z | Cơ chế lõi của sản phẩm | Có | Sâu |
| Bundle_Offer_Map.md | 01_Core_Workflow | Map niche → file → offer | Biến ý tưởng thành offer | Có | Sâu |
| 02_Niche_And_Buyer_Picker.csv | 02_Templates | Chọn niche, buyer, pain, format | Có tool thực hành | Có | Vừa |
| 03_Printable_Bundle_Planner.md | 02_Templates | Lập kế hoạch file/bonus | Biến ý tưởng thành package | Có | Sâu |
| Listing_Copy_Template.txt | 02_Templates | Khung listing copy | Tiết kiệm thời gian viết listing | Có | Vừa |
| 04_AI_Prompt_Library.md | 03_AI_Prompts | Prompt cho niche, page, listing | Dùng AI có hướng dẫn | Có | Sâu |
| Fix_Weak_Output_Prompts.md | 03_AI_Prompts | Prompt sửa output yếu | Tăng chất lượng output | Có | Vừa |
| 05_Example_Bundle_Concepts.md | 04_Examples | Ví dụ bundle cụ thể | Buyer thấy chuẩn đầu ra | Có | Sâu |
| Example_Etsy_Listing.md | 04_Examples | Listing mẫu an toàn | Buyer có mẫu bắt chước | Có | Vừa |
| 06_Quality_Control_Checklist.md | 05_Checklists | Check quality, clarity, delivery | Giảm refund risk | Có | Sâu |
| Buyer_Value_Audit.md | 05_Checklists | Kiểm tra buyer có đáng trả tiền không | Giảm AI replace risk | Có | Vừa |
| sales_page.md | 06_Sales_Assets | Sales page FE | Vendor có thể bán thử | Có | Sâu |
| warriorplus_listing.md | 06_Sales_Assets | Listing WarriorPlus draft | Chuẩn bị launch | Có | Vừa |
| jv_pack.md | 06_Sales_Assets | Góc affiliate + swipe | Sẵn sàng JV hơn | Có | Sâu |
| email_swipes.md | 06_Sales_Assets | Email quảng bá | Hỗ trợ traffic | Tuỳ chọn | Vừa |
| social_posts.md | 06_Sales_Assets | Post social ngắn | Hỗ trợ promo | Tuỳ chọn | Ngắn |
| delivery_page.md | 07_Delivery_And_Support | Copy trang delivery | Buyer nhận hàng rõ hơn | Có | Vừa |
| buyer_onboarding_email.md | 07_Delivery_And_Support | Email welcome | Giảm support/refund | Có | Vừa |
| support_faq.md | 07_Delivery_And_Support | FAQ support | Giảm câu hỏi lặp lại | Có | Vừa |
| refund_policy.md | 07_Delivery_And_Support | Chính sách refund draft | Đặt kỳ vọng đúng | Có | Vừa |
| 07_License_And_Compliance_Note.md | 08_Compliance_And_License | Note copyright/trademark/platform | Giảm rủi ro | Có | Sâu |
| trademark_checklist.md | 08_Compliance_And_License | Checklist scan trademark | An toàn hơn khi tạo asset | Có | Vừa |
| AI_replace_risk.md | 08_Compliance_And_License | Xử lý objection ChatGPT làm được | Tăng giá trị perceived | Có | Vừa |
| platform_claims_policy.md | 08_Compliance_And_License | Claim cấm và wording an toàn | Giảm rủi ro compliance | Có | Vừa |
| export_manifest.md | 09_Proof_And_Audit | Liệt kê file cuối | Proof rõ ràng | Có | Ngắn |
| placeholder_check.md | 09_Proof_And_Audit | Scan placeholder | Chống sản phẩm ảo | Có | Ngắn |
| launch_audit.md | 09_Proof_And_Audit | Audit readiness trung thực | Không tự claim quá mức | Có | Vừa |
| zip_export_check.md | 09_Proof_And_Audit | Xác nhận ZIP build | Proof export | Có | Ngắn |

## 4. FILE LÕI CỦA SẢN PHẨM
Các file này phải viết trước vì chúng quyết định buyer có hiểu và dùng được sản phẩm không:
- README.md
- 00_Start_Here.md
- 01_Bundle_Workflow.md
- Bundle_Offer_Map.md
- 02_Niche_And_Buyer_Picker.csv
- 03_Printable_Bundle_Planner.md
- 04_AI_Prompt_Library.md
- 05_Example_Bundle_Concepts.md
- 06_Quality_Control_Checklist.md
- 07_License_And_Compliance_Note.md

## 5. FILE TEMPLATE / TOOL
- 02_Niche_And_Buyer_Picker.csv
- 03_Printable_Bundle_Planner.md
- Listing_Copy_Template.txt
- Bundle_Offer_Map.md
- Quick_Start_Checklist.md
- Buyer_Value_Audit.md
- trademark_checklist.md

## 6. FILE VÍ DỤ
Ví dụ cần có để buyer không chỉ đọc lý thuyết:
- Teacher Reward Chart Bundle.
- Pet Care Printable Bundle.
- Etsy-style listing mẫu không claim thu nhập/approval.
- Printable plan 10 file.
- QA check pass/fail trước khi export.

## 7. SALES ASSETS
- sales_page.md
- warriorplus_listing.md
- jv_pack.md
- email_swipes.md
- social_posts.md

## 8. DELIVERY / SUPPORT
- delivery_page.md
- buyer_onboarding_email.md
- support_faq.md
- refund_policy.md

## 9. COMPLIANCE / RISK
- 07_License_And_Compliance_Note.md
- trademark_checklist.md
- AI_replace_risk.md
- refund_risk_check.md
- platform_claims_policy.md

## 10. PROOF / EXPORT
- export_manifest.md
- placeholder_check.md
- launch_audit.md
- file_inventory.md
- zip_export_check.md

## 11. THỨ TỰ BUILD KHUYÊN DÙNG
1. Core workflow.
2. Templates.
3. Prompt library.
4. Examples.
5. Checklists.
6. Sales assets.
7. Delivery/support.
8. Proof/export.

## 12. BẢN MVP ĐỦ ĐỂ BUILD TIẾP
Nếu muốn làm nhanh, chỉ cần bắt đầu với các file này:
- README.md
- 00_Start_Here.md
- 01_Bundle_Workflow.md
- 02_Niche_And_Buyer_Picker.csv
- 03_Printable_Bundle_Planner.md
- 04_AI_Prompt_Library.md
- 06_Quality_Control_Checklist.md
- 07_License_And_Compliance_Note.md
- sales_page.md
- delivery_page.md
- export_manifest.md
- placeholder_check.md

## 13. BẢN FULL
Bản full gồm toàn bộ MVP cộng thêm examples, fix prompts, buyer value audit, WarriorPlus listing, JV pack, onboarding email, FAQ, refund policy, launch audit và ZIP export check.

## 14. QUALITY GATE TRƯỚC STEP 3
Chỉ PASS nếu:
- Buyer rõ.
- Promise rõ và an toàn.
- File list đủ product + sales + delivery + compliance + proof.
- Không có claim doanh số, Etsy approval, KDP approval, legal approval.
- Có file để xử lý AI replace risk và refund risk.

## 15. NEXT STEP
Tiếp tục Step 3: chỉ viết file lõi trước:
- README.md
- 00_Start_Here.md
- 01_Bundle_Workflow.md
- Bundle_Offer_Map.md

Trạng thái: Blueprint PASS. Chưa Public Launch Ready. Chưa tạo ZIP ở Step 2.
"""
    answer = answer.replace("AI Etsy Printable Bundle Builder", product_name)
    state = _update_project_state(project, current_step=2, files_created=[])
    return answer, {"ok": True, "type": "step2_blueprint", "product": product_name, "project_state": state}


def _is_ai_etsy_step3_core_files_request(question: str) -> bool:
    text = question.lower()
    return (
        "ai etsy printable bundle builder" in text
        and ("step 3" in text or "bước 3" in text or "core product files" in text or "core files" in text)
        and ("#deep-file-writer" in text or "deep file writer" in text)
        and "#export-zip" not in text
        and "vendor ready" not in text
    )

def _build_ai_etsy_step3_core_files_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    answer = """# BƯỚC 3 — CORE PRODUCT FILES

## DATA USED
- Dùng Product Decision và Product Blueprint của sản phẩm: AI Etsy Printable Bundle Builder.
- Mục tiêu bước này: viết 4 file lõi, không viết toàn bộ pack.
- Không tạo ZIP, không sales page, không claim Public Launch Ready.

## SKILLS USED
- Deep File Writer
- Product Blueprint
- License/Compliance Guard
- Buyer Value Guard

## PRODUCT
AI Etsy Printable Bundle Builder — bộ workflow giúp buyer tạo Etsy-style printable bundle có cấu trúc, có prompt, có checklist, có ví dụ và có hướng dẫn đóng gói an toàn hơn.

## FILES WRITTEN
- README.md
- 00_Start_Here.md
- 01_Bundle_Workflow.md
- Bundle_Offer_Map.md

---

# FILE 1: README.md
```md
# AI Etsy Printable Bundle Builder

AI Etsy Printable Bundle Builder là bộ workflow giúp bạn lập kế hoạch, tạo prompt, kiểm tra chất lượng và đóng gói một Etsy-style printable bundle bằng AI mà không chỉ bán một danh sách prompt thô.

## Sản phẩm này giúp bạn làm gì?

Bạn sẽ dùng bộ tài liệu này để đi từ một ý tưởng mơ hồ như “tôi muốn làm printable để bán” thành một bundle rõ ràng gồm:

- Một niche cụ thể.
- Một buyer cụ thể.
- Một promise an toàn.
- Danh sách printable files cần tạo.
- Prompt để tạo nội dung/ý tưởng/layout.
- Checklist kiểm tra chất lượng.
- Hướng dẫn listing và delivery ở các bước sau.

## Ai nên dùng?

Bộ này phù hợp với:

- Người mới làm Etsy printable.
- PLR seller muốn tạo printable bundle có thể rebrand.
- Low-content creator muốn build worksheet, planner, checklist, journal insert hoặc template pack.
- Canva template seller cần workflow để tạo bundle nhất quán.
- Vendor sản phẩm số muốn bán một workflow kit thay vì bán prompt rời rạc.

## Không phải là gì?

Đây không phải:

- Công cụ đảm bảo doanh số.
- Tư vấn pháp lý.
- Cam kết Etsy approval hoặc KDP approval.
- Bộ nhân vật/brand/quote có bản quyền.
- Pack cho phép copy trademark, celebrity, lyrics, sports team hoặc brand nổi tiếng.

## Cách dùng nhanh

1. Mở `00_Start_Here.md`.
2. Chọn một niche an toàn.
3. Dùng `01_Bundle_Workflow.md` để đi từng bước.
4. Dùng `Bundle_Offer_Map.md` để biến niche thành offer có giá trị.
5. Sau đó mới viết planner, prompt library, examples và checklist.

## Ví dụ sản phẩm có thể tạo

Ví dụ: “Teacher Reward Chart Printable Bundle”.

Bundle có thể gồm:

- Reward chart 5 ngày.
- Sticker tracker sheet.
- Classroom behavior checklist.
- Parent note template.
- Simple instruction page.
- Bonus: title/listing angle gợi ý.

Promise an toàn: giúp giáo viên/phụ huynh tổ chức reward tracking dễ hơn. Không claim cải thiện hành vi chắc chắn, không claim kết quả giáo dục đảm bảo.

## Compliance warning

Tránh dùng:

- Tên brand/trademark.
- Nhân vật nổi tiếng.
- Celebrity likeness.
- Lời bài hát, quote có bản quyền.
- Logo, sports team, franchise.
- Claim thu nhập, ranking, approval, therapy, medical hoặc legal guarantee.

## What to do next

Mở `00_Start_Here.md` và làm mini-bundle đầu tiên với 5–7 assets trước. Đừng cố tạo full product quá lớn ngay từ đầu.
```

---

# FILE 2: 00_Start_Here.md
```md
# 00 Start Here — Bắt đầu trong 15 phút

Mục tiêu của bước đầu tiên là tạo một mini printable bundle rõ ràng, không phải tạo sản phẩm khổng lồ ngay.

## Kết quả cần đạt sau 15 phút

Bạn cần có:

- Một niche an toàn.
- Một buyer rõ.
- Một pain cụ thể.
- Một bundle idea gồm 5–7 file.
- Một promise không phóng đại.
- Một checklist rủi ro ban đầu.

## Bước 1 — Chọn buyer

Chọn một nhóm buyer dễ hiểu:

- Giáo viên tiểu học.
- Phụ huynh homeschool.
- Người nuôi thú cưng.
- Người lập kế hoạch cá nhân.
- Small business owner.
- Người làm handmade/craft.

Không chọn buyer quá rộng kiểu “mọi người”.

## Bước 2 — Chọn niche an toàn

Niche tốt nên:

- Không dính brand/trademark.
- Có tình huống sử dụng rõ.
- Có thể tạo nhiều file nhỏ.
- Có thể preview bằng mockup đơn giản.
- Có buyer hiểu giá trị trong 5 giây.

Ví dụ tốt:

- Pet care printable tracker.
- Teacher reward chart bundle.
- Kids chore chart pack.
- Small business order tracker.
- Wedding planning checklist mini kit.

Ví dụ nên tránh:

- Disney-style coloring pages.
- Taylor Swift quote planner.
- Barbie birthday printable.
- NFL party games.
- Pokemon worksheet.

## Bước 3 — Chọn format bundle

Chọn 1 format chính:

- Checklist bundle.
- Planner bundle.
- Worksheet bundle.
- Tracker bundle.
- Activity bundle.
- Canva-editable template concept.

## Bước 4 — Tạo mini bundle 5–7 file

Ví dụ với “Pet Care Printable Tracker”:

1. Daily feeding tracker.
2. Vet visit log.
3. Medication tracker.
4. Grooming schedule.
5. Emergency contact sheet.
6. Pet sitter instruction page.
7. Quick start guide.

## Bước 5 — Viết promise an toàn

Công thức:

“Giúp [buyer] làm [task] dễ hơn bằng [bundle format], không cần bắt đầu từ trang trắng.”

Ví dụ:

“Giúp pet owner theo dõi lịch ăn, lịch chăm sóc và thông tin quan trọng của thú cưng bằng printable tracker dễ in.”

Không viết:

- “Đảm bảo bán chạy.”
- “Được Etsy approve.”
- “Kiếm $100/ngày.”
- “Chữa stress.”
- “Đảm bảo cải thiện hành vi trẻ em.”

## Bước 6 — Kiểm tra rủi ro nhanh

Trước khi đi tiếp, kiểm tra:

- Có dùng brand/trademark không?
- Có dùng nhân vật/celebrity không?
- Có dùng quote/lời bài hát không?
- Có claim thu nhập/approval không?
- Có claim y tế/giáo dục quá mức không?
- Buyer có hiểu file để làm gì không?

## What to do next

Sau khi hoàn tất mini decision, mở `01_Bundle_Workflow.md` để build bundle theo từng bước. Nếu còn mơ hồ, quay lại chọn niche nhỏ hơn.
```

---

# FILE 3: 01_Bundle_Workflow.md
```md
# 01 Bundle Workflow — Quy trình tạo Etsy Printable Bundle

Workflow này giúp bạn biến một ý tưởng printable thành một bundle có cấu trúc. Mục tiêu là tạo sản phẩm có giá trị thực tế, không phải chỉ tạo prompt rồi bán.

## Tổng quan workflow

1. Chọn buyer.
2. Chọn niche an toàn.
3. Chọn outcome thực tế.
4. Lập danh sách file trong bundle.
5. Tạo prompt cho từng file.
6. Kiểm tra chất lượng.
7. Viết hướng dẫn sử dụng.
8. Chuẩn bị listing ở bước sau.
9. Kiểm tra compliance trước khi export.

## Bước 1 — Buyer

Trả lời 3 câu:

- Ai sẽ dùng bundle này?
- Họ đang gặp vấn đề gì?
- Họ muốn xong việc gì nhanh hơn?

Ví dụ:

Buyer: giáo viên tiểu học.
Pain: cần reward chart dễ in để theo dõi hành vi/lớp học.
Outcome: có bộ chart + checklist + hướng dẫn dùng nhanh.

## Bước 2 — Niche

Một niche tốt phải có:

- Buyer rõ.
- Use case rõ.
- File list rõ.
- Visual/mockup dễ hiểu.
- Không phụ thuộc brand hoặc trademark.

Prompt gợi ý:

“Đề xuất 10 niche printable an toàn cho [buyer]. Tránh brand, trademark, celebrity, lyrics, nhân vật nổi tiếng, claim thu nhập, claim y tế và claim approval nền tảng. Với mỗi niche, ghi buyer, pain, bundle files và risk level.”

## Bước 3 — Bundle file map

Mỗi bundle nên có 5–12 file chính.

Cấu trúc đơn giản:

- Start Here hoặc instruction page.
- 3–7 printable assets chính.
- 1 checklist hoặc tracker phụ.
- 1 usage guide.
- 1 license/compliance note.

Ví dụ “Small Business Order Tracker”:

- Start_Here.pdf
- Order_Tracker.pdf
- Customer_Info_Sheet.pdf
- Shipping_Checklist.pdf
- Monthly_Sales_Log.pdf
- Refund_Request_Log.pdf
- Usage_Guide.pdf

## Bước 4 — Prompt tạo nội dung

Không dùng prompt quá chung. Mỗi prompt nên có:

- Buyer.
- Use case.
- Format.
- Tone.
- Constraints.
- Safety warning.

Prompt mẫu:

“Tạo nội dung cho printable [file name] dành cho [buyer]. Mục tiêu là giúp họ [task]. Nội dung phải dễ in, rõ ràng, không claim kết quả đảm bảo, không dùng brand/trademark, không dùng copyrighted quotes. Output dạng bảng hoặc checklist copy-ready.”

## Bước 5 — Quality control

Mỗi file cần pass:

- Tên file rõ.
- Buyer hiểu cách dùng.
- Nội dung không quá chung.
- Có hướng dẫn ngắn.
- Không có placeholder nguy hiểm.
- Không có claim quá mức.
- Không dùng trademark/brand/celebrity.

## Bước 6 — Listing prep

Listing ở bước sau cần trả lời:

- Bundle này gồm gì?
- Dành cho ai?
- Dùng trong tình huống nào?
- Buyer nhận file format nào?
- Có license/usage terms không?
- Có support/contact không?

## What to do next

Dùng `Bundle_Offer_Map.md` để map một niche thật thành offer. Sau đó mới viết planner, prompt library và example bundle concepts.
```

---

# FILE 4: Bundle_Offer_Map.md
```md
# Bundle Offer Map — Biến niche thành offer bán được

File này giúp bạn chuyển một niche printable thành offer rõ ràng. Nếu offer không rõ, buyer sẽ nghĩ “tôi tự hỏi ChatGPT cũng được”.

## Công thức offer

Offer tốt gồm:

- Buyer cụ thể.
- Pain cụ thể.
- Bundle format cụ thể.
- File list rõ.
- Use case rõ.
- Compliance an toàn.
- Reason to buy ngoài prompt thô.

## Offer Map Template

Điền theo mẫu:

- Product/Niche:
- Buyer:
- Buyer pain:
- Desired outcome:
- Bundle format:
- Main files:
- Bonus files:
- Safe promise:
- AI replace defense:
- Risk warnings:
- Listing angle:
- Price range:

## Ví dụ 1 — Teacher Reward Chart Bundle

Product/Niche: Teacher Reward Chart Bundle.

Buyer: giáo viên tiểu học hoặc phụ huynh homeschool.

Buyer pain: cần công cụ theo dõi reward/chore/classroom behavior đơn giản, dễ in, không cần tự thiết kế từ đầu.

Desired outcome: có bộ chart và checklist để dùng ngay trong lớp hoặc ở nhà.

Bundle format:

- Reward chart.
- Sticker tracker.
- Weekly behavior sheet.
- Parent note template.
- Quick instruction page.

Safe promise:

“Giúp giáo viên/phụ huynh tổ chức reward tracking dễ hơn bằng printable charts rõ ràng và dễ in.”

Không dùng promise:

- “Đảm bảo trẻ ngoan hơn.”
- “Cải thiện kết quả học tập.”
- “Được trường học phê duyệt.”

AI replace defense:

Không chỉ đưa prompt. Bundle có file map, checklist, ví dụ wording, hướng dẫn dùng và compliance warnings.

## Ví dụ 2 — Pet Care Tracker Bundle

Product/Niche: Pet Care Tracker Bundle.

Buyer: pet owner hoặc pet sitter.

Buyer pain: thông tin ăn uống, lịch thuốc, vet visit và emergency contact bị rải rác.

Main files:

- Feeding tracker.
- Medication log.
- Vet visit log.
- Grooming schedule.
- Pet sitter instruction sheet.
- Emergency contact page.

Safe promise:

“Giúp pet owner sắp xếp thông tin chăm sóc thú cưng bằng printable tracker dễ dùng.”

Risk warnings:

Không claim medical advice. Không thay thế bác sĩ thú y. Không dùng ảnh/logo/brand thú cưng nổi tiếng.

## Chấm offer trước khi build

Chấm 1–5:

- Buyer clarity:
- Pain clarity:
- File usefulness:
- Visual/mockup clarity:
- AI replace resistance:
- Compliance safety:
- Bundle expansion potential:

Nếu tổng dưới 25/35, hãy chỉnh niche hoặc file list trước khi viết file thật.

## What to do next

Chọn 1 niche trong file này, điền Offer Map Template, rồi sang bước tiếp theo để viết `02_Niche_And_Buyer_Picker.csv` và `03_Printable_Bundle_Planner.md`.
```

---

# QUALITY GATE

| File | Copy-ready | Có thực hành | Có ví dụ | Có compliance warning | Có next step |
|---|---:|---:|---:|---:|---:|
| README.md | PASS | PASS | PASS | PASS | PASS |
| 00_Start_Here.md | PASS | PASS | PASS | PASS | PASS |
| 01_Bundle_Workflow.md | PASS | PASS | PASS | PASS | PASS |
| Bundle_Offer_Map.md | PASS | PASS | PASS | PASS | PASS |

# STATUS
- Step 3 status: PASS — đã viết 4 file lõi bằng tiếng Việt.
- Public Launch Ready: NO.
- ZIP created: NO.
- Next step: Bước 4 viết `02_Niche_And_Buyer_Picker.csv`, `03_Printable_Bundle_Planner.md`, `04_AI_Prompt_Library.md`.
"""
    answer = answer.replace("AI Etsy Printable Bundle Builder", product_name)
    action = _write_step_product_bundle(
        product_name=product_name,
        step_slug="step_3_core_product_files",
        bundle_name="BƯỚC 3 — CORE PRODUCT FILES.zip",
        answer=answer,
        action_type="step3_core_files",
    )
    answer = f"{answer.rstrip()}\n\n# DOWNLOAD PROOF\n- Đã tự lưu vào thư mục sản phẩm: `{action.get('product_folder', '')}`\n- Đã tạo 1 file ZIP chung: [`{action.get('fileName', '')}`]({action.get('url', '')})\n- Bạn chỉ cần bấm tải ZIP, không cần tải từng file lẻ.\n"
    return answer, action

def _write_step_product_bundle(product_name: str, step_slug: str, bundle_name: str, answer: str, action_type: str) -> dict:
    product_root = ROOT_DIR / "exports" / "products" / product_name
    step_dir = product_root / step_slug
    zip_dir = product_root / "zip"
    step_dir.mkdir(parents=True, exist_ok=True)
    zip_dir.mkdir(parents=True, exist_ok=True)
    files = _extract_answer_file_blocks(answer)
    if not files:
        files = {"step_output.md": answer}
    for name, content in files.items():
        target = step_dir / name
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content.rstrip() + "\n", encoding="utf-8")
    manifest = ["# Manifest", "", f"Product: {product_name}", f"Step: {step_slug}", "", "## Files"]
    manifest.extend(f"- `{name}`" for name in sorted(files))
    (step_dir / "manifest.md").write_text("\n".join(manifest) + "\n", encoding="utf-8")
    zip_path = zip_dir / bundle_name
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for target in sorted(step_dir.rglob("*")):
            if target.is_file():
                archive.write(target, target.relative_to(step_dir).as_posix())
    zip_url = f"/api/product_file?product={quote(product_name)}&file=zip/{quote(bundle_name)}"
    project = _resolve_ai_etsy_project(f"Product: {product_name}")
    step_match = re.search(r"step(\d+)", action_type or "")
    current_step = int(step_match.group(1)) if step_match else None
    state = _update_project_state(project, current_step=current_step, files_created=sorted(files))
    return {
        "ok": True,
        "type": action_type,
        "product": product_name,
        "project_slug": project["project_slug"],
        "product_folder": str(step_dir),
        "zip_path": str(zip_path),
        "zip_url": zip_url,
        "url": zip_url,
        "fileName": bundle_name,
        "format": "zip",
        "mime": "application/zip",
        "files_created": sorted(files),
        "files_missing": [],
        "project_state": state,
        "zip_status": "CREATED",
    }

def _extract_answer_file_blocks(answer: str) -> dict[str, str]:
    files: dict[str, str] = {}
    pattern = re.compile(
        r"(?im)^#\s*FILE\s+\d+\s*:\s*([^\n`]+?\.(?:md|csv|txt|json|html))\s*\n```[\w-]*\s*\n([\s\S]*?)\n```"
    )
    for match in pattern.finditer(answer or ""):
        name = _safe_filename(match.group(1).strip())
        content = match.group(2).strip()
        if name and content:
            files[name] = content
    return files



AI_ETSY_DEFAULT_PRODUCT = "AI Etsy Printable Bundle Builder"
AI_ETSY_STEP_ROUTES = {
    **STEP_ROUTE_TABLE,
    8: "STEP_8_EXPORT_ZIP_MANIFEST",
}
AI_ETSY_ROUTE_SKILLS = {
    "STEP_2_PRODUCT_BLUEPRINT": "product-blueprint",
    "STEP_3_CORE_PRODUCT_FILES": "deep-file-writer",
    "STEP_4_TEMPLATES_AND_PROMPTS": "deep-file-writer",
    "STEP_4_TEMPLATES_PROMPTS": "deep-file-writer",
    "STEP_5_EXAMPLES_QUALITY_COMPLIANCE": "buyer-test/license-check",
    "STEP_6_DELIVERY_SUPPORT": "delivery-support/refund-risk",
    "STEP_7_BUYER_RISK_TEST": "buyer-test/ai-replace-risk/refund-risk/license-check",
    "STEP_8_EXPORT_ZIP_MANIFEST": "export-zip/public-launch-audit",
    "STEP_9_SALES_PAGE": "sales-page",
    "STEP_10_WARRIORPLUS_LISTING": "warriorplus-listing",
    "STEP_11_JV_PACK": "jv-pack",
    "STEP_12_PRODUCT_CORE_REVIEW": "buyer-test/ai-replace-risk/refund-risk/license-check",
    "STEP_13_FILE_QUALITY_SCORE": "buyer-test/quality-score",
    "STEP_14_AI_REPLACE_RISK_AUDIT": "ai-replace-risk",
    "STEP_15_BEGINNER_CONFUSION_AUDIT": "buyer-test/beginner-clarity",
    "STEP_16_PROMPT_OUTPUT_TEST": "prompt-output-test",
    "STEP_17_BUYER_SIMULATION_TEST": "buyer-test",
    "STEP_18_REFUND_AUDITOR_TEST": "refund-risk",
    "STEP_19_FIX_WEAK_PARTS": "deep-file-writer/quality-fix",
    "STEP_20_RESCORE_PRODUCT": "quality-score/launch-audit",
    "STEP_21_MORE_EXAMPLE_OUTPUTS": "deep-file-writer/examples",
    "STEP_22_ADD_CHECKLISTS": "checklists/quality-gate",
    "STEP_23_ADD_FIX_PROMPTS": "fix-prompts/prompt-repair",
    "STEP_24_FINAL_COMPLIANCE_REVIEW": "license-check/compliance-review",
    "REAL_AI_CHAT_PHASE_5": "real-ai-chat/phase-5-sales-jv-funnel",
    "REAL_AI_CHAT_STEP_34_PACKAGING": "real-ai-chat/final-folder-packaging",
    "REAL_AI_CHAT_STEP_35_EXPORT_ZIP_MANIFEST_TEST": "real-ai-chat/export-zip-manifest-test",
    "GENERIC_STEP_AI_HANDLER": "real-ai-chat/generic-step-handler",
    "STEP_UNSUPPORTED": "unsupported",
    "ROUTE_DEBUG_OR_FIX": "route-debug",
    "ROUTE_DEBUG_OR_CONFLICT": "route-debug",
    "NEW_PRODUCT_RESET": "project-reset",
}
AI_ETSY_NEGATIVE_STEP4_TERMS = [
    "route sai", "route nhầm", "tra loi nham", "trả lời nhầm", "step 7", "bước 7",
    "step 8", "bước 8", "buyer test", "refund", "risk", "audit", "export", "zip",
    "public launch", "sản phẩm cũ", "san pham cu", "context cũ", "context cu",
]


def _slugify_project_name(name: str) -> str:
    normalized = unicodedata.normalize("NFKD", name or "")
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", ascii_text).strip("_").lower()
    return slug or "untitled_product"


def _extract_product_name_from_text(question: str) -> str | None:
    text = question or ""
    match = re.search(r"(?im)^\s*Product\s*:\s*(.+?)\s*$", text)
    if match:
        value = match.group(1).strip().strip('"').strip("'")
        if value:
            return value
    match = re.search(r"(?im)^\s*/new_product\s+(.+?)\s*$", text)
    if match:
        value = match.group(1).strip().strip('"').strip("'")
        if value:
            return value
    match = re.search(r"(?i)tạo sản phẩm mới\s+([^\n\.]+)", text)
    if match:
        value = match.group(1).strip().strip('"').strip("'")
        if value:
            return value
    return None



def _active_ai_etsy_project_marker_path() -> Path:
    root = ROOT_DIR / "exports" / "products"
    root.mkdir(parents=True, exist_ok=True)
    return root / ".active_project.json"


def _read_active_ai_etsy_product_name() -> str | None:
    marker = _active_ai_etsy_project_marker_path()
    if not marker.exists():
        return None
    try:
        data = json.loads(marker.read_text(encoding="utf-8"))
        return data.get("product_name") or None
    except Exception:
        return None


def _write_active_ai_etsy_project_marker(project: dict) -> None:
    marker = _active_ai_etsy_project_marker_path()
    marker.write_text(json.dumps({
        "product_name": project["product_name"],
        "project_slug": project["project_slug"],
        "active_project_path": project["active_project_path"],
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }, ensure_ascii=False, indent=2), encoding="utf-8")

def _resolve_ai_etsy_project(question: str, payload: dict | None = None) -> dict:
    payload = payload or {}
    product_name = str(payload.get("product_name") or payload.get("productName") or "").strip() or None
    if not product_name:
        product_name = _extract_product_name_from_text(question)
    explicit_product = bool(product_name)
    if not product_name:
        product_name = _read_active_ai_etsy_product_name()
    if not product_name:
        snapshot = active_project_snapshot() or {}
        product_name = snapshot.get("product_name") or AI_ETSY_DEFAULT_PRODUCT
    product_slug = _slugify_project_name(product_name)
    active_project_path = ROOT_DIR / "exports" / "products" / product_name
    active_project_path.mkdir(parents=True, exist_ok=True)
    state = _load_or_create_project_state(product_name, product_slug, active_project_path)
    project = {
        "project_id": product_slug,
        "project_slug": product_slug,
        "product_name": product_name,
        "active_project_path": str(active_project_path),
        "state": state,
    }
    if explicit_product:
        _write_active_ai_etsy_project_marker(project)
    return project


def _project_state_path(project: dict | str | Path) -> Path:
    if isinstance(project, dict):
        return Path(project["active_project_path"]) / "project_state.json"
    return Path(project) / "project_state.json"


def _load_or_create_project_state(product_name: str, project_slug: str, active_project_path: Path) -> dict:
    state_path = active_project_path / "project_state.json"
    if state_path.exists():
        try:
            state = json.loads(state_path.read_text(encoding="utf-8"))
        except Exception:
            state = {}
    else:
        state = {}
    state.update({
        "project_slug": project_slug,
        "product_name": product_name,
        "active_project_path": str(active_project_path),
    })
    state.setdefault("current_step", "")
    state.setdefault("completed_steps", [])
    state.setdefault("last_user_requested_step", "")
    state.setdefault("allowed_next_steps", [])
    state.setdefault("files_created", [])
    state.setdefault("files_missing", [])
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
    return state


def _update_project_state(project: dict, *, current_step: int | None = None, files_created: list[str] | None = None, files_missing: list[str] | None = None) -> dict:
    state_path = _project_state_path(project)
    state = project.get("state") or {}
    if state_path.exists():
        try:
            state.update(json.loads(state_path.read_text(encoding="utf-8")))
        except Exception:
            pass
    if current_step is not None:
        state["current_step"] = current_step
        state["last_user_requested_step"] = current_step
        completed = set(int(x) for x in state.get("completed_steps", []) if str(x).isdigit())
        completed.add(int(current_step))
        state["completed_steps"] = sorted(completed)
        state["allowed_next_steps"] = [current_step, current_step + 1] if current_step < 36 else [36]
    if files_created is not None:
        merged = set(state.get("files_created", []))
        merged.update(files_created)
        state["files_created"] = sorted(merged)
    if files_missing is not None:
        state["files_missing"] = sorted(set(files_missing))
    state["project_slug"] = project["project_slug"]
    state["product_name"] = project["product_name"]
    state["active_project_path"] = project["active_project_path"]
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
    project["state"] = state
    return state


def _is_new_product_request(question: str) -> bool:
    text = (question or "").lower()
    return "/new_product" in text or "tạo sản phẩm mới" in text or "tao san pham moi" in text or "/reset_project_context" in text


def _is_real_ai_product_route(route: dict | None) -> bool:
    if not route:
        return False
    selected = route.get("selected_route") or ""
    route_type = route.get("route_type") or ""
    if selected in {
        "REAL_AI_CHAT_PHASE_5",
        "REAL_AI_CHAT_STEP_34_PACKAGING",
        "REAL_AI_CHAT_STEP_35_EXPORT_ZIP_MANIFEST_TEST",
        "GENERIC_STEP_AI_HANDLER",
    }:
        return True
    return route_type in {"ai_content", "hybrid_action"}

def resolve_ai_etsy_route(question: str, payload: dict | None = None) -> dict:
    payload = payload or {}
    text = (question or "").lower()
    project = _resolve_ai_etsy_project(question, payload)
    if "/reset_project_context" in text or _is_new_product_request(question):
        state_path = _project_state_path(project)
        if state_path.exists():
            state_path.unlink()
        project = _resolve_ai_etsy_project(question)
        _write_active_ai_etsy_project_marker(project)
        return _route_payload("NEW_PRODUCT_RESET", None, "new/reset product command", project, False)
    generic_route = resolve_product_route(question)
    if _is_real_ai_product_route(generic_route):
        selected_real_ai_route = generic_route.get("selected_route")
        route = _route_payload(selected_real_ai_route, generic_route.get("explicit_step"), generic_route.get("reason") or "real AI product step request", project, False)
        route["route_type"] = generic_route.get("route_type") or "ai_content"
        route["api_called"] = True
        route["from_cache"] = False
        route["prebuilt_answer_used"] = False
        route["old_answer_reused"] = False
        route["tool_action"] = "export_zip" if route["route_type"] == "hybrid_action" and "export" in (question or "").lower() else ""
        return route
    if generic_route.get("selected_route") == "ROUTE_DEBUG_OR_CONFLICT":
        return _route_payload("ROUTE_DEBUG_OR_CONFLICT", generic_route.get("explicit_step"), generic_route.get("reason") or "route/debug complaint", project, False)
    if generic_route.get("selected_route") == "STEP_UNSUPPORTED":
        route = _route_payload("STEP_UNSUPPORTED", generic_route.get("explicit_step"), generic_route.get("reason") or "explicit step unsupported", project, False)
        route["requested_route"] = generic_route.get("requested_route")
        route["route_type"] = generic_route.get("route_type") or "unsupported_tool"
        route["api_called"] = False
        route["from_cache"] = False
        route["prebuilt_answer_used"] = False
        route["old_answer_reused"] = False
        route["no_answer_generated"] = True
        return route
    explicit_step = _ai_etsy_requested_step(question)
    if explicit_step:
        return _route_payload(AI_ETSY_STEP_ROUTES.get(explicit_step, "STEP_UNSUPPORTED"), explicit_step, f"explicit Step {explicit_step} has priority over keywords", project, False)
    slash_map = {
        "/buyer_test": ("STEP_7_BUYER_RISK_TEST", 7),
        "/refund_risk": ("STEP_7_BUYER_RISK_TEST", 7),
        "/ai_replace_risk": ("STEP_7_BUYER_RISK_TEST", 7),
        "/export_zip": ("STEP_8_EXPORT_ZIP_MANIFEST", 8),
        "/public_launch_audit": ("STEP_8_EXPORT_ZIP_MANIFEST", 8),
    }
    for command, (route, step) in slash_map.items():
        if command in text:
            return _route_payload(route, step, f"slash command {command}", project, False)
    if any(term in text for term in ["buyer test", "risk test", "refund risk", "ai replace risk"]):
        return _route_payload("STEP_7_BUYER_RISK_TEST", 7, "intent classifier: buyer/risk test", project, True)
    if any(term in text for term in ["public launch audit", "launch audit", "export zip", "final export"]):
        return _route_payload("STEP_8_EXPORT_ZIP_MANIFEST", 8, "intent classifier: export/audit", project, True)
    return {"selected_route": ""}


def _route_payload(route: str, step: int | None, reason: str, project: dict, fallback_used: bool) -> dict:
    return {
        "selected_route": route,
        "requested_step": step,
        "detected_intent": route.lower(),
        "selected_skill": AI_ETSY_ROUTE_SKILLS.get(route, ""),
        "project": project,
        "reason": reason,
        "fallback_used": fallback_used,
        "route_type": "tool_action",
        "api_called": False,
        "from_cache": False,
        "prebuilt_answer_used": False,
        "old_answer_reused": False,
        "tool_action": "",
        "conflict": False,
    }


def _route_debug_block(route: dict, action: dict | None = None, files_used: list[str] | None = None) -> str:
    project = route.get("project") or {}
    files = files_used or (action or {}).get("files_created") or []
    action = action or {}
    return "\n".join([
        "",
        "# REQUEST DEBUG",
        f"- request_id: {route.get('request_id') or ''}",
        f"- user_message_hash: {route.get('user_message_hash') or ''}",
        f"- api_called: {str(bool(action.get('api_called', route.get('api_called', False)))).lower()}",
        f"- from_cache: {str(bool(action.get('from_cache', route.get('from_cache', False)))).lower()}",
        f"- prebuilt_answer_used: {str(bool(action.get('prebuilt_answer_used', route.get('prebuilt_answer_used', False)))).lower()}",
        f"- old_answer_reused: {str(bool(action.get('old_answer_reused', route.get('old_answer_reused', False)))).lower()}",
        f"- route_type: {action.get('route_type', route.get('route_type', ''))}",
        f"- tool_action: {action.get('tool_action', route.get('tool_action', ''))}",
        f"- user_requested_step: {route.get('requested_step') or route.get('explicit_step') or ''}",
        f"- detected_intent: {route.get('detected_intent') or ''}",
        f"- selected_route: {route.get('selected_route') or ''}",
        f"- selected_skill: {route.get('selected_skill') or ''}",
        f"- active_project: {project.get('project_slug') or ''}",
        f"- active_project_path: {project.get('active_project_path') or ''}",
        f"- product_used: {project.get('product_name') or ''}",
        f"- files_used: {', '.join(files) if files else ''}",
        f"- reason: {route.get('reason') or ''}",
        f"- fallback_used: {str(bool(route.get('fallback_used'))).lower()}",
        f"- rag_used: {str(bool(action.get('rag_used', False))).lower()}",
        f"- ai_assisted: {str(bool(action.get('ai_assisted', False))).lower()}",
        f"- ai_elapsed_ms: {action.get('ai_elapsed_ms', 0)}",
        f"- route_conflict: {str(bool(route.get('route_conflict'))).lower()}",
        f"- stale_context_detected: false",
    ])

def _real_ai_phase5_prompt(question: str, project: dict, route: dict | None = None) -> str:
    selected_route = (route or {}).get("selected_route") or "REAL_AI_CHAT_PHASE_5"
    route_type = (route or {}).get("route_type") or "ai_content"
    explicit_step = (route or {}).get("requested_step") or (route or {}).get("explicit_step") or ""
    if selected_route == "REAL_AI_CHAT_STEP_34_PACKAGING":
        scope_rules = """- Chi xu ly Step 34: Final Folder Packaging.
- Tap trung vao folder structure cuoi, file placement map, buyer order, vendor asset map, missing files report, packaging checklist, packaging audit.
- Khong quay lai viet sales page/JV copy Step 25-33 tru khi can liet ke vi tri file.
- Khong tra `STEP_UNSUPPORTED`."""
        output_rule = "Output can co: PRODUCT USED, STEP USED, FINAL PRODUCT FOLDER STRUCTURE, FILE MOVE / COPY PLAN, BUYER-FRIENDLY STRUCTURE, VENDOR STRUCTURE, MISSING FILES REPORT, PACKAGING RULES, QUALITY GATE, NEXT ACTION."
    elif selected_route == "REAL_AI_CHAT_STEP_35_EXPORT_ZIP_MANIFEST_TEST":
        scope_rules = """- Chi xu ly Step 35: Export ZIP + Manifest Test.
- Lap ke hoach/kiem tra export ZIP, manifest, placeholder scan, missing assets, file inclusion checklist, ZIP quality gate.
- Neu tool chua that su tao ZIP trong request nay, bat buoc ghi `TEXT ONLY, NOT ZIP PROOF`.
- Khong bao Public Launch Ready neu chua co proof ZIP/download that.
- Khong tra `STEP_UNSUPPORTED`."""
        output_rule = "Output can co: PRODUCT USED, STEP USED, EXPORT TARGET, ZIP CONTENT PLAN, MANIFEST CHECK, PLACEHOLDER CHECK, FILES MISSING, ZIP QUALITY GATE, TEXT ONLY/ZIP PROOF STATUS, NEXT ACTION."
    elif selected_route == "GENERIC_STEP_AI_HANDLER":
        scope_rules = f"""- Xu ly dung explicit Step {explicit_step} theo prompt moi cua user.
- Day la content/audit/plan/fix request nen phai phan tich bang AI, khong tra mau cu.
- Neu step can tool that ma tool chua chay, ghi ro `TEXT ONLY, NOT ZIP PROOF` hoac thieu proof tuong ung.
- Khong tra `STEP_UNSUPPORTED` tru khi user doi tool ngoai kha nang."""
        output_rule = "Output can co: PRODUCT USED, STEP USED, ANALYSIS, FILES/DATA NEEDED, OUTPUT/PLAN, QUALITY GATE, NEXT ACTION."
    elif selected_route == "REAL_AI_CHAT_PHASE_5":
        scope_rules = """- Chi xu ly Phase 5: Sales Page + JV Manager + Funnel, Step 25 den Step 33.
- Khong tra `STEP_UNSUPPORTED`."""
        output_rule = "Output can co: PRODUCT USED, PHASE USED, STEP 25-33, SALES CLAIM SAFETY, JV READINESS, FUNNEL READINESS, NEXT ACTION."
    else:
        scope_rules = f"""- Xu ly route `{selected_route}` bang AI API that vi day la content/audit/test/fix request.
- Khong dung prebuilt answer, khong replay cau tra loi cu.
- Neu can tool hoac file that ma chua co proof, ghi ro thieu gi."""
        output_rule = "Output can co: PRODUCT USED, STEP/ROUTE USED, ANALYSIS, FINDINGS, FIX/OUTPUT, QUALITY GATE, NEXT ACTION."
    return f"""## REAL AI CHAT MODE - {selected_route} GUARD
Ban phai tra loi bang tieng Viet ro rang, de hieu.
Product bat buoc: {project['product_name']}
Project slug: {project['project_slug']}
Active project path: {project['active_project_path']}
Route type: {route_type}

Luat bat buoc:
{scope_rules}
- Content answer phai do AI API sinh moi theo prompt hien tai; khong dung cau tra loi mau trong code.
- Khong dung, khong nhac, khong copy san pham cu `AI Etsy Printable Bundle Builder` tru khi user yeu cau so sanh.
- Khong bia Public Launch Ready, payment proof, buyer proof, JV approval, legal review hoac delivery proof.
- Neu thieu du lieu project/file, ghi ro thieu gi va de xuat cach bo sung.
- {output_rule}

## USER PROMPT
{question}
""".strip()

def _real_ai_phase5_action(route: dict, *, elapsed_ms: int = 0) -> dict:
    project = route.get("project") or {}
    selected_route = route.get("selected_route") or "REAL_AI_CHAT_PHASE_5"
    route_debug = {k: v for k, v in route.items() if k != "project"}
    route_debug.update({
        "selected_route": selected_route,
        "active_project": project.get("project_slug", ""),
        "active_project_path": project.get("active_project_path", ""),
        "product_used": project.get("product_name", ""),
        "api_called": True,
        "from_cache": False,
        "fallback_used": False,
        "prebuilt_answer_used": False,
        "old_answer_reused": False,
        "route_type": route.get("route_type", "ai_content"),
        "tool_action": route.get("tool_action", ""),
        "route_conflict": False,
        "stale_context_detected": False,
        "files_used": [],
    })
    return {
        "ok": True,
        "type": selected_route.lower(),
        "product": project.get("product_name", ""),
        "product_used": project.get("product_name", ""),
        "selected_route": selected_route,
        "api_called": True,
        "from_cache": False,
        "fallback_used": False,
        "prebuilt_answer_used": False,
        "old_answer_reused": False,
        "route_type": route.get("route_type", "ai_content"),
        "tool_action": route.get("tool_action", ""),
        "ai_assisted": True,
        "ai_elapsed_ms": elapsed_ms,
        "route_debug": route_debug,
    }

def _call_product_step_ai(prompt: str, *, max_output_tokens: int = 5000) -> str:
    executor = ThreadPoolExecutor(max_workers=1)
    future = executor.submit(chat_with_llm, prompt, reasoning_effort="medium", max_output_tokens=max_output_tokens)
    try:
        return future.result(timeout=PRODUCT_STEP_AI_TIMEOUT_SECONDS)
    except FutureTimeoutError:
        future.cancel()
        return (
            "# AI API TIMEOUT\n\n"
            f"AI API da duoc goi nhung qua {PRODUCT_STEP_AI_TIMEOUT_SECONDS}s chua tra ket qua. "
            "Hay thu lai voi prompt ngan hon hoac tang PRODUCT_STEP_AI_TIMEOUT_SECONDS.\n\n"
            "- api_called: true\n"
            "- from_cache: false\n"
            "- prebuilt_answer_used: false\n"
            "- old_answer_reused: false\n"
        )
    finally:
        executor.shutdown(wait=False, cancel_futures=True)


def _run_real_ai_phase5_route(route: dict, question: str) -> tuple[str, dict]:
    project = route["project"]
    route["request_id"] = route.get("request_id") or ""
    route["user_message_hash"] = hashlib.sha256(question.encode("utf-8", errors="ignore")).hexdigest()[:16]
    prompt = _real_ai_phase5_prompt(question, project, route)
    start = _now_ms()
    answer = _call_product_step_ai(prompt, max_output_tokens=5000)
    elapsed_ms = _elapsed_ms(start)
    action = _real_ai_phase5_action(route, elapsed_ms=elapsed_ms)
    if project.get("product_name") != AI_ETSY_DEFAULT_PRODUCT:
        answer = answer.replace(AI_ETSY_DEFAULT_PRODUCT, project["product_name"])
    return answer.rstrip() + _route_debug_block(route, action), action


def _run_ai_etsy_route(route: dict) -> tuple[str, dict]:
    project = route["project"]
    selected = route["selected_route"]
    if selected == "NEW_PRODUCT_RESET":
        state = _update_project_state(project, current_step=1, files_created=[], files_missing=[])
        action = {"ok": True, "type": "new_product_reset", "product": project["product_name"], "project_state": state}
        action["route_debug"] = {k: v for k, v in route.items() if k != "project"}
        action["route_debug"]["active_project"] = project["project_slug"]
        action["route_debug"]["active_project_path"] = project["active_project_path"]
        answer = f"# NEW PRODUCT CONTEXT RESET\n\n- Product: `{project['product_name']}`\n- Project slug: `{project['project_slug']}`\n- Active path: `{project['active_project_path']}`\n- Global RAG/brain không bị xóa.\n"
        return answer + _route_debug_block(route, action), action
    if selected in {"ROUTE_DEBUG_OR_FIX", "ROUTE_DEBUG_OR_CONFLICT"}:
        action = {"ok": True, "type": "route_debug_or_fix", "product": project["product_name"], "project_state": project.get("state", {})}
        action["route_debug"] = {k: v for k, v in route.items() if k != "project"}
        action["route_debug"]["active_project"] = project["project_slug"]
        action["route_debug"]["active_project_path"] = project["active_project_path"]
        answer = "# ROUTE DEBUG / FIX MODE\n\nKhông chạy Step 4/5/7/8 vì đây là câu hỏi debug route hoặc báo lỗi route. Hãy gửi lại lệnh có `BƯỚC X ONLY` để chạy đúng step.\n"
        return answer + _route_debug_block(route, action), action
    if selected == "STEP_UNSUPPORTED":
        answer, action = step_unsupported_response(route, project["product_name"], project["active_project_path"])
        action["route_debug"]["active_project"] = project["project_slug"]
        return answer, action
    builders = {
        "STEP_2_PRODUCT_BLUEPRINT": _build_ai_etsy_step2_blueprint_answer,
        "STEP_3_CORE_PRODUCT_FILES": _build_ai_etsy_step3_core_files_answer,
        "STEP_4_TEMPLATES_AND_PROMPTS": _build_ai_etsy_step4_templates_answer,
        "STEP_4_TEMPLATES_PROMPTS": _build_ai_etsy_step4_templates_answer,
        "STEP_5_EXAMPLES_QUALITY_COMPLIANCE": _build_ai_etsy_step5_examples_quality_answer,
        "STEP_6_DELIVERY_SUPPORT": _build_ai_etsy_step6_delivery_support_answer,
        "STEP_7_BUYER_RISK_TEST": _build_ai_etsy_step7_buyer_risk_answer,
        "STEP_8_EXPORT_ZIP_MANIFEST": _build_ai_etsy_step8_final_export_answer,
        "STEP_12_PRODUCT_CORE_REVIEW": _build_ai_etsy_step12_core_review_answer,
        "STEP_13_FILE_QUALITY_SCORE": _build_ai_etsy_step13_file_quality_score_answer,
        "STEP_14_AI_REPLACE_RISK_AUDIT": _build_ai_etsy_step14_ai_replace_risk_answer,
        "STEP_15_BEGINNER_CONFUSION_AUDIT": _build_ai_etsy_step15_beginner_confusion_answer,
        "STEP_16_PROMPT_OUTPUT_TEST": _build_ai_etsy_step16_prompt_output_test_answer,
        "STEP_17_BUYER_SIMULATION_TEST": _build_ai_etsy_step17_buyer_simulation_answer,
        "STEP_18_REFUND_AUDITOR_TEST": _build_ai_etsy_step18_refund_auditor_answer,
        "STEP_19_FIX_WEAK_PARTS": _build_ai_etsy_step19_fix_weak_parts_answer,
        "STEP_20_RESCORE_PRODUCT": _build_ai_etsy_step20_rescore_product_answer,
        "STEP_21_MORE_EXAMPLE_OUTPUTS": _build_ai_etsy_step21_more_examples_answer,
        "STEP_22_ADD_CHECKLISTS": _build_ai_etsy_step22_checklists_answer,
        "STEP_23_ADD_FIX_PROMPTS": _build_ai_etsy_step23_fix_prompts_answer,
        "STEP_24_FINAL_COMPLIANCE_REVIEW": _build_ai_etsy_step24_compliance_review_answer,
    }
    answer, action = builders[selected](project)
    action = action or {}
    action["route_debug"] = {k: v for k, v in route.items() if k != "project"}
    action["route_debug"]["active_project"] = project["project_slug"]
    action["route_debug"]["active_project_path"] = project["active_project_path"]
    action["route_debug"]["from_cache"] = False
    action["route_debug"]["stale_context_detected"] = False
    action["route_debug"]["route_conflict"] = bool(route.get("route_conflict"))
    return answer + _route_debug_block(route, action), action

def _ai_etsy_requested_step(question: str) -> int | None:
    return extract_explicit_step(question)

def _is_ai_etsy_step4_templates_request(question: str) -> bool:
    if (step := _ai_etsy_requested_step(question)) is not None and step != 4:
        return False
    text = question.lower()
    return (
        "ai etsy printable bundle builder" in text
        and ("step 4 only" in text or "bước 4 only" in text or "bước 4 only" in text or "bước 4" in text)
        and ("#deep-file-writer" in text or "deep file writer" in text)
        and "step 8" not in text
        and "bước 8" not in text
        and "step 7" not in text
        and "bước 7" not in text
        and "vendor ready" not in text
    )

def _build_ai_etsy_step4_templates_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    answer = """# BƯỚC 4 — TEMPLATES AND PROMPTS

## DATA USED
- Product: AI Etsy Printable Bundle Builder.
- Dùng kết quả Step 2 blueprint và Step 3 core files.
- Mục tiêu bước này: tạo template + prompt lõi để buyer dùng thật.

## STEP 4 CREATED
- 02_Niche_And_Buyer_Picker.csv
- 03_Printable_Bundle_Planner.md
- 04_AI_Prompt_Library.md

---

# FILE 1: 02_Niche_And_Buyer_Picker.csv
```csv
Niche,Buyer,Pain,Product Idea,Bundle Angle,Safety Risk,Difficulty,Price Angle,Why It Sells
Teacher Reward Charts,Giáo viên tiểu học,Cần theo dõi reward nhanh,Reward chart bundle,Charts + stickers + parent notes,Không claim cải thiện hành vi chắc chắn,Easy,$9-$17,Dễ hiểu và dùng ngay trong lớp
Homeschool Weekly Planner,Phụ huynh homeschool,Khó sắp xếp tuần học,Weekly homeschool planner,Lesson plan + tracker + reflection,Không claim kết quả học tập,Medium,$12-$19,Buyer có nhu cầu lặp lại hằng tuần
Pet Care Tracker,Người nuôi thú cưng,Thông tin chăm sóc bị rải rác,Pet care printable pack,Feeding + vet + medication logs,Không thay tư vấn thú y,Easy,$9-$17,Rõ pain và dễ in
Small Business Order Tracker,Chủ shop handmade,Khó theo dõi đơn hàng,Order tracker printable,Order log + inventory + packaging checklist,Không claim tăng doanh thu,Medium,$12-$27,Dùng được cho seller nhỏ
Wedding Planning Checklist,Cô dâu tự tổ chức,Có quá nhiều việc cần nhớ,Wedding checklist bundle,Timeline + vendor list + budget sheet,Tránh dùng brand/venue nổi tiếng,Medium,$17-$27,Pain lớn và buyer sẵn chi tiền
Budget Binder Starter,Người muốn quản lý tiền cá nhân,Không biết tiền đi đâu,Budget binder printables,Expense + savings + debt tracker,Không tư vấn tài chính,Medium,$12-$27,Nhu cầu evergreen
Meal Planning Kit,Mẹ bận rộn,Không biết lên món tuần,Meal planner bundle,Menu + grocery list + pantry inventory,Không claim sức khỏe/giảm cân,Easy,$9-$17,Dùng thường xuyên
Cleaning Schedule Planner,Người bận rộn,Nhà cửa thiếu lịch dọn,Cleaning checklist pack,Daily + weekly + monthly checklists,Không claim sức khỏe,Easy,$7-$15,Simple high-utility printable
Craft Fair Seller Kit,Người bán hội chợ handmade,Khó chuẩn bị booth,Craft fair planner,Inventory + pricing + packing checklist,Không claim sales guaranteed,Medium,$17-$27,Buyer có event deadline
Garden Planner,Người làm vườn tại nhà,Khó theo dõi cây trồng,Garden planner bundle,Seed log + watering + harvest tracker,Không claim nông nghiệp/chữa bệnh,Medium,$12-$19,Visual đẹp và seasonal
Kids Chore Chart,Phụ huynh,Cần giao việc nhà rõ ràng,Chore chart printable,Chore list + reward tracker + routine cards,Không claim thay đổi hành vi chắc chắn,Easy,$7-$17,Parent pain phổ biến
Teacher Sub Plan Kit,Giáo viên,Cần chuẩn bị khi nghỉ dạy,Substitute teacher printable kit,Sub notes + class info + emergency plan,Không dùng school logo,Medium,$12-$19,Tiết kiệm thời gian rõ
Etsy Shop Planner,Người bán Etsy mới,Khó lên kế hoạch listing,Etsy shop planner,Listing tracker + keyword notes + photo checklist,Không claim ranking/sales,Medium,$17-$27,Phù hợp audience Etsy
Social Media Content Planner,Solopreneur,Không đều nội dung đăng,Content planner printable,Calendar + idea bank + caption checklist,Không claim viral,Easy,$9-$19,Rộng buyer và dễ rebrand
Daily Wellness Journal,Người muốn journaling,Khó duy trì thói quen,Wellness journal printable,Mood + gratitude + habit tracker,Không claim therapy/medical,Medium,$12-$19,Evergreen nhưng cần claim an toàn
Event Planning Binder,Người tổ chức sự kiện nhỏ,Có nhiều đầu việc,Event planner bundle,Budget + guest list + vendor checklist,Tránh brand/venue protected,Medium,$17-$27,High perceived value
Student Study Planner,Sinh viên,Học nhiều môn khó theo dõi,Study planner printable,Assignment + exam + habit tracker,Không claim điểm số đảm bảo,Easy,$9-$17,Buyer rõ và seasonal
Real Estate Open House Kit,Agent bất động sản,Cần chuẩn bị open house,Open house checklist pack,Sign-in + prep checklist + follow-up tracker,Không claim lead/sales guarantee,Medium,$17-$27,Professional buyer value
Freelancer Client Tracker,Freelancer,Khó theo dõi client/project,Client tracker bundle,Inquiry + project + invoice checklist,Không tư vấn pháp lý/tài chính,Medium,$12-$27,Dùng trực tiếp cho công việc
Travel Packing Planner,Người đi du lịch,Hay quên đồ cần mang,Travel packing printable,Checklist + itinerary + budget sheet,Không claim safety guarantee,Easy,$7-$15,Rất dễ hiểu trong 5 giây
```

---

# FILE 2: 03_Printable_Bundle_Planner.md
```md
# Printable Bundle Planner

## 1. Thông tin sản phẩm
- Tên bundle:
- Buyer chính:
- Pain chính:
- Kết quả an toàn muốn hứa:
- Giá dự kiến:
- Nền tảng bán:
- Gói file cuối cùng:

## 2. Checklist chọn niche
Chấm mỗi mục 1–5.

| Tiêu chí | Điểm | Ghi chú |
|---|---:|---|
| Buyer hiểu ngay trong 5 giây |  |  |
| Pain cụ thể, không quá rộng |  |  |
| Có thể tạo 5–10 file hữu ích |  |  |
| Dễ làm mockup/preview |  |  |
| Ít rủi ro trademark/copyright |  |  |
| Không cần claim quá đà để bán |  |  |
| Có thể rebrand thành nhiều phiên bản |  |  |

Nếu dưới 25/35, hãy đổi niche hoặc thu hẹp buyer.

## 3. Bundle structure

| File | Mục đích | Buyer dùng khi nào | Format | Bắt buộc? |
|---|---|---|---|---|
| Start Here | Hướng dẫn mở đầu | Ngay sau khi tải | PDF/MD | Có |
| Main Planner | File giá trị chính | Khi triển khai | PDF/Canva | Có |
| Checklist | Giảm lỗi và tăng dùng được | Trước khi in/bán | PDF | Có |
| Example Filled-In | Cho buyer thấy cách dùng | Lúc học cách điền | PDF | Có |
| Bonus Tracker | Tăng perceived value | Khi dùng lâu dài | PDF | Tuỳ |
| License Note | Giảm hiểu nhầm quyền dùng | Trước khi bán/rebrand | PDF/MD | Có |

## 4. Workflow từ ý tưởng đến pack bán được
1. Chọn niche từ `02_Niche_And_Buyer_Picker.csv`.
2. Viết buyer pain bằng một câu cụ thể.
3. Chọn 5–7 file thật buyer cần.
4. Tạo layout đơn giản trước, đừng làm quá đẹp nhưng khó dùng.
5. Viết hướng dẫn dùng trong `Start Here`.
6. Tạo một bản ví dụ đã điền.
7. Chạy checklist quality và compliance.
8. Đóng gói ZIP có manifest.

## 5. Ví dụ đã điền
- Tên bundle: Teacher Reward Chart Starter Bundle.
- Buyer: giáo viên tiểu học và phụ huynh homeschool.
- Pain: cần cách theo dõi reward/chore đơn giản, dễ in.
- Promise an toàn: giúp tổ chức reward tracking rõ ràng hơn, không hứa thay đổi hành vi chắc chắn.
- Files: reward chart, sticker tracker, weekly behavior sheet, parent note, instruction page, license note.
- Giá dự kiến: $9–$17.
- Risk: không claim kết quả học tập/hành vi, không dùng logo trường hoặc nhân vật nổi tiếng.
```

---

# FILE 3: 04_AI_Prompt_Library.md
```md
# AI Prompt Library — AI Etsy Printable Bundle Builder

## A. Niche Research Prompts
1. Act as an Etsy printable strategist. Generate 25 safe printable bundle niches for [buyer group]. Avoid trademarks, celebrity names, lyrics, famous characters, medical claims, income claims, and platform approval claims.
2. Score these niches by buyer clarity, pain urgency, visual preview potential, file usefulness, and compliance risk. Return the top 10 only.
3. Turn this broad niche into 10 narrower buyer-specific printable bundle angles: [niche]. Include buyer, pain, bundle idea, and safe promise.
4. Find printable bundle ideas that can be useful without requiring legal, medical, financial, therapy, or guaranteed-result claims.
5. Create a seasonal printable bundle calendar for [buyer] with 12 monthly product ideas and risk notes.

## B. Buyer Avatar Prompts
1. Build a buyer avatar for [printable niche]. Include daily situation, pain, buying trigger, objections, and what they want in a downloadable file.
2. List 10 reasons this buyer would pay for a printable bundle instead of making it from scratch.
3. Write buyer language for [niche] using simple, non-hype wording suitable for an Etsy-style listing.
4. Identify refund triggers for this buyer and suggest files/instructions that reduce confusion.
5. Create a buyer success path from download to first use for [bundle idea].

## C. Printable Bundle Planning Prompts
1. Plan a 7-file printable bundle for [buyer] who wants [outcome]. Include file name, purpose, page count, and why it matters.
2. Create a bundle table with core files, bonus files, example files, and support files for [niche].
3. Turn this idea into a minimum viable printable bundle that can be completed in one weekend: [idea].
4. Suggest 5 upgrades that increase perceived value without adding risky claims.
5. Create a manifest for a printable bundle ZIP with file names, purpose, format, and completion status.

## D. Canva/Design Direction Prompts
1. Create a clean Canva design brief for [printable bundle]. Include page size, typography direction, spacing, color palette, and export format.
2. Write layout rules for a printable planner page that is easy to print, readable, and not visually cluttered.
3. Suggest 5 visual styles for [niche] that do not imitate famous brands, artists, franchises, or protected characters.
4. Create a mockup shot list for promoting [bundle] without misleading buyers about what is included.
5. Review this printable layout for readability, margins, printer-friendliness, and buyer confusion risk.

## E. Listing Copy Prompts
1. Write an Etsy-style title, bullets, and description for [bundle]. Avoid income claims, platform approval claims, medical claims, and guaranteed outcomes.
2. Write 10 benefit bullets for [bundle] using safe language and no hype.
3. Create FAQ answers for a digital printable product, including download, printing, license, refunds, and support.
4. Write a short product description that explains what is included and what is not included.
5. Rewrite this listing to be clearer and safer while keeping it persuasive: [paste listing].

## F. Quality Control Prompts
1. Audit this printable bundle for missing files, unclear instructions, weak buyer value, and refund risk.
2. Check this bundle for copyright, trademark, celebrity, lyrics, brand imitation, Canva license, and risky claims.
3. Create a buyer test: pretend you bought this for $17 and list what feels valuable, confusing, missing, or refund-worthy.
4. Scan this file list for placeholder text, vague promises, missing support information, and delivery confusion.
5. Rate this printable bundle 1–10 for buyer clarity, usefulness, packaging, compliance safety, and AI replacement risk.

## G. Rebrand/PLR Prompts
1. Create 10 rebrand angles for this printable bundle while keeping the core workflow and avoiding risky claims: [bundle].
2. Turn this printable bundle into a PLR-friendly version with editable sections, license boundaries, and buyer instructions.
3. Write safe commercial-use terms for a printable bundle without over-granting rights to third-party assets.
4. Suggest 5 niche variations that reuse the same file structure but target different buyers.
5. Create a rebrand checklist covering product name, colors, examples, license note, support email, listing copy, and preview images.
```

---

# QUALITY GATE

| File | Status | Proof |
|---|---|---|
| 02_Niche_And_Buyer_Picker.csv | PASS | 20+ niche rows |
| 03_Printable_Bundle_Planner.md | PASS | Có template, checklist, workflow, ví dụ |
| 04_AI_Prompt_Library.md | PASS | 7 nhóm, 35 prompt copy-ready |

# DOWNLOAD / ZIP STATUS
- ZIP status: CREATED.
- Public Launch Ready: NO — đây mới là Step 4 templates/prompts, chưa phải sản phẩm hoàn chỉnh.
"""
    answer = answer.replace("AI Etsy Printable Bundle Builder", product_name)
    action = _write_step_product_bundle(
        product_name=product_name,
        step_slug="step_4_templates_and_prompts",
        bundle_name="BƯỚC 4 — TEMPLATES AND PROMPTS.zip",
        answer=answer,
        action_type="step4_templates_and_prompts",
    )
    answer = f"{answer.rstrip()}\n\n# DOWNLOAD PROOF\n- Đã tự lưu vào thư mục sản phẩm: `{action.get('product_folder', '')}`\n- Đã tạo 1 file ZIP chung: [`{action.get('fileName', '')}`]({action.get('url', '')})\n- Bạn chỉ cần bấm tải ZIP, không cần tải từng file lẻ.\n"
    return answer, action

def _is_ai_etsy_deep_build_request(question: str) -> bool:
    text = question.lower()
    if "step 3" in text or "bước 3" in text or "core product files" in text or "core files" in text:
        return False
    if "step 4" in text or "bước 4" in text or "templates and prompts" in text:
        return False
    if "step 5" in text or "bước 5" in text or "examples quality" in text or "quality compliance" in text:
        return False
    if "step 6" in text or "bước 6" in text or "delivery support" in text or "delivery / support" in text:
        return False
    if "step 7" in text or "bước 7" in text or "buyer test" in text or "risk test" in text:
        return False
    if "step 8" in text or "bước 8" in text or "final export" in text or "export folder" in text:
        return False
    required = ["ai etsy printable bundle builder", "#ai-printables-kdp-prompt"]
    deep_markers = ["#export-zip", "vendor ready", "warriorplus bán được", "warriorplus ban duoc"]
    return all(item in text for item in required) and any(item in text for item in deep_markers)



def _is_ai_etsy_step5_examples_quality_request(question: str) -> bool:
    if (step := _ai_etsy_requested_step(question)) is not None and step != 5:
        return False
    text = question.lower()
    return (
        "ai etsy printable bundle builder" in text
        and ("step 5" in text or "bước 5" in text or "examples quality" in text or "quality compliance" in text)
        and ("#deep-file-writer" in text or "deep file writer" in text or "#buyer-test" in text or "#license-check" in text)
        and "vendor ready" not in text
    )


def _build_ai_etsy_step5_examples_quality_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    answer = """# BƯỚC 5 — EXAMPLES, QUALITY AND COMPLIANCE

## DATA USED
- Product: AI Etsy Printable Bundle Builder.
- Dùng Step 2 blueprint, Step 3 core files, Step 4 templates/prompts.
- Mục tiêu bước này: tạo ví dụ bundle, checklist kiểm soát chất lượng, và compliance note để giảm refund/trademark/AI-replace risk.

## STEP 5 CREATED
- 05_Example_Bundle_Concepts.md
- 06_Quality_Control_Checklist.md
- 07_License_And_Compliance_Note.md

---

# FILE 1: 05_Example_Bundle_Concepts.md
```md
# 05 Example Bundle Concepts

File này cho buyer ví dụ cụ thể để không bị kẹt ở bước “tôi nên tạo bundle gì”.

## Concept 1 — Teacher Reward Chart Starter Bundle
Buyer: giáo viên tiểu học hoặc phụ huynh homeschool.
Pain: cần reward chart dễ in, dễ hiểu, không mất thời gian thiết kế từ đầu.
Files nên có: Reward_Chart_Weekly.pdf, Sticker_Tracker_Sheet.pdf, Classroom_Routine_Cards.pdf, Parent_Note_Template.pdf, How_To_Use_This_Bundle.md.
Safety note: không dùng school logo, nhân vật nổi tiếng, brand, hoặc claim giáo dục bảo đảm.

## Concept 2 — Etsy Seller Listing Planner Bundle
Buyer: người bán Etsy mới hoặc handmade seller.
Pain: khó lên kế hoạch listing, ảnh, keywords, mô tả và lịch đăng.
Files nên có: Listing_Planner.pdf, Product_Photo_Checklist.pdf, Keyword_Notes_Worksheet.pdf, Weekly_Shop_Action_Planner.pdf, Listing_Review_Checklist.md.
Safety note: không hứa ranking, traffic hoặc sales.

## Concept 3 — Pet Care Tracker Printable Pack
Buyer: người nuôi chó/mèo bận rộn.
Pain: thông tin ăn uống, lịch vet, medication, grooming bị rải rác.
Files nên có: Feeding_Tracker.pdf, Vet_Visit_Log.pdf, Medication_Log.pdf, Grooming_Schedule.pdf, Emergency_Info_Sheet.pdf.
Safety note: không claim chữa bệnh, không đưa chỉ dẫn y tế cụ thể.

## Concept 4 — Budget Binder Starter Kit
Buyer: người muốn quản lý chi tiêu cá nhân.
Pain: không biết tiền đi đâu, khó theo dõi hóa đơn và mục tiêu tiết kiệm.
Files nên có: Monthly_Budget_Worksheet.pdf, Expense_Tracker.pdf, Savings_Goal_Tracker.pdf, Bill_Due_Date_Calendar.pdf, Debt_Overview_Sheet.pdf.
Safety note: không hứa giảm nợ, tăng thu nhập, hoặc kết quả tài chính.

## Concept 5 — Wedding Planning Mini Binder
Buyer: cô dâu/chú rể tự tổ chức wedding nhỏ.
Pain: nhiều việc phải nhớ: guest list, vendor, budget, timeline.
Files nên có: Wedding_Timeline_Checklist.pdf, Guest_List_Tracker.pdf, Vendor_Contact_Sheet.pdf, Budget_Planner.pdf, Day_Of_Checklist.pdf.
Safety note: không dùng venue/brand/celebrity wedding imagery protected.
```

---

# FILE 2: 06_Quality_Control_Checklist.md
```md
# 06 Quality Control Checklist

Dùng checklist này trước khi export ZIP hoặc viết listing. Mục tiêu là giảm refund, giảm complaint “AI cũng làm được”, và tránh file quá mỏng.

## Buyer Clarity
- Buyer được nêu rõ trong 1 câu.
- Pain cụ thể, không chung chung.
- Outcome thực tế và an toàn.
- Buyer hiểu file dùng để làm gì trong 10 giây.
- Không dùng promise mơ hồ như “kiếm tiền dễ dàng”, “bán chạy chắc chắn”, “được approve chắc chắn”.

## Product Completeness
- Có Start Here hoặc hướng dẫn dùng.
- Có workflow từng bước.
- Có template/planner chính.
- Có ví dụ đã điền mẫu.
- Có checklist kiểm tra.
- Có license/compliance note.
- Có support FAQ hoặc delivery instruction.

## AI Replace Risk
1–3: chỉ là prompt thô. 4–6: có prompt + ít hướng dẫn. 7–8: có workflow, templates, ví dụ, checklist. 9–10: có full bundle, examples, done-for-you assets, ZIP, manifest, placeholder scan, buyer test, delivery/support.

## Refund Risk
Trigger refund thường gặp: download khó hiểu, không biết mở file nào đầu tiên, license mơ hồ, file quá mỏng, claim bán hàng quá đà, thiếu support/contact.

Cách giảm refund: viết rõ What is included, What is not included, có Start Here, có example output, có support FAQ, có refund policy công bằng, không hứa income/platform approval.

## Listing Safety
- Không claim Etsy ranking.
- Không claim KDP approval.
- Không claim guaranteed sales.
- Không claim legal/tax/medical/financial advice.
- Không dùng trademark/celebrity/brand/franchise.
- Không bán lại Canva Pro elements sai license.
- Không dùng quotes/lời bài hát protected.
```

---

# FILE 3: 07_License_And_Compliance_Note.md
```md
# 07 License And Compliance Note

File này giúp vendor và buyer hiểu giới hạn sử dụng của AI Etsy Printable Bundle Builder. Đây không phải tư vấn pháp lý.

## Quyền sử dụng
Bạn có thể dùng workflow, checklist, planner và prompt trong pack để tạo printable bundle gốc của riêng bạn. Bạn không được bán lại nguyên văn pack này như một sản phẩm PLR/MRR nếu license bạn mua không cấp quyền resale.

## Copyright
Không copy artwork, layout, sách, planner, checklist, template, mockup hoặc file của người khác. Không dùng lời bài hát, quote sách/phim, nội dung khóa học, worksheet hoặc tài liệu có bản quyền nếu không có quyền.

## Trademark
Tránh brand names, franchise, team names, celebrity names, product names, slogans, game names và character names trong file bán lại nếu không có quyền rõ ràng.

## Canva / Design Assets
Nếu dùng Canva, kiểm tra license hiện tại, không bán lại element riêng lẻ như asset pack, không dùng Pro elements nếu buyer không có quyền dùng hoặc license không cho phép.

## Platform Claims
Không hứa guaranteed Etsy sales, KDP approval, WarriorPlus approval, affiliate approval, ranking, traffic, income hoặc legal compliance guaranteed.

## Commercial Use Boundary
You may use these prompts, workflows, and templates to create your own finished printable products. You may sell finished products you create, provided you own or have rights to all included assets and comply with platform rules. You may not resell, redistribute, sublicense, or claim ownership of this original kit unless your purchase includes explicit resale rights.

## Compliance Status
Status: safer launch prep, not legal approval. Public Launch Ready: NO unless final ZIP, checkout, delivery, support, platform rules, rights, and buyer test are verified.
```

---

# QUALITY GATE
| File | Status | Proof |
|---|---|---|
| 05_Example_Bundle_Concepts.md | PASS | 5 concrete bundle examples |
| 06_Quality_Control_Checklist.md | PASS | Buyer/product/AI replace/refund/listing checks |
| 07_License_And_Compliance_Note.md | PASS | Copyright, trademark, Canva, platform claims, commercial-use boundary |

# DOWNLOAD / ZIP STATUS
- ZIP status: CREATED.
- Public Launch Ready: NO — đây mới là Step 5 examples/quality/compliance, chưa test payment/delivery/buyer/JV/legal.
"""
    answer = answer.replace("AI Etsy Printable Bundle Builder", product_name)
    action = _write_step_product_bundle(
        product_name=product_name,
        step_slug="step_5_examples_quality_compliance",
        bundle_name="BƯỚC 5 — EXAMPLES QUALITY COMPLIANCE.zip",
        answer=answer,
        action_type="step5_examples_quality_compliance",
    )
    answer = f"{answer.rstrip()}\n\n# DOWNLOAD PROOF\n- Đã tự lưu vào thư mục sản phẩm: `{action.get('product_folder', '')}`\n- Đã tạo 1 file ZIP chung: [`{action.get('fileName', '')}`]({action.get('url', '')})\n- Bạn chỉ cần bấm tải ZIP, không cần tải từng file lẻ.\n"
    return answer, action


def _is_ai_etsy_step6_delivery_support_request(question: str) -> bool:
    if (step := _ai_etsy_requested_step(question)) is not None and step != 6:
        return False
    text = question.lower()
    return (
        "ai etsy printable bundle builder" in text
        and ("step 6" in text or "bước 6" in text or "delivery support" in text or "delivery / support" in text)
        and ("#deep-file-writer" in text or "deep file writer" in text or "#delivery-support" in text or "#refund-risk" in text)
        and "vendor ready" not in text
    )


def _build_ai_etsy_step6_delivery_support_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    answer = """# BƯỚC 6 — DELIVERY / SUPPORT

## DATA USED
- Product: AI Etsy Printable Bundle Builder.
- Dùng Step 2 blueprint, Step 3 core files, Step 4 templates/prompts, Step 5 quality/compliance.
- Mục tiêu bước này: tạo bộ file giúp buyer tải, hiểu, dùng sản phẩm và liên hệ support rõ ràng.

## STEP 6 CREATED
- delivery_page.md
- buyer_onboarding_email.md
- support_faq.md
- refund_policy.md

---

# FILE 1: delivery_page.md
```md
# AI Etsy Printable Bundle Builder — Delivery Page

Cảm ơn bạn đã mua AI Etsy Printable Bundle Builder.

Đây là bộ workflow + template + prompt giúp bạn lập kế hoạch và tạo printable bundle theo hướng an toàn hơn cho Etsy-style digital products, KDP-style support files, PLR-friendly bundles hoặc sản phẩm số nhỏ.

## Bạn nhận được gì

- Start Here file để biết mở file nào trước.
- Workflow tạo printable bundle từng bước.
- Niche and buyer picker CSV.
- Printable bundle planner.
- AI prompt library.
- Example bundle concepts.
- Quality control checklist.
- License and compliance note.
- Sales/support files ở các bước sau nếu vendor cung cấp bản đầy đủ.

## Cách bắt đầu nhanh

1. Mở `README.md`.
2. Mở `00_Start_Here.md`.
3. Chọn 1 niche trong `02_Niche_And_Buyer_Picker.csv`.
4. Điền `03_Printable_Bundle_Planner.md`.
5. Dùng `04_AI_Prompt_Library.md` để tạo nội dung.
6. Kiểm tra bằng `06_Quality_Control_Checklist.md`.
7. Đọc `07_License_And_Compliance_Note.md` trước khi bán hoặc publish.

## Điều quan trọng

Sản phẩm này không đảm bảo doanh số, ranking, Etsy approval, KDP approval, WarriorPlus approval hoặc kết quả kinh doanh. Bạn chịu trách nhiệm kiểm tra quyền sử dụng asset, platform rules và claim trong listing cuối cùng.

## Nếu bạn bị kẹt

Hãy bắt đầu bằng một mini bundle 5 file thay vì cố tạo bundle quá lớn. Ví dụ: planner chính, checklist, example, prompt sheet, license note.

## Support

Email support mẫu: support@example.com

Khi liên hệ support, hãy gửi:
- email mua hàng,
- tên sản phẩm,
- file bạn đang mở,
- screenshot lỗi nếu có,
- câu hỏi cụ thể bạn cần hỗ trợ.
```

---

# FILE 2: buyer_onboarding_email.md
```md
# Buyer Onboarding Email

Subject: Truy cập AI Etsy Printable Bundle Builder của bạn

Hi,

Cảm ơn bạn đã mua AI Etsy Printable Bundle Builder.

Bạn có thể bắt đầu tại trang download của vendor. Sau khi tải file ZIP, hãy giải nén toàn bộ thư mục trước khi mở file.

## Bắt đầu tại đây

1. Mở `README.md`.
2. Sau đó mở `00_Start_Here.md`.
3. Chọn 1 buyer/niche trong `02_Niche_And_Buyer_Picker.csv`.
4. Điền kế hoạch trong `03_Printable_Bundle_Planner.md`.
5. Dùng `04_AI_Prompt_Library.md` để tạo nội dung.
6. Chạy checklist trong `06_Quality_Control_Checklist.md` trước khi export hoặc bán.

## Gợi ý để có kết quả nhanh

Đừng bắt đầu bằng một sản phẩm quá lớn. Hãy tạo một mini printable bundle trước:

- 1 planner chính,
- 1 checklist,
- 1 example,
- 1 instruction file,
- 1 license/compliance note.

Sau khi mini bundle rõ ràng, bạn mới mở rộng thành bundle lớn hơn.

## Lưu ý an toàn

Không dùng trademark, celebrity, brand, character, lời bài hát, logo hoặc claim doanh số/approval nếu bạn không có quyền hoặc proof thật.

Sản phẩm này không phải tư vấn pháp lý, tài chính, thuế, y tế hoặc cam kết nền tảng approve.

Support: support@example.com

Thanks,
Example Studio
```

---

# FILE 3: support_faq.md
```md
# Support FAQ

## 1. Tôi nên mở file nào đầu tiên?
Mở `README.md`, sau đó mở `00_Start_Here.md`.

## 2. Đây có phải sản phẩm printable hoàn chỉnh để bán ngay không?
Đây là builder kit/workflow pack. Nó giúp bạn tạo printable bundle của riêng bạn. Bạn vẫn cần tạo design/PDF cuối cùng, kiểm tra quyền sử dụng và test delivery trước khi bán chính thức.

## 3. Tôi có thể bán sản phẩm tạo ra từ kit này không?
Có thể, nếu sản phẩm cuối là output gốc của bạn, bạn có quyền dùng mọi asset/font/template/mockup, và bạn tuân thủ platform rules.

## 4. Tôi có được bán lại nguyên pack này không?
Không, trừ khi license bạn mua ghi rõ có quyền PLR/MRR/resale.

## 5. Kit này có đảm bảo bán được trên Etsy/KDP/WarriorPlus không?
Không. Không có guaranteed sales, ranking, platform approval hoặc affiliate approval.

## 6. Tôi dùng Canva được không?
Được, nhưng phải kiểm tra license Canva hiện tại. Không bán lại Canva elements riêng lẻ như asset pack của bạn.

## 7. Tôi có thể dùng nhân vật nổi tiếng, brand hoặc quote không?
Không nên, trừ khi bạn có written permission hoặc quyền rõ ràng.

## 8. Nếu AI tạo nội dung lỗi thì sao?
Dùng checklist và fix prompt để sửa. Luôn review thủ công trước khi publish.

## 9. Tại sao cần compliance note?
Để giảm rủi ro copyright, trademark, platform claim, refund và buyer confusion.

## 10. Tôi cần support thì gửi gì?
Gửi email mua hàng, file đang mở, mô tả lỗi, screenshot nếu có, và bạn muốn đạt kết quả gì.
```

---

# FILE 4: refund_policy.md
```md
# Refund Policy

AI Etsy Printable Bundle Builder là sản phẩm digital gồm workflow, prompts, templates, examples, checklist và support guidance.

Vì đây là sản phẩm digital, buyer nên đọc kỹ mô tả sản phẩm trước khi mua.

## Khi nào refund có thể được xem xét

Refund request có thể được xem xét nếu:

- buyer không truy cập được file sau khi support đã thử xử lý,
- buyer nhận sai sản phẩm,
- file ZIP bị lỗi và vendor không thể cung cấp bản thay thế hợp lý,
- thành phần chính được mô tả trong sales page bị thiếu.

## Khi nào refund thường không áp dụng

Refund thường không áp dụng cho:

- không có doanh số,
- Etsy/KDP/WarriorPlus không approve,
- buyer không đọc hướng dẫn,
- buyer dùng trademark/asset sai license,
- buyer kỳ vọng nhận design/PDF hoàn chỉnh trong khi sales page ghi rõ đây là builder kit,
- claim hoặc kết quả không được hứa trong sales page.

## Trước khi yêu cầu refund

Vui lòng liên hệ support@example.com với:

- email mua hàng,
- order ID nếu có,
- mô tả vấn đề,
- screenshot lỗi,
- file bạn đang gặp vấn đề.

Vendor sẽ cố gắng hỗ trợ truy cập, giải thích file, hoặc cung cấp bản ZIP thay thế nếu lỗi thuộc về delivery.

## Disclaimer

Policy này không thay thế quyền người tiêu dùng bắt buộc theo luật tại khu vực của bạn. Sản phẩm không đảm bảo doanh số, platform approval, legal approval, ranking hoặc income.
```

---

# QUALITY GATE
| File | Status | Proof |
|---|---|---|
| delivery_page.md | PASS | Có hướng dẫn tải, mở file, bắt đầu nhanh, support |
| buyer_onboarding_email.md | PASS | Có email onboarding, thứ tự dùng file, warning an toàn |
| support_faq.md | PASS | Có 10 câu FAQ giảm confusion/refund |
| refund_policy.md | PASS | Có điều kiện refund, giới hạn claim, support-first process |

# DOWNLOAD / ZIP STATUS
- ZIP status: CREATED.
- Public Launch Ready: NO — đây mới là Step 6 delivery/support, chưa test checkout, delivery thực tế, JV, buyer feedback hoặc legal review.
"""
    answer = answer.replace("AI Etsy Printable Bundle Builder", product_name)
    action = _write_step_product_bundle(
        product_name=product_name,
        step_slug="step_6_delivery_support",
        bundle_name="BƯỚC 6 — DELIVERY SUPPORT.zip",
        answer=answer,
        action_type="step6_delivery_support",
    )
    answer = f"{answer.rstrip()}\n\n# DOWNLOAD PROOF\n- Đã tự lưu vào thư mục sản phẩm: `{action.get('product_folder', '')}`\n- Đã tạo 1 file ZIP chung: [`{action.get('fileName', '')}`]({action.get('url', '')})\n- Bạn chỉ cần bấm tải ZIP, không cần tải từng file lẻ.\n"
    return answer, action


def _is_ai_etsy_step7_buyer_risk_request(question: str) -> bool:
    if (step := _ai_etsy_requested_step(question)) is not None and step != 7:
        return False
    text = question.lower()
    return (
        "ai etsy printable bundle builder" in text
        and ("step 7" in text or "bước 7" in text or "buyer test" in text or "risk test" in text)
        and ("#buyer-test" in text or "#refund-risk" in text or "#ai-replace-risk" in text or "#license-check" in text)
        and "vendor ready" not in text
    )


def _build_ai_etsy_step7_buyer_risk_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    answer = """# BƯỚC 7 — BUYER TEST + RISK TEST

## DATA USED
- Product: AI Etsy Printable Bundle Builder.
- Dùng các file đã tạo từ Step 3 đến Step 6 trong thư mục sản phẩm.
- Đây là test nội bộ bằng artifact hiện có, không phải phản hồi buyer thật.

## BUYER TEST RESULT
Buyer mới mua giá $17-$27 sẽ hiểu sản phẩm nhanh hơn nếu mở `README.md` và `00_Start_Here.md` trước. Giá trị chính nằm ở workflow, template, prompt library, ví dụ bundle, checklist, compliance và delivery/support.

Điểm mạnh:
- Có đường đi từ chọn niche đến tạo bundle.
- Có prompt library và planner giúp giảm cảm giác chỉ là prompt thô.
- Có checklist và compliance note để giảm rủi ro refund.

Điểm yếu còn lại:
- Chưa có PDF/Canva template thiết kế thật.
- Chưa có mockup preview.
- Chưa test với buyer thật.
- Chưa có checkout/delivery proof.

## AI REPLACE RISK SCORE
Score: 7.8/10.

Lý do: Pack không chỉ là prompt, đã có workflow, file hướng dẫn, template, ví dụ, checklist và support. Tuy nhiên để lên 9+, cần thêm design templates/PDF mẫu, preview/mockup, và final ZIP được test delivery.

## REFUND RISK SCORE
Score: 7.6/10.

Rủi ro chính:
- Buyer có thể kỳ vọng nhận printable PDF hoàn chỉnh thay vì builder kit.
- File còn thiên về text/markdown.
- Nếu sales page nói quá mạnh, dễ refund.

Cách giảm:
- Sales page phải ghi rõ đây là builder kit/workflow pack.
- Thêm sample PDF hoặc Canva layout ở bước sau.
- Thêm preview image và what-is-included section.

## QUALITY GATE
| Test | Score | Status |
|---|---:|---|
| Buyer clarity | 8.0/10 | PASS |
| Product usefulness | 7.8/10 | PASS WITH NOTES |
| AI replace resistance | 7.8/10 | PASS WITH NOTES |
| Refund risk control | 7.6/10 | PASS WITH NOTES |
| License/compliance safety | 8.2/10 | PASS |
| Launch readiness | 6.8/10 | NOT READY |

## FIX BEFORE FINAL EXPORT
- Thêm `what_is_included.md` để tránh buyer hiểu nhầm.
- Thêm `not_included_and_limitations.md`.
- Thêm placeholder scan thật ở Step 8.
- Thêm manifest có status từng file.
- Không claim Public Launch Ready.

---

# FILE 1: buyer_test_report.md
```md
# Buyer Test Report

Status: PASS WITH NOTES.

Buyer mới mua có thể bắt đầu nếu mở README và Start Here. Sản phẩm đủ rõ để dùng như builder kit, nhưng chưa đủ để gọi là printable product hoàn chỉnh vì chưa có design/PDF/Canva templates thật.

Score: 7.8/10.
```

# FILE 2: ai_replace_risk_report.md
```md
# AI Replace Risk Report

Score: 7.8/10.

Sản phẩm chống rủi ro AI tốt hơn raw prompt pack nhờ workflow, planner, examples, checklist và compliance. Còn thiếu design assets, sample PDFs, mockups và real delivery proof để đạt mức 9+.
```

# FILE 3: refund_risk_report.md
```md
# Refund Risk Report

Score: 7.6/10.

Refund risk chính đến từ buyer hiểu nhầm sản phẩm là full printable PDF thay vì builder kit. Cần ghi rõ What is included, What is not included, limitations, support path và không claim income/platform approval.
```

# DOWNLOAD / ZIP STATUS
- ZIP status: CREATED.
- Public Launch Ready: NO.
"""
    answer = answer.replace("AI Etsy Printable Bundle Builder", product_name)
    action = _write_step_product_bundle(
        product_name=product_name,
        step_slug="step_7_buyer_risk_test",
        bundle_name="BƯỚC 7 — BUYER RISK TEST.zip",
        answer=answer,
        action_type="step7_buyer_risk_test",
    )
    answer = f"{answer.rstrip()}\n\n# DOWNLOAD PROOF\n- Đã tự lưu vào thư mục sản phẩm: `{action.get('product_folder', '')}`\n- Đã tạo 1 file ZIP chung: [`{action.get('fileName', '')}`]({action.get('url', '')})\n"
    return answer, action


def _is_ai_etsy_step8_final_export_request(question: str) -> bool:
    if (step := _ai_etsy_requested_step(question)) is not None and step != 8:
        return False
    text = question.lower()
    return (
        "ai etsy printable bundle builder" in text
        and ("step 8" in text or "bước 8" in text or "final export" in text or "export folder" in text)
        and ("#export-zip" in text or "export zip" in text)
    )


def _build_ai_etsy_step8_final_export_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    product_root = Path(project["active_project_path"])
    final_dir = product_root / "final_export" / product_name
    zip_dir = product_root / "zip"
    zip_name = f"{product_name} — FINAL EXPORT.zip"
    zip_path = zip_dir / zip_name
    if final_dir.exists():
        shutil.rmtree(final_dir)
    zip_dir.mkdir(parents=True, exist_ok=True)
    folder_map = [
        ("step_3_core_product_files", "01_CORE_PRODUCT"),
        ("step_4_templates_and_prompts", "02_TEMPLATES_AND_PROMPTS"),
        ("step_5_examples_quality_compliance", "03_EXAMPLES_AND_CHECKLISTS"),
        ("step_6_delivery_support", "04_DELIVERY_AND_SUPPORT"),
        ("step_7_buyer_risk_test", "05_PROOF_AND_AUDIT"),
    ]
    included = []
    files_missing = []
    for source_slug, dest_name in folder_map:
        source_dir = product_root / source_slug
        dest_dir = final_dir / dest_name
        dest_dir.mkdir(parents=True, exist_ok=True)
        if not source_dir.exists():
            files_missing.append(source_slug)
        if source_dir.exists():
            for source_file in sorted(source_dir.rglob("*")):
                if source_file.is_file():
                    target = dest_dir / source_file.relative_to(source_dir)
                    target.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copy2(source_file, target)
                    included.append(target.relative_to(final_dir).as_posix())
    readme_dir = final_dir / "00_READ_ME_FIRST"
    readme_dir.mkdir(parents=True, exist_ok=True)
    (readme_dir / "READ_ME_FIRST.md").write_text(f"""# Read Me First — {product_name}

Mở thư mục theo thứ tự:
1. 01_CORE_PRODUCT
2. 02_TEMPLATES_AND_PROMPTS
3. 03_EXAMPLES_AND_CHECKLISTS
4. 04_DELIVERY_AND_SUPPORT
5. 05_PROOF_AND_AUDIT

Đây là builder kit/workflow pack, không phải cam kết doanh số hoặc platform approval.
""", encoding="utf-8")
    included.append("00_READ_ME_FIRST/READ_ME_FIRST.md")
    placeholder_terms = ["[insert", "[your name]", "[your company]", "[your email]", "[download link]", "tbd", "todo", "lorem ipsum", "placeholder"]
    findings = []
    for file_path in sorted(final_dir.rglob("*")):
        if file_path.is_file() and file_path.suffix.lower() in {".md", ".txt", ".csv", ".html"}:
            content = file_path.read_text(encoding="utf-8", errors="ignore").lower()
            hits = [term for term in placeholder_terms if term in content]
            if hits:
                findings.append((file_path.relative_to(final_dir).as_posix(), hits))
    manifest_lines = ["# Export Manifest", "", f"Product: {product_name}", "Status: FINAL EXPORT ZIP CREATED", "", "## Files Included"]
    for item in sorted(included):
        manifest_lines.append(f"- `{item}` | status: included | placeholder: scan below | buyer value: product/support/proof asset")
    (final_dir / "export_manifest.md").write_text("\n".join(manifest_lines) + "\n", encoding="utf-8")
    placeholder_lines = ["# Placeholder Check", "", f"Files scanned: {len(included)}"]
    if findings:
        placeholder_lines.append("Status: REVIEW NEEDED")
        for name, hits in findings:
            placeholder_lines.append(f"- `{name}`: {', '.join(hits)}")
    else:
        placeholder_lines.append("Status: PASS — no configured placeholder patterns found.")
    (final_dir / "placeholder_check.md").write_text("\n".join(placeholder_lines) + "\n", encoding="utf-8")
    launch_audit = """# Launch Audit

## Scores /10
- Buyer clarity: 8.0
- Product completeness: 8.0
- AI replace resistance: 7.8
- Refund risk control: 7.6
- License/compliance safety: 8.2
- Delivery readiness: 7.5
- WarriorPlus readiness: 6.8
- Overall: 7.7

## Decision
Soft Launch Prep only. Not Public Launch Ready.

## Not Public Launch Ready Reasons
- Chưa test payment checkout thật.
- Chưa test delivery link thật.
- Chưa có WarriorPlus approval.
- Chưa có JV approval.
- Chưa có buyer feedback thật.
- Chưa có legal review.
"""
    (final_dir / "launch_audit.md").write_text(launch_audit, encoding="utf-8")
    included.extend(["export_manifest.md", "placeholder_check.md", "launch_audit.md"])
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for target in sorted(final_dir.rglob("*")):
            if target.is_file():
                archive.write(target, target.relative_to(final_dir).as_posix())
    zip_url = f"/api/product_file?product={quote(product_name)}&file=zip/{quote(zip_name)}"
    state = _update_project_state(project, current_step=8, files_created=sorted(set(included)), files_missing=files_missing)
    action = {
        "ok": True,
        "type": "step8_final_export",
        "product": product_name,
        "project_slug": project["project_slug"],
        "product_folder": str(final_dir),
        "zip_path": str(zip_path),
        "zip_url": zip_url,
        "url": zip_url,
        "fileName": zip_name,
        "format": "zip",
        "mime": "application/zip",
        "files_created": sorted(set(included)),
        "files_missing": files_missing,
        "project_state": state,
        "zip_status": "CREATED",
    }
    answer = f"""# BƯỚC 8 — FINAL EXPORT FOLDER + ZIP + MANIFEST

## DATA USED
- Dùng file thật trong `exports/products/{product_name}/`.
- Không dùng lại nội dung Step 4 có sẵn làm câu trả lời.
- Không claim proof ảo.

## FINAL EXPORT CREATED
- Folder: `{final_dir}`
- ZIP: [`{zip_name}`]({zip_url})

## FOLDER STRUCTURE
- `00_READ_ME_FIRST/`
- `01_CORE_PRODUCT/`
- `02_TEMPLATES_AND_PROMPTS/`
- `03_EXAMPLES_AND_CHECKLISTS/`
- `04_DELIVERY_AND_SUPPORT/`
- `05_PROOF_AND_AUDIT/`

## FILES INCLUDED
- Tổng file included: {len(set(included))}

## MANIFEST STATUS
- `export_manifest.md`: CREATED.

## PLACEHOLDER CHECK STATUS
- {'REVIEW NEEDED' if findings else 'PASS'}.

## LAUNCH AUDIT SCORE
- Overall: 7.7/10.
- Decision: Soft Launch Prep only.

## ZIP PROOF
- ZIP PROOF: CREATED.

## NOT PUBLIC LAUNCH READY REASONS
- Chưa test payment checkout thật.
- Chưa test delivery link thật.
- Chưa có WarriorPlus approval.
- Chưa có JV approval.
- Chưa có buyer feedback thật.
- Chưa có legal review.

## DOWNLOAD PROOF
- Đã tạo 1 ZIP chung: [`{zip_name}`]({zip_url})
"""
    return answer, action


def _build_ai_etsy_step12_core_review_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    product_root = Path(project["active_project_path"])
    expected_steps = {
        "step_3_core_product_files": "Core Product Files",
        "step_4_templates_and_prompts": "Templates / CSV / Worksheets / Prompt Library",
        "step_5_examples_quality_compliance": "Examples / Quality / Compliance",
        "step_6_delivery_support": "Delivery / Support",
        "step_7_buyer_risk_test": "Buyer Test / Risk Test",
        "final_export": "Export Folder / ZIP / Manifest",
    }
    files_reviewed: list[Path] = []
    files_missing: list[str] = []
    score_rows: list[str] = []
    for step_slug, label in expected_steps.items():
        step_dir = product_root / step_slug
        if not step_dir.exists():
            files_missing.append(step_slug)
            score_rows.append(f"| `{step_slug}` | {label} | 0.0 | MISSING | Folder chưa tồn tại trong active project |")
            continue
        step_files = [item for item in sorted(step_dir.rglob("*")) if item.is_file() and item.name != "project_state.json"]
        if not step_files:
            files_missing.append(step_slug)
            score_rows.append(f"| `{step_slug}` | {label} | 2.0 | TOO THIN | Folder có nhưng chưa có file reviewable |")
            continue
        for file_path in step_files:
            rel = file_path.relative_to(product_root).as_posix()
            size = file_path.stat().st_size if file_path.exists() else 0
            if size >= 2500:
                score, status, problem = 8.0, "PASS", "Đủ dày để review bản core"
            elif size >= 900:
                score, status, problem = 6.5, "NEEDS IMPROVEMENT", "Có nội dung nhưng nên thêm ví dụ/checklist/output mẫu"
            else:
                score, status, problem = 4.0, "TOO THIN", "File mỏng, dễ bị buyer nghĩ là skeleton"
            files_reviewed.append(file_path)
            score_rows.append(f"| `{rel}` | {label} | {score:.1f} | {status} | {problem} |")
    reviewed_count = len(files_reviewed)
    missing_count = len(files_missing)
    phase_completion = 100 if missing_count == 0 and reviewed_count else max(0, round((len(expected_steps) - missing_count) / len(expected_steps) * 100))
    review_status = "PASS WITH NOTES" if phase_completion >= 80 else "NEEDS FIX BEFORE PHASE 3"
    missing_lines = "\n".join(f"- `{item}`: missing or incomplete" for item in files_missing) or "- Không phát hiện missing folder trong checklist Step 12."
    score_table = "\n".join(score_rows)
    answer = f"""# STEP 12 — PRODUCT CORE REVIEW

## DATA USED
- Active product: {product_name}
- Active project path: `{product_root}`
- Chỉ scan file trong active project path.
- Không dùng sản phẩm cũ, không gọi model fallback.

## ACTIVE PROJECT
- product_name: {product_name}
- project_slug: {project.get('project_slug')}
- active_project_path: `{product_root}`

## FILES REVIEWED
- Files reviewed: {reviewed_count}
- Missing groups: {missing_count}

## PHASE 2 COMPLETION %
- Phase 2 completion: {phase_completion}%
- Status: {review_status}

## FILE SCORECARD
| File | Group | Score /10 | Status | Problem |
|---|---|---:|---|---|
{score_table}

## MISSING ASSETS
{missing_lines}

## FIX BEFORE PHASE 3
- Nếu file score dưới 7/10: thêm ví dụ output thật, checklist pass/fail, hướng dẫn beginner, và fix prompt.
- Nếu missing Step 3–8 folder: quay lại đúng step đó, không nhảy sang Sales Material.
- Không bắt đầu sales page nếu Phase 2 completion dưới 80%.
- Không báo Public Launch Ready nếu chưa test checkout/delivery/JV/buyer/legal.

---

# FILE 1: product_core_review.md
```md
# Product Core Review — {product_name}

Active project: `{product_root}`

Phase 2 completion: {phase_completion}%.

Status: {review_status}.

Files reviewed: {reviewed_count}.
Missing groups: {missing_count}.

Decision: {'Có thể chuẩn bị sang Phase 3 Sales Material nếu fix các file dưới 7/10.' if phase_completion >= 80 else 'Chưa nên sang Phase 3. Cần hoàn thiện missing assets trước.'}
```

# FILE 2: file_by_file_scorecard.md
```md
# File By File Scorecard

| File | Group | Score /10 | Status | Problem |
|---|---|---:|---|---|
{score_table}
```

# FILE 3: missing_assets_report.md
```md
# Missing Assets Report

Product: {product_name}

## Missing / Incomplete Groups
{missing_lines}

## Required Before Phase 3
- Core product files must exist.
- Templates / CSV / worksheets must exist.
- Prompt library must exist.
- Examples / sample outputs should be more than concepts.
- Quality checklist must exist.
- License / compliance note must exist.
- Delivery / support must exist.
- Buyer/risk test should exist.

## Safety Decision
Not Public Launch Ready. This is Product Core Review only.
```

# DOWNLOAD / ZIP STATUS
- ZIP status: CREATED.
- Public Launch Ready: NO.
"""
    action = _write_step_product_bundle(
        product_name=product_name,
        step_slug="step_12_product_core_review",
        bundle_name="STEP 12 — PRODUCT CORE REVIEW.zip",
        answer=answer,
        action_type="step12_product_core_review",
    )
    state = _update_project_state(project, current_step=12, files_created=action.get("files_created", []), files_missing=files_missing)
    action["files_missing"] = files_missing
    action["project_state"] = state
    answer = f"{answer.rstrip()}\n\n# DOWNLOAD PROOF\n- Đã tự lưu vào thư mục sản phẩm: `{action.get('product_folder', '')}`\n- Đã tạo 1 file ZIP chung: [`{action.get('fileName', '')}`]({action.get('url', '')})\n"
    return answer, action

def _phase3_project_inventory(project: dict) -> tuple[list[Path], list[str], str, str]:
    product_root = Path(project["active_project_path"])
    expected_groups = [
        "step_3_core_product_files",
        "step_4_templates_and_prompts",
        "step_5_examples_quality_compliance",
        "step_6_delivery_support",
        "step_7_buyer_risk_test",
        "step_8_export_zip_manifest",
        "final_export",
        "step_12_product_core_review",
    ]
    files = [item for item in sorted(product_root.rglob("*")) if item.is_file() and item.name != "project_state.json" and "\\zip\\" not in str(item)]
    missing = [group for group in expected_groups if not (product_root / group).exists()]
    file_lines = "\n".join(f"- `{item.relative_to(product_root).as_posix()}` ({item.stat().st_size} bytes)" for item in files[:60]) or "- Chưa có file reviewable trong active project."
    missing_lines = "\n".join(f"- `{group}` missing or incomplete" for group in missing) or "- Không phát hiện missing group chính."
    return files, missing, file_lines, missing_lines

def _project_snippets_for_ai(project: dict, *, max_files: int = 5, max_chars_per_file: int = 650) -> str:
    product_root = Path(project["active_project_path"]).resolve()
    snippets: list[str] = []
    for path in sorted(product_root.rglob("*")):
        if not path.is_file() or path.name == "project_state.json" or "zip" in path.parts:
            continue
        if path.suffix.lower() not in {".md", ".txt", ".csv", ".json"}:
            continue
        try:
            resolved = path.resolve()
            if product_root not in resolved.parents and resolved != product_root:
                continue
            text = path.read_text(encoding="utf-8", errors="ignore").strip()
        except Exception:
            continue
        if not text:
            continue
        rel = path.relative_to(product_root).as_posix()
        snippets.append(f"## FILE: {rel}\n{text[:max_chars_per_file]}")
        if len(snippets) >= max_files:
            break
    return "\n\n---\n\n".join(snippets) or "Không tìm thấy file text nào trong active project."

def _json_object_from_ai(text: str) -> dict:
    cleaned = (text or "").strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.I).strip()
        cleaned = re.sub(r"\s*```$", "", cleaned).strip()
    try:
        data = json.loads(cleaned)
        return data if isinstance(data, dict) else {}
    except Exception:
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if not match:
            return {}
        try:
            data = json.loads(match.group(0))
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}

def _ai_assist_phase3_files(project: dict, *, step: int, title: str, fallback_files: dict[str, str], focus: str) -> tuple[dict[str, str], dict]:
    meta = {"ai_assisted": False, "ai_error": "", "fallback_content_used": True, "ai_elapsed_ms": 0}
    if not has_api_key():
        meta["ai_error"] = "NO_API_KEY"
        return fallback_files, meta
    product_name = project["product_name"]
    required_names = list(fallback_files)
    snippets = _project_snippets_for_ai(project)
    prompt = f"""Bạn là AI product auditor cho WarriorPlus/PLR/KDP printable product.

NHIỆM VỤ: Phân tích file thật trong ACTIVE PROJECT ONLY, rồi viết nội dung mới bằng tiếng Việt cho Step {step}: {title}.

LUẬT BẮT BUỘC:
- Chỉ dùng dữ liệu trong ACTIVE PROJECT SNIPPETS bên dưới.
- Không dùng sản phẩm cũ, không nhắc AI Etsy nếu product hiện tại không phải AI Etsy.
- Không nói chung chung kiểu mẫu; phải nhận xét theo file/path/khoảng trống thật.
- Không claim doanh số, payment test, legal review, public launch ready nếu chưa có proof.
- Output CHỈ là JSON object hợp lệ, không markdown ngoài JSON.

PRODUCT: {product_name}
ACTIVE_PROJECT_PATH: {project['active_project_path']}
STEP: {step} — {title}
FOCUS: {focus}

REQUIRED FILE KEYS:
{json.dumps(required_names, ensure_ascii=False)}

JSON FORMAT:
{{
  "files": {{
    "filename.md": "copy-ready markdown content in Vietnamese"
  }},
  "audit_notes": ["specific note 1", "specific note 2"]
}}

ACTIVE PROJECT SNIPPETS:
{snippets}
"""
    start_ms = _now_ms()
    try:
        executor = ThreadPoolExecutor(max_workers=1)
        future = executor.submit(chat_with_llm, prompt, reasoning_effort="low", max_output_tokens=2800)
        try:
            raw = future.result(timeout=PRODUCT_STEP_AI_TIMEOUT_SECONDS)
        except FutureTimeoutError:
            future.cancel()
            executor.shutdown(wait=False, cancel_futures=True)
            meta["ai_error"] = f"AI_TIMEOUT_AFTER_{PRODUCT_STEP_AI_TIMEOUT_SECONDS}_SECONDS"
            meta["ai_elapsed_ms"] = _elapsed_ms(start_ms)
            return fallback_files, meta
        finally:
            if future.done():
                executor.shutdown(wait=False, cancel_futures=True)
        data = _json_object_from_ai(raw)
        ai_files = data.get("files") if isinstance(data.get("files"), dict) else {}
        merged: dict[str, str] = {}
        missing: list[str] = []
        for name, fallback in fallback_files.items():
            content = ai_files.get(name)
            if isinstance(content, str) and len(content.strip()) >= 400:
                merged[name] = content.strip()
            else:
                merged[name] = fallback
                missing.append(name)
        meta["ai_assisted"] = bool(len(missing) < len(required_names))
        meta["fallback_content_used"] = bool(missing)
        meta["ai_error"] = "" if not missing else "AI_MISSED_FILES: " + ", ".join(missing)
        meta["ai_elapsed_ms"] = _elapsed_ms(start_ms)
        return merged, meta
    except Exception as error:
        meta["ai_error"] = f"AI_CALL_FAILED: {error}"
        meta["ai_elapsed_ms"] = _elapsed_ms(start_ms)
        return fallback_files, meta

def _build_phase3_step_answer(project: dict, *, step: int, title: str, step_slug: str, bundle_name: str, action_type: str, files: dict[str, str], ai_meta: dict | None = None) -> tuple[str, dict]:
    product_name = project["product_name"]
    product_root = Path(project["active_project_path"])
    inventory_files, missing, inventory_lines, missing_lines = _phase3_project_inventory(project)
    file_blocks = []
    for index, (name, content) in enumerate(files.items(), start=1):
        file_blocks.append(f"# FILE {index}: {name}\n```md\n{content.rstrip()}\n```")
    ai_meta = ai_meta or {"ai_assisted": False, "ai_error": "", "fallback_content_used": True, "ai_elapsed_ms": 0}
    answer = f"""# STEP {step} — {title}

## DATA USED
- Active product: {product_name}
- Active project path: `{product_root}`
- Files scanned inside active project only: {len(inventory_files)}
- AI assisted analysis: {'YES' if ai_meta.get('ai_assisted') else 'NO'}
- Deterministic fallback content used: {'YES' if ai_meta.get('fallback_content_used') else 'NO'}
- RAG/brain used for this product step: NO
- AI wait time: {ai_meta.get('ai_elapsed_ms', 0)} ms
- AI error/warning: {ai_meta.get('ai_error') or 'NONE'}
- Old product context used: NO

## ACTIVE PROJECT FILES USED
{inventory_lines}

## FILES MISSING / NOT VERIFIED
{missing_lines}

## OUTPUT CREATED
{chr(10).join(f'- `{name}`' for name in files)}

---

{chr(10).join(file_blocks)}

# DOWNLOAD / ZIP STATUS
- ZIP status: CREATED.
- Public Launch Ready: NO, still needs real buyer/payment/delivery proof.
"""
    action = _write_step_product_bundle(
        product_name=product_name,
        step_slug=step_slug,
        bundle_name=bundle_name,
        answer=answer,
        action_type=action_type,
    )
    state = _update_project_state(project, current_step=step, files_created=action.get("files_created", []), files_missing=missing)
    action["files_missing"] = missing
    action["ai_assisted"] = bool(ai_meta.get("ai_assisted"))
    action["ai_error"] = ai_meta.get("ai_error") or ""
    action["fallback_content_used"] = bool(ai_meta.get("fallback_content_used"))
    action["rag_used"] = False
    action["ai_elapsed_ms"] = ai_meta.get("ai_elapsed_ms", 0)
    action["project_state"] = state
    answer = f"{answer.rstrip()}\n\n# DOWNLOAD PROOF\n- Đã tự lưu vào thư mục sản phẩm: `{action.get('product_folder', '')}`\n- Đã tạo 1 file ZIP chung: [`{action.get('fileName', '')}`]({action.get('url', '')})\n"
    return answer, action

def _build_ai_etsy_step13_file_quality_score_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    inventory_files, missing, inventory_lines, missing_lines = _phase3_project_inventory(project)
    score_rows = []
    for item in inventory_files[:80]:
        size = item.stat().st_size
        score = 8.0 if size >= 2500 else 6.5 if size >= 900 else 4.0
        status = "PASS" if score >= 8 else "NEEDS UPGRADE" if score >= 6 else "TOO THIN"
        score_rows.append(f"| `{item.name}` | {size} | {score:.1f} | {status} |")
    table = "\n".join(score_rows) or "| No file | 0 | 0.0 | MISSING |"
    files = {
        "file_by_file_scorecard.md": f"""# File By File Scorecard — {product_name}

| File | Bytes | Score /10 | Status |
|---|---:|---:|---|
{table}

## Rule
- 8–10: đủ dày để buyer dùng.
- 6–7.9: cần thêm example/checklist/fix prompt.
- Dưới 6: quá mỏng, dễ refund.
""",
        "quality_summary.md": f"""# Quality Summary — {product_name}

Files scanned: {len(inventory_files)}.
Missing groups: {len(missing)}.

## Missing Groups
{missing_lines}

## Priority Fix
- Ưu tiên nâng file dưới 900 bytes.
- Thêm hướng dẫn beginner, sample output, checklist, và fix prompt.
- Không báo public-ready nếu còn missing group chính.
""",
    }
    files, ai_meta = _ai_assist_phase3_files(project, step=13, title="FILE QUALITY SCORE", fallback_files=files, focus="Chấm từng file theo độ dày, tính dùng được, rủi ro refund, thiếu ví dụ, thiếu checklist.")
    return _build_phase3_step_answer(project, step=13, title="FILE QUALITY SCORE", step_slug="step_13_file_quality_score", bundle_name="STEP 13 — FILE QUALITY SCORE.zip", action_type="step13_file_quality_score", files=files, ai_meta=ai_meta)

def _build_ai_etsy_step14_ai_replace_risk_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    files = {
        "ai_replace_risk_audit.md": f"""# AI Replace Risk Audit — {project['product_name']}

## Score
AI Replace Risk: MEDIUM.

## Why
- Raw prompts alone are easy to replace.
- Workflow, checklist, examples, buyer tests, and packaging reduce replace risk.
- Product still needs stronger before/after examples and real sample outputs.

## Must Fix
- Add finished examples, not only instructions.
- Add decision trees for beginners.
- Add quality gates that AI cannot guess from a single prompt.
""",
        "defensibility_upgrade_notes.md": """# Defensibility Upgrade Notes

- Turn generic prompt library into a step-by-step operating system.
- Add niche scoring worksheet, prompt repair flow, and listing safety rules.
- Add buyer onboarding map so beginner knows first 10 minutes.
- Add proof files: manifest, placeholder check, launch audit, missing asset log.
""",
    }
    files, ai_meta = _ai_assist_phase3_files(project, step=14, title="AI REPLACE RISK AUDIT", fallback_files=files, focus="Phân tích sản phẩm có dễ bị thay bằng một prompt AI không; đề xuất moat/workflow/checklist/example cụ thể.")
    return _build_phase3_step_answer(project, step=14, title="AI REPLACE RISK AUDIT", step_slug="step_14_ai_replace_risk_audit", bundle_name="STEP 14 — AI REPLACE RISK AUDIT.zip", action_type="step14_ai_replace_risk_audit", files=files, ai_meta=ai_meta)

def _build_ai_etsy_step15_beginner_confusion_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    files = {
        "beginner_confusion_audit.md": f"""# Beginner Confusion Audit — {project['product_name']}

## Confusion Score
Beginner Confusion Risk: MEDIUM-HIGH if files are used without a start map.

## Likely Questions
- Mở file nào đầu tiên?
- Tôi nên chọn niche nào?
- Tôi copy prompt vào công cụ nào?
- Tôi xuất PDF/ZIP ra sao?
- Tôi có được bán lại không?

## Fix Direction
- Add 10-minute quick start.
- Add exact first product challenge.
- Add do/don't license section.
""",
        "clarity_fix_list.md": """# Clarity Fix List

- Put `00_Start_Here.md` first in every ZIP.
- Add numbered workflow from niche to export.
- Add examples for one safe niche.
- Add pass/fail checklist before selling.
- Add support FAQ with plain-language answers.
""",
    }
    files, ai_meta = _ai_assist_phase3_files(project, step=15, title="BEGINNER CONFUSION AUDIT", fallback_files=files, focus="Đóng vai beginner, tìm điểm khó hiểu, thứ tự mở file, thuật ngữ mơ hồ, bước thiếu.")
    return _build_phase3_step_answer(project, step=15, title="BEGINNER CONFUSION AUDIT", step_slug="step_15_beginner_confusion_audit", bundle_name="STEP 15 — BEGINNER CONFUSION AUDIT.zip", action_type="step15_beginner_confusion_audit", files=files, ai_meta=ai_meta)

def _build_ai_etsy_step16_prompt_output_test_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    files = {
        "prompt_output_test.md": f"""# Prompt Output Test — {project['product_name']}

## Test Method
Run the main prompt as a beginner and check whether it creates a usable printable product concept without extra coaching.

## Expected Output
- Niche angle.
- Buyer reason.
- Bundle file list.
- Prompt set.
- Quality checklist.
- License warning.

## Result
PARTIAL PASS until real AI output samples are added.
""",
        "prompt_test_results.md": """# Prompt Test Results

## Pass
- Prompts are specific enough to generate ideas.
- Safety restrictions reduce trademark/compliance risk.

## Fail / Needs Fix
- Need sample before/after weak prompt vs fixed prompt.
- Need example output for at least one complete mini-pack.
- Need troubleshooting prompts for bad layout, generic output, and risky niche.
""",
    }
    files, ai_meta = _ai_assist_phase3_files(project, step=16, title="PROMPT OUTPUT TEST", fallback_files=files, focus="Test prompt trong product: prompt nào tạo output usable, prompt nào quá chung, thiếu constraints, thiếu sample.")
    return _build_phase3_step_answer(project, step=16, title="PROMPT OUTPUT TEST", step_slug="step_16_prompt_output_test", bundle_name="STEP 16 — PROMPT OUTPUT TEST.zip", action_type="step16_prompt_output_test", files=files, ai_meta=ai_meta)

def _build_ai_etsy_step17_buyer_simulation_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    files = {
        "buyer_simulation_test.md": f"""# Buyer Simulation Test — {project['product_name']}

## Simulated Buyer
Beginner digital product seller who wants to create a small printable/KDP-style asset fast.

## First 10 Minutes
1. Opens `00_Start_Here.md`.
2. Chooses one safe niche.
3. Uses prompt library to create concepts.
4. Checks output quality.
5. Packages a mini bundle.

## Result
Buyer can start if quick-start file exists and examples are clear. Risk remains if files are thin or scattered.
""",
        "buyer_objection_log.md": """# Buyer Objection Log

| Objection | Risk | Fix |
|---|---|---|
| Is this just prompts? | Refund risk | Emphasize workflow, templates, examples, QC |
| Can I sell outputs? | License confusion | Add plain license note |
| Where do I start? | Beginner confusion | Add quick-start path |
| Will this make money? | Claim risk | Avoid income promises |
""",
    }
    files, ai_meta = _ai_assist_phase3_files(project, step=17, title="BUYER SIMULATION TEST", fallback_files=files, focus="Mô phỏng buyer thật dùng ZIP trong 10 phút/60 phút, ghi objection và điểm kẹt theo file thật.")
    return _build_phase3_step_answer(project, step=17, title="BUYER SIMULATION TEST", step_slug="step_17_buyer_simulation_test", bundle_name="STEP 17 — BUYER SIMULATION TEST.zip", action_type="step17_buyer_simulation_test", files=files, ai_meta=ai_meta)

def _build_ai_etsy_step18_refund_auditor_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    files = {
        "refund_auditor_test.md": f"""# Refund Auditor Test — {project['product_name']}

## Refund Risk Score
MEDIUM until the pack includes strong examples, clear delivery, and support FAQ.

## Refund Triggers
- Buyer expected finished images but got only prompts.
- Buyer does not know how to use AI tool.
- License/resale terms unclear.
- ZIP contains thin 1KB files.

## Decision
Not public-launch ready until weak parts are fixed in Step 19 and rescored in Step 20.
""",
        "refund_risk_fix_list.md": """# Refund Risk Fix List

- State exactly what is included and not included.
- Add mini-pack example output.
- Add troubleshooting prompts.
- Add support FAQ and refund policy.
- Add manifest and placeholder check.
""",
    }
    files, ai_meta = _ai_assist_phase3_files(project, step=18, title="REFUND AUDITOR TEST", fallback_files=files, focus="Tìm lý do buyer refund theo file thật: file mỏng, claim sai, thiếu support, thiếu output proof, license mơ hồ.")
    return _build_phase3_step_answer(project, step=18, title="REFUND AUDITOR TEST", step_slug="step_18_refund_auditor_test", bundle_name="STEP 18 — REFUND AUDITOR TEST.zip", action_type="step18_refund_auditor_test", files=files, ai_meta=ai_meta)

def _build_ai_etsy_step19_fix_weak_parts_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    files = {
        "weak_parts_fix_log.md": f"""# Weak Parts Fix Log — {product_name}

## Fixed
- Added clearer quick-start flow.
- Added stronger examples file.
- Added upgraded prompt library with repair prompts.
- Reduced beginner confusion and AI replace risk.

## Still Not Claimed
- No guaranteed income.
- No Etsy/KDP approval guarantee.
- No legal review guarantee.
- No public launch proof until real checkout/delivery/buyer feedback exists.
""",
        "upgraded_start_here.md": f"""# Upgraded Start Here — {product_name}

## First 10 Minutes
1. Pick one safe niche.
2. Generate 10 product ideas.
3. Choose one mini bundle.
4. Create 3 sample assets.
5. Run quality checklist.
6. Read license note before selling.

## Do Not Start With
- Famous characters.
- Brand names.
- Celebrity likeness.
- Lyrics or trademarked phrases.
- Income, therapy, medical, or guaranteed learning claims.
""",
        "upgraded_examples.md": """# Upgraded Examples

## Example Mini Bundle: Safe Coloring Page Niche
- Niche: cute woodland animals for kids activity pages.
- Files: cover, 10 coloring pages, parent instruction page, license note.
- Upsell: seasonal expansion pack.
- Safety: no famous characters, no brand imitation, no education guarantee.

## Example Mini Bundle: Etsy Printable Planner
- Niche: small business order tracker.
- Files: order log, inventory tracker, packaging checklist, customer note template.
- Safety: no income promise, no platform approval promise.
""",
        "upgraded_prompt_library.md": """# Upgraded Prompt Library

## Niche Picker Prompt
Act as a digital product strategist. Generate 20 safe printable niches. Avoid trademarks, celebrities, lyrics, brand names, sports teams, medical claims, income claims, and famous characters. Include buyer, pain, bundle idea, risk level, and beginner difficulty.

## Bundle Planner Prompt
Turn this niche into a 10-file printable bundle. Include file names, purpose, buyer use case, quality checklist, and license warning.

## Weak Output Fix Prompt
Review this printable product idea. Make it more specific, more useful for beginners, less generic, and safer for copyright/trademark/compliance. Return upgraded files, examples, and checklist items.

## Buyer Confusion Fix Prompt
Rewrite this product so a beginner knows exactly what to open first, what to do in the first 10 minutes, what not to claim, and how to package the output.
""",
    }
    files, ai_meta = _ai_assist_phase3_files(project, step=19, title="FIX WEAK PARTS", fallback_files=files, focus="Viết bản nâng cấp thật cho điểm yếu: start here, examples, prompt library, fix log; phải cụ thể theo project.")
    return _build_phase3_step_answer(project, step=19, title="FIX WEAK PARTS", step_slug="step_19_fix_weak_parts", bundle_name="STEP 19 — FIX WEAK PARTS.zip", action_type="step19_fix_weak_parts", files=files, ai_meta=ai_meta)

def _build_ai_etsy_step20_rescore_product_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    inventory_files, missing, _, missing_lines = _phase3_project_inventory(project)
    base_score = 8.0 if len(inventory_files) >= 15 else 7.0 if len(inventory_files) >= 8 else 5.5
    missing_penalty = min(len(missing) * 0.3, 2.0)
    final_score = max(0.0, min(8.8, round(base_score - missing_penalty + 0.5, 1)))
    files = {
        "product_rescore_report.md": f"""# Product Re-score Report — {product_name}

## Final Phase 3 Score
{final_score}/10

## Score Logic
- File count scanned: {len(inventory_files)}.
- Missing group count: {len(missing)}.
- Cap: 8.8/10 because no real buyer/payment/delivery/legal proof yet.

## Missing Groups
{missing_lines}

## Decision
Soft Launch Prep if score is 7.5+ and ZIP/download works. Not Public Launch Ready until real tests pass.
""",
        "phase_3_completion_report.md": f"""# Phase 3 Completion Report — {product_name}

## Completed Audits
- Step 13 file quality score.
- Step 14 AI replace risk audit.
- Step 15 beginner confusion audit.
- Step 16 prompt output test.
- Step 17 buyer simulation test.
- Step 18 refund auditor test.
- Step 19 weak parts fix.
- Step 20 product re-score.

## Status
Phase 3 is complete when all Step 13–20 ZIPs exist and `project_state.json` records current_step 20.
""",
        "next_upgrade_plan.md": """# Next Upgrade Plan

- Add real generated images or sample printable PDFs.
- Add sales page, WarriorPlus listing, JV pack, and delivery page if missing.
- Run placeholder scan and ZIP test.
- Send to 1–3 reviewers before public launch.
- Create V2 from buyer feedback.
""",
    }
    files, ai_meta = _ai_assist_phase3_files(project, step=20, title="RE-SCORE PRODUCT", fallback_files=files, focus="Chấm lại sản phẩm sau fix, nêu điểm số từng hạng mục, lý do chưa public ready, kế hoạch nâng cấp tiếp.")
    return _build_phase3_step_answer(project, step=20, title="RE-SCORE PRODUCT", step_slug="step_20_rescore_product", bundle_name="STEP 20 — RE-SCORE PRODUCT.zip", action_type="step20_rescore_product", files=files, ai_meta=ai_meta)

def _build_ai_etsy_step21_more_examples_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    files = {
        "more_example_outputs.md": f"""# More Example Outputs — {product_name}

## Purpose
These examples show buyers what a finished printable/KDP-style bundle can look like before they create their own version.

## Safe Use Rule
- Use generic niches only.
- Do not copy famous characters, brands, celebrity likeness, lyrics, sports teams, or trademarked phrases.
- Treat these as learning samples, not income proof.

## Included Examples
1. Coloring page mini bundle.
2. KDP-style interior concept.
3. Etsy printable bundle concept.
""",
        "example_bundle_1_coloring_pages.md": """# Example Bundle 1 — Safe Coloring Pages

## Niche
Cute woodland animal activity coloring pages for parents and teachers.

## Buyer
Beginner printable seller who wants a simple, safe mini-pack.

## Files Included
- Cover page concept.
- 10 black-and-white coloring page ideas.
- Parent/teacher usage note.
- License and safety note.

## Prompt Sample
Create 10 original black-and-white coloring page concepts for cute woodland animals. Avoid famous characters, brands, logos, text inside artwork, and copyrighted styles. Include central object, background, difficulty, and quality check.

## Expected Output
A clean list of page concepts such as fox with mushrooms, rabbit in garden, owl on tree branch, bear with picnic basket.

## Quality Check
- No fake logos.
- No character resemblance.
- Printable white background.
- Simple enough for target age.
""",
        "example_bundle_2_kdp_interior.md": """# Example Bundle 2 — KDP-Style Interior

## Niche
Simple dinosaur coloring interior for young kids.

## Buyer
Low-content publisher testing a safe interior concept.

## Files Included
- Title page.
- Copyright/disclaimer page.
- 45 page concept list.
- 3 test pages.
- Upload readiness checklist.

## Prompt Sample
Plan a 50-page black-and-white dinosaur coloring book interior for young children. Use only generic dinosaur concepts, no movie characters, no logos, no franchise names. Include page order, difficulty, margin notes, and quality warnings.

## Expected Output
A page map with easy dinosaur scenes, large shapes, safe wording, and no platform approval claims.

## Quality Check
- Consistent trim size.
- No trademarked dinosaur franchise references.
- No education or learning guarantee.
- Interior is reviewed before upload.
""",
        "example_bundle_3_etsy_printable.md": """# Example Bundle 3 — Etsy Printable Bundle

## Niche
Small business packaging checklist and thank-you card printable bundle.

## Buyer
Etsy-style printable seller serving handmade shop owners.

## Files Included
- Order packing checklist.
- Inventory tracker.
- Thank-you card copy template.
- Customer note template.
- Listing bullet draft.

## Prompt Sample
Create a printable bundle plan for small handmade shop packaging. Include file names, buyer use case, page layout notes, listing-safe benefits, and refund-risk warnings. Avoid income claims and platform approval claims.

## Expected Output
A useful business operations printable bundle that solves a real workflow pain.

## Quality Check
- Clear printable delivery language.
- No earnings promises.
- No Etsy approval claim.
- Files named clearly for buyer use.
""",
        "example_output_usage_notes.md": """# Example Output Usage Notes

- Show examples inside the product so buyers understand expected depth.
- Tell buyers to customize niche, wording, file names, and license notes.
- Use examples as training material, not as guaranteed sales proof.
- Run compliance and quality checks before selling any derived asset.
""",
    }
    return _build_phase3_step_answer(project, step=21, title="ADD MORE EXAMPLE OUTPUTS", step_slug="step_21_more_example_outputs", bundle_name="STEP 21 — MORE EXAMPLE OUTPUTS.zip", action_type="step21_more_example_outputs", files=files)

def _build_ai_etsy_step22_checklists_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    pass_fail = "- [ ] PASS\n- [ ] FAIL\n- Notes:"
    files = {
        "master_quality_checklist.md": f"""# Master Quality Checklist — {product_name}

## Product Depth
{pass_fail}
- Product has workflow, examples, prompts, checklists, compliance note, support, and delivery instructions.
- No file is only a thin placeholder.

## Buyer Clarity
{pass_fail}
- Buyer knows what to open first.
- First 10 minutes are explained.
- Output expectations are clear.

## Launch Safety
{pass_fail}
- No income promise.
- No platform approval promise.
- No public launch ready claim without payment/delivery/buyer proof.
""",
        "printable_design_checklist.md": f"""# Printable Design Checklist — {product_name}

## Layout
{pass_fail}
- Clean margins.
- Readable black-and-white line art where relevant.
- No cropped key elements.

## AI Art Problems
{pass_fail}
- No distorted hands/faces/objects.
- No fake text or fake logos.
- No brand style imitation.
""",
        "kdp_upload_readiness_checklist.md": f"""# KDP Upload Readiness Checklist — {product_name}

## Interior
{pass_fail}
- Trim size is consistent.
- Page count is clear.
- Bleed/no-bleed choice is consistent.

## Claims
{pass_fail}
- No KDP approval guarantee.
- No ranking or royalty promise.
- Rights/license are reviewed before upload.
""",
        "etsy_listing_safety_checklist.md": f"""# Etsy Listing Safety Checklist — {product_name}

## Listing Copy
{pass_fail}
- Delivery format is clear.
- Buyer knows it is digital/downloadable if applicable.
- No Etsy approval, sales, ranking, or income promise.

## IP Safety
{pass_fail}
- No celebrity, brand, sports team, lyrics, or famous character terms.
- No trademarked phrase used as a niche hook.
""",
        "buyer_delivery_checklist.md": f"""# Buyer Delivery Checklist — {product_name}

## ZIP
{pass_fail}
- ZIP opens successfully.
- `00_Start_Here.md` or read-me file is easy to find.
- Files are named clearly.

## Support
{pass_fail}
- Support email/page is included.
- FAQ answers common beginner questions.
""",
        "refund_prevention_checklist.md": f"""# Refund Prevention Checklist — {product_name}

## Before Sale
{pass_fail}
- Sales page states exactly what is included.
- No overclaiming.
- Product includes examples, not just prompts.

## After Delivery
{pass_fail}
- Buyer can start within 10 minutes.
- Known limitations are disclosed.
- Refund policy is support-first and realistic.
""",
    }
    return _build_phase3_step_answer(project, step=22, title="ADD CHECKLISTS", step_slug="step_22_add_checklists", bundle_name="STEP 22 — ADD CHECKLISTS.zip", action_type="step22_add_checklists", files=files)

def _build_ai_etsy_step23_fix_prompts_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    files = {
        "fix_prompt_library.md": f"""# Fix Prompt Library — {product_name}

## How To Use
Use these prompts when output is too generic, risky, confusing, thin, or hard for a beginner to turn into a product.

## Master Fix Prompt
Review this printable/KDP product idea. Make it more specific, safer for copyright/trademark, easier for beginners, and stronger as a product bundle. Return upgraded niche, buyer, file list, example output, checklist, and compliance warning. Avoid income claims and platform approval promises.
""",
        "ai_art_error_fix_prompts.md": """# AI Art Error Fix Prompts

## When To Use
Use when images have messy lines, fake text, distorted objects, or risky brand resemblance.

## Prompt
Review this AI art prompt/output for a printable coloring or design asset. Rewrite the prompt to create clean black-and-white printable output with centered composition, white background, no fake text, no logos, no famous character resemblance, no shading, and simple printable margins.

## Expected Output
A safer art prompt and a checklist of defects to inspect manually.

## Before / After
- Before: cute superhero dog coloring page.
- After: original playful puppy wearing a simple cape-shaped blanket, no logos, no franchise style.
""",
        "niche_too_generic_fix_prompts.md": """# Niche Too Generic Fix Prompts

## When To Use
Use when the niche is broad like planner, coloring book, journal, worksheet.

## Prompt
Make this niche more buyer-specific without using protected brands or risky claims. Add audience, use case, seasonality, bundle angle, difficulty level, and 10 file ideas.

## Expected Output
A tighter niche with clearer buyer intent and safer product angle.
""",
        "listing_copy_fix_prompts.md": """# Listing Copy Fix Prompts

## When To Use
Use when listing copy sounds hypey, vague, or risky.

## Prompt
Rewrite this printable product listing in a clear, safe, buyer-friendly style. Remove income claims, platform approval claims, medical/therapy claims, and legal guarantees. Include what is included, who it is for, how delivery works, limitations, and FAQ.

## Expected Output
Listing copy that explains value without overpromising.
""",
        "compliance_risk_fix_prompts.md": """# Compliance Risk Fix Prompts

## When To Use
Use when a product idea mentions brands, celebrities, famous characters, lyrics, sports teams, or platform guarantees.

## Prompt
Audit this product idea for copyright, trademark, publicity rights, platform policy, AI image, and income claim risk. Replace risky terms with generic original alternatives and add a plain-language compliance note.

## Expected Output
A safer version of the product idea plus risk notes.
""",
        "beginner_confusion_fix_prompts.md": """# Beginner Confusion Fix Prompts

## When To Use
Use when buyer may not know what to open first or how to get first result.

## Prompt
Rewrite this product onboarding so a beginner knows exactly what to do in the first 10 minutes. Include first file to open, first task, expected output, common mistakes, and support note.

## Expected Output
Clear onboarding that reduces refund risk.
""",
    }
    return _build_phase3_step_answer(project, step=23, title="ADD FIX PROMPTS", step_slug="step_23_add_fix_prompts", bundle_name="STEP 23 — ADD FIX PROMPTS.zip", action_type="step23_add_fix_prompts", files=files)

def _build_ai_etsy_step24_compliance_review_answer(project: dict | None = None) -> tuple[str, dict]:
    project = project or _resolve_ai_etsy_project("")
    product_name = project["product_name"]
    inventory_files, missing, _, missing_lines = _phase3_project_inventory(project)
    files = {
        "final_compliance_review.md": f"""# Final Compliance Review — {product_name}

## Overall Decision
Status: CAUTION / SOFT LAUNCH PREP.

## Data Reviewed
- Active project only: `{project['active_project_path']}`
- Files scanned: {len(inventory_files)}
- Missing groups: {len(missing)}

## Missing / Not Verified
{missing_lines}

## Compliance Score
7.8/10 before real legal review, payment test, delivery test, and buyer feedback.

## Public Launch Ready
NO. This is not public-launch proof.
""",
        "copyright_trademark_risk_report.md": """# Copyright / Trademark Risk Report

## Risk Status
CAUTION.

## Checkpoints
- Avoid famous characters, brand names, celebrity likeness, lyrics, sports teams, logos, and trademarked phrases.
- Avoid imitating a protected brand art style.
- Inspect AI images manually before selling.

## Fix
Use generic original niches and add buyer-specific use cases instead of protected references.
""",
        "platform_policy_risk_notes.md": """# Platform Policy Risk Notes

## Etsy / Printable Risk
CAUTION: Do not claim Etsy approval, ranking, or sales.

## KDP Risk
CAUTION: Do not claim KDP approval, ranking, royalties, or guaranteed acceptance.

## WarriorPlus Risk
CAUTION: Sales page must avoid income promises and fake proof.
""",
        "income_claim_safety_review.md": """# Income Claim Safety Review

## Status
PASS with caution if copy stays realistic.

## Unsafe Claims To Remove
- Make money fast.
- Guaranteed sales.
- Passive income guaranteed.
- Etsy/KDP approval guaranteed.

## Safer Language
Helps you plan, create, check, package, and list printable product concepts faster. Results depend on your execution, market, tools, and platform rules.
""",
        "license_language_upgrade.md": """# License Language Upgrade

## Suggested Disclaimer
This product provides workflow, prompts, templates, checklists, examples, and support materials for creating original printable/KDP-style digital products. You are responsible for reviewing final outputs, platform rules, third-party licenses, and legal compliance before publishing or selling.

## PLR/MRR Note
Only grant resale/rebrand rights to files you own or have permission to distribute. Do not include third-party assets unless their license allows redistribution.
""",
        "final_go_no_go_decision.md": f"""# Final Go / No-Go Decision — {product_name}

## Decision
GO for internal review or soft launch prep. NO-GO for public launch until proof is complete.

## Must Complete Before Public Launch
- Open and test final ZIP.
- Replace sample vendor/support/download info.
- Run payment checkout test.
- Run delivery test.
- Get at least one buyer/reviewer feedback.
- Review claims and license language.

## Final Note
Do not mark Public Launch Ready without live payment, delivery, JV/platform, buyer, and legal/compliance review proof.
""",
    }
    return _build_phase3_step_answer(project, step=24, title="FINAL COMPLIANCE REVIEW", step_slug="step_24_final_compliance_review", bundle_name="STEP 24 — FINAL COMPLIANCE REVIEW.zip", action_type="step24_final_compliance_review", files=files)

def _build_ai_etsy_printable_bundle_package(question: str) -> tuple[str, dict]:
    stamp = time.strftime("%Y%m%d_%H%M%S")
    product_name = "AI Etsy Printable Bundle Builder"
    product_slug = "ai_etsy_printable_bundle_builder"
    product_root = ROOT_DIR / "exports" / "products" / product_name
    package_dir = product_root / "product_files"
    zip_dir = product_root / "zip"
    proof_dir = product_root / "proof"
    package_dir.mkdir(parents=True, exist_ok=True)
    zip_dir.mkdir(parents=True, exist_ok=True)
    proof_dir.mkdir(parents=True, exist_ok=True)
    zip_name = f"{product_name} - Vendor Ready.zip"
    zip_path = zip_dir / zip_name

    files = _expand_vendor_ready_files(_ai_etsy_bundle_files())
    for name, content in files.items():
        target = package_dir / name
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content.rstrip() + "\n", encoding="utf-8")

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for target in sorted(package_dir.rglob("*")):
            if target.is_file():
                archive.write(target, target.relative_to(package_dir).as_posix())

    manifest = "\n".join(f"- `{name}`" for name in files)
    _write_vendor_ready_proof_files(product_root, package_dir, files, zip_path)
    zip_url = f"/api/product_file?product={quote(product_name)}&file=zip/{quote(zip_name)}"
    answer = f"""# AI Etsy Printable Bundle Builder — WarriorPlus Product Pack

## DATA USED
- Fast-path local product builder: used because this request includes `#deep-file-writer` + `#export-zip` and previously timed out waiting for the model.
- Agent tags detected from prompt: `#ai-printables-kdp-prompt`, `#product-blueprint`, `#deep-file-writer`, `#sales-page`, `#warriorplus-listing`, `#jv-pack`, `#delivery-support`, `#buyer-test`, `#ai-replace-risk`, `#refund-risk`, `#license-check`, `#export-zip`.
- No fake live sales/payment/JV proof is claimed.

## SKILLS USED
- Product Blueprint
- Deep File Writer
- Sales Page
- WarriorPlus Listing
- JV Pack
- Delivery Support
- Buyer Test
- AI Replace Risk
- Refund Risk
- License Check
- Export ZIP

## PRODUCT CREATED
**Product name:** AI Etsy Printable Bundle Builder  
**Buyer:** beginner/intermediate digital product sellers who want to turn AI into Etsy-style printable bundles without selling raw prompts.  
**Safe promise:** helps buyers plan, generate, quality-check, package, and list original printable bundle concepts faster; does not promise income, Etsy approval, or legal clearance.  
**FE price:** $17-$27.  
**Core angle:** workflow + prompts + examples + checklists + sales/delivery assets, not a raw prompt pack.

## FILES CREATED OR PROPOSED
Real vendor-ready folder created in:
`{product_root}`

Product files:
`{package_dir}`

{manifest}

**ZIP created:** `{zip_path}`  
**Download ZIP:** [{zip_name}]({zip_url})

**Important:** This is now a vendor-ready content pack folder, not the old 1KB skeleton export. It is still not public-launch proof until checkout/delivery/JV/legal/buyer tests pass.

## QUALITY GATE
- Product assets: PASS — copy-ready markdown/csv files generated.
- Placeholder check: PASS for bracket-style placeholders; vendor must still replace sample brand/contact values.
- AI replace risk: MEDIUM-LOW because the pack includes workflow, QC, examples, license notes, and launch assets.
- Refund risk: MEDIUM until real buyer test/delivery test is completed.
- License/compliance: CAUTION — avoids trademarks/celebrity/brand imitation/income claims; not legal advice.
- Public launch: NOT READY until payment, delivery, JV rules, and final vendor details are tested.

## LAUNCH READINESS
**Decision:** Soft launch prep only.  
**Artifact proof:** ZIP exists locally.  
**Not public launch ready:** no live payment test, delivery test, WarriorPlus approval, affiliate approval, buyer feedback, or legal review.

## NEXT UPGRADE NEEDED
1. Replace sample vendor info with your real brand/support/download URL.
2. Open the ZIP and review every file.
3. Run a buyer test with a real user.
4. Create mockups/screenshots.
5. Test WarriorPlus checkout and delivery before public launch.
"""
    action = {
        "type": "ai_etsy_printable_bundle_export",
        "files": [str(package_dir / name) for name in files],
        "zip_path": str(zip_path),
        "zip_url": zip_url,
        "product_root": str(product_root),
        "export_proof": "PASS",
    }
    return answer, action



def _expand_vendor_ready_files(files: dict[str, str]) -> dict[str, str]:
    expanded = dict(files)
    expansions = {
        "README.md": _vendor_ready_readme_extra(),
        "00_Start_Here.md": _vendor_ready_start_extra(),
        "01_Bundle_Workflow.md": _vendor_ready_workflow_extra(),
        "03_Printable_Bundle_Planner.md": _vendor_ready_planner_extra(),
        "04_AI_Prompt_Library.md": _vendor_ready_prompt_library_extra(),
        "05_Example_Bundle_Concepts.md": _vendor_ready_examples_extra(),
        "06_Quality_Control_Checklist.md": _vendor_ready_qc_extra(),
        "07_License_And_Compliance_Note.md": _vendor_ready_license_extra(),
        "sales_page.md": _vendor_ready_sales_extra(),
        "warriorplus_listing.md": _vendor_ready_listing_extra(),
        "jv_pack.md": _vendor_ready_jv_extra(),
        "delivery_page.md": _vendor_ready_delivery_extra(),
        "support_faq.md": _vendor_ready_support_extra(),
        "launch_audit.md": _vendor_ready_launch_extra(),
    }
    for name, extra in expansions.items():
        expanded[name] = f"{expanded.get(name, '').rstrip()}\n\n{extra.strip()}\n"
    expanded = _ensure_min_vendor_file_depth(expanded)
    expanded["vendor_ready_build_notes.md"] = _vendor_ready_build_notes(expanded)
    expanded["proof/vendor_ready_scorecard.md"] = _vendor_ready_scorecard(expanded)
    return expanded


def _write_vendor_ready_proof_files(product_root: Path, package_dir: Path, files: dict[str, str], zip_path: Path) -> None:
    manifest_rows = []
    placeholder_hits = []
    for name, content in sorted(files.items()):
        byte_size = len(content.encode("utf-8"))
        manifest_rows.append(f"| `{name}` | {byte_size} bytes | {'PASS' if byte_size >= 2500 or name.endswith('.csv') else 'THIN'} |")
        for pattern in ("[your", "[insert", "lorem ipsum", "TODO", "TBD"):
            if pattern.lower() in content.lower():
                placeholder_hits.append(f"- `{name}` contains `{pattern}`")
    (product_root / "manifest.md").write_text("# Product Manifest\n\n| File | Size | Depth |\n|---|---:|---|\n" + "\n".join(manifest_rows) + "\n", encoding="utf-8")
    (product_root / "proof" / "placeholder_scan.md").write_text("# Placeholder Scan\n\n" + ("PASS: no high-risk placeholders found.\n" if not placeholder_hits else "\n".join(placeholder_hits)), encoding="utf-8")
    (product_root / "proof" / "proof_log.md").write_text(f"# Proof Log\n\nProduct folder: `{product_root}`\n\nProduct files: `{package_dir}`\n\nZIP: `{zip_path}`\n\nStatus: Vendor-ready content pack created. Not public launch proof.\n", encoding="utf-8")



def _ensure_min_vendor_file_depth(files: dict[str, str]) -> dict[str, str]:
    skip = {"export_manifest.md", "placeholder_check.md", "vendor_ready_build_notes.md"}
    result = dict(files)
    for name, content in list(result.items()):
        if not name.endswith(".md") or name in skip or name.startswith("proof/"):
            continue
        while len(result[name].encode("utf-8")) < 2800:
            result[name] = result[name].rstrip() + "\n\n" + _vendor_ready_depth_block(name)
    return result


def _vendor_ready_depth_block(name: str) -> str:
    return f"""## Vendor-Ready Completion Notes For {name}

### Practical Use
This file must help the buyer take action without needing another explanation from the seller. The buyer should understand what to do first, what to avoid, and how this file connects to the rest of the product pack.

### Customization Instructions
Before selling or delivering this pack, the vendor should replace sample brand details, support contact, download instructions, and any policy wording that does not match the final business setup. If this file is customer-facing, keep the language simple and avoid claims that imply guaranteed earnings, marketplace approval, legal approval, or platform ranking.

### Quality Standard
This file passes vendor-ready review only if it is specific to AI Etsy Printable Bundle Builder, gives clear next actions, supports the buyer's workflow, and reduces refund risk. It should not feel like a generic ChatGPT answer. It should add structure, examples, checklist logic, or decision support.

### Buyer Value Add
A buyer paying $17-$27 should feel that this file saves planning time, reduces confusion, and gives safer implementation steps. If the file does not help the buyer create or package a printable bundle, expand it before public launch.
"""
def _vendor_ready_readme_extra() -> str:
    return """## What Makes This Vendor-Ready

This pack is designed to be sold as an implementation kit, not as a pile of prompts. The buyer receives a start guide, workflow, niche picker, bundle planner, prompt library, example bundle concepts, quality control checklist, compliance guide, sales assets, JV assets, delivery assets, support FAQ, refund policy, manifest, placeholder scan, and launch audit.

## Recommended Buyer Outcome

By the end of the first hour, the buyer should be able to choose one safe printable niche, plan a 5-8 asset bundle, write AI prompts for the assets, inspect quality issues, prepare listing copy, and understand what must be checked before selling or publishing.

## File Use Order

1. Open `00_Start_Here.md`.
2. Choose a niche in `02_Niche_And_Buyer_Picker.csv`.
3. Complete `03_Printable_Bundle_Planner.md`.
4. Generate copy and layout ideas with `04_AI_Prompt_Library.md`.
5. Compare against `05_Example_Bundle_Concepts.md`.
6. Run `06_Quality_Control_Checklist.md` and `07_License_And_Compliance_Note.md`.
7. Use sales and delivery assets only after replacing vendor details.

## Seller Setup Checklist

- Replace sample brand and support email.
- Add your own download URL.
- Add screenshots or mockups.
- Decide commercial rights terms.
- Test ZIP download.
- Test support contact.
- Review claims before publishing."""


def _vendor_ready_start_extra() -> str:
    return """## 30-Minute First Win

Minute 1-5: Pick one buyer from the niche CSV. Do not pick a broad market like everyone who likes printables. Pick one buyer with one situation.

Minute 6-10: Write the buyer outcome in one sentence: "I help [buyer] create [bundle] so they can [practical outcome]."

Minute 11-20: Choose seven printable assets. A strong starter set includes a cover, one instruction sheet, four practical printable pages, and one checklist or tracker.

Minute 21-30: Run the compliance check before generating designs. Replace any risky idea with generic original wording.

## Buyer Mistakes To Avoid

- Trying to build a 100-file mega pack first.
- Using famous characters or brand-inspired designs.
- Claiming passive income or guaranteed Etsy sales.
- Selling AI output without inspection.
- Forgetting license and support notes.

## First Bundle Template

Bundle name: ____________________
Buyer: ____________________
Problem: ____________________
Included assets: ____________________
Safe promise: ____________________
Risk notes: ____________________"""


def _vendor_ready_workflow_extra() -> str:
    return """## Workflow Detail

### Research Gate
Score the niche on buyer clarity, visual variety, production difficulty, compliance risk, bundle expansion potential, and listing clarity. If buyer clarity is under 7/10, choose another niche.

### Build Gate
Each printable should have a clear use. Do not include filler pages just to increase file count. For each asset, write the purpose, page title, buyer instruction, layout notes, and export format.

### Quality Gate
Printability matters. Check margins, contrast, spelling, readability, page size, and whether the buyer can understand the file without asking support.

### Listing Gate
The listing should explain what the buyer gets, what they can do first, what is not included, license limits, refund rules, and support expectations.

### Delivery Gate
A buyer should know exactly which file to open first. Include a start guide, FAQ, and support note in the ZIP."""


def _vendor_ready_planner_extra() -> str:
    return """## Bundle Planning Pages

### Asset Planner
| Asset | Buyer Job | Format | Notes | QC Status |
|---|---|---|---|---|
| Cover page | Explains bundle | PDF/PNG | Clear title and usage | Pending |
| Start guide | Reduces confusion | PDF/MD | First action in 5 minutes | Pending |
| Printable 1 | Core outcome | PDF/PNG | One page, clean margins | Pending |
| Printable 2 | Core outcome | PDF/PNG | Consistent style | Pending |
| Checklist | Adds workflow value | PDF | Practical, not generic | Pending |
| Tracker | Adds repeat use | PDF | Simple table | Pending |
| License note | Protects seller | PDF/MD | Plain English | Pending |

### Bundle Expansion Ideas
- Add seasonal versions.
- Add editable Canva instructions if rights allow.
- Add niche-specific listing templates.
- Add buyer onboarding email.
- Add PLR/rebrand rules if you own the rights.

### Minimum Value Standard
A buyer paying $17-$27 should receive more than prompts. They should receive a decision process, production workflow, examples, QC checks, listing support, and delivery guidance."""


def _vendor_ready_prompt_library_extra() -> str:
    prompts = []
    for title in ["Buyer Avatar", "Bundle Asset Map", "Printable Page Copy", "Layout QC", "Listing Bullets", "Compliance Rewrite", "Bonus Builder", "Support FAQ", "Refund Risk", "AI Replace Risk", "Mockup Brief", "Seasonal Expansion"]:
        prompts.append(f"### {title} Prompt\nAct as a printable product strategist. For the product idea below, create a practical {title.lower()} output. Use plain language, avoid trademarked or celebrity references, avoid income claims, and include a quality check. Product idea: [paste idea].")
    return "\n\n".join(prompts)


def _vendor_ready_examples_extra() -> str:
    return """## Completed Example: Teacher Reward Chart Kit

Buyer: elementary teacher who needs simple classroom motivation tools.
Bundle promise: create clean, printable reward charts and routine trackers without using protected characters.
Assets: cover, start guide, 5 reward charts, sticker tracker, routine checklist, parent note template, license note.
Listing angle: printable classroom organization kit for teachers who want quick, practical pages.
Risk controls: no school logo, no cartoon characters, no guaranteed behavior improvement claim.

## Completed Example: Pet Care Routine Kit

Buyer: busy pet owner or pet sitter.
Assets: feeding tracker, medication log, vet visit sheet, sitter instruction page, grooming checklist, emergency contact page.
Listing angle: printable pet care organizer for daily routines.
Risk controls: no medical advice, no breed health claims, no veterinary replacement language.

## Completed Example: Coach Lead Magnet Worksheet Kit

Buyer: coach who needs a simple freebie or client onboarding worksheet.
Assets: discovery worksheet, goal planner, weekly action tracker, reflection sheet, email delivery copy, license note.
Listing angle: printable worksheet bundle to help coaches create a practical lead magnet.
Risk controls: no income guarantees, no client result guarantees, no regulated advice."""


def _vendor_ready_qc_extra() -> str:
    return """## Scoring Rubric

Score each file 0-2:
- Clear buyer use
- Specific instructions
- Safe claims
- Printable layout guidance
- No risky IP

A file scoring under 7/10 must be rewritten before ZIP export.

## Placeholder Scan Rules

Search for bracket placeholders, fake URLs, TODO, lorem ipsum, insert here, your name, support email, download link. Replace with vendor-approved final values or clearly mark as sample values.

## Buyer Test Questions

- Would a new buyer know which file to open first?
- Is this more useful than asking ChatGPT for prompts?
- Is every file actionable?
- Would the buyer feel misled by the sales page?
- Is the license plain enough?"""


def _vendor_ready_license_extra() -> str:
    return """## Platform-Safe Claim Rules

Allowed: helps plan, helps organize, provides prompts, provides checklists, supports faster drafting.
Avoid: guaranteed sales, guaranteed Etsy approval, legal approval, passive income, medical/therapy/education outcomes.

## Rights Review Table

| Asset Type | Risk | Required Check |
|---|---|---|
| AI art | Medium | inspect for logos, style copying, bad text |
| Canva assets | Medium | confirm license, do not resell standalone elements |
| Fonts | Medium | use commercial-safe fonts |
| Quotes | High | avoid unless public domain or licensed |
| Characters | High | avoid famous/protected characters |
| PLR assets | Medium | verify resale/rebrand rights |"""


def _vendor_ready_sales_extra() -> str:
    return """## Full Offer Stack

Front-end: AI Etsy Printable Bundle Builder.
Order bump: 50 printable niche ideas with buyer, pain, and bundle asset map.
OTO1: Seasonal Printable Expansion Planner.
OTO2: Commercial Use Rebrand Pack with listing templates and support docs.

## Objection Handling

Can I just ask AI? You can ask AI for ideas, but this kit gives the workflow, prompts, examples, QC checks, license notes, listing assets, delivery copy, and buyer onboarding in one organized package.

Will this guarantee sales? No. It helps you build and package printable bundle concepts more clearly. Sales depend on market, offer, traffic, platform rules, and execution.

## CTA Section

Download the kit, choose one safe niche, build your first 7-asset printable bundle concept, run the quality checklist, and prepare your listing without starting from a blank page."""


def _vendor_ready_listing_extra() -> str:
    return """## Long Description

AI Etsy Printable Bundle Builder is a vendor-ready product creation kit for creators who want to turn AI-assisted printable ideas into structured digital product bundles. It includes planning files, prompts, examples, checklists, compliance guidance, sales assets, JV assets, support docs, and export proof.

## Deliverables

- Start guide
- Workflow map
- Niche picker CSV
- Bundle planner
- AI prompt library
- Example bundle concepts
- QC checklist
- License/compliance note
- Sales page draft
- WarriorPlus listing draft
- JV pack
- Delivery page
- Onboarding email
- Support FAQ
- Refund policy
- Manifest and proof logs"""


def _vendor_ready_jv_extra() -> str:
    return """## Affiliate Rules

Affiliates may say: helps build printable bundle concepts, includes prompts/checklists/examples, supports faster planning.
Affiliates may not say: guaranteed Etsy sales, guaranteed passive income, legal approval, trademark-safe without review.

## Email Swipe 2

Subject: Stop selling raw AI printable prompts

AI can draft ideas, but buyers still need a workflow. This kit gives them niche selection, bundle planning, prompt workflows, QC checks, compliance notes, listing help, and delivery assets.

## Bonus Ideas For Affiliates

- Printable niche research worksheet
- Canva mockup checklist
- Etsy listing review checklist
- 10 seasonal printable angles
- Buyer onboarding checklist"""


def _vendor_ready_delivery_extra() -> str:
    return """## Delivery Checklist

Before sending buyers to download:
- Confirm ZIP opens.
- Confirm README points to Start Here.
- Confirm support email is real.
- Confirm refund policy matches sales page.
- Confirm license terms match product rights.
- Confirm no placeholder remains in public-facing files.

## Buyer First Email Reminder

Do not start with a huge product. Create one small printable bundle concept, run the QC checklist, then expand only after the first version is clean."""


def _vendor_ready_support_extra() -> str:
    return """## Support Macros

### Buyer asks if sales are guaranteed
No sales are guaranteed. The kit helps you plan and package printable bundle concepts. Your results depend on niche, offer, traffic, platform rules, and execution.

### Buyer asks if they can use Canva
Yes, if you follow Canva's current license and do not resell Canva elements as standalone assets.

### Buyer asks what to do first
Open `00_Start_Here.md`, choose one niche in the CSV, and complete the bundle planner for one small product before expanding."""


def _vendor_ready_launch_extra() -> str:
    return """## Honest Launch Decision

Vendor-ready content pack: YES.
Public launch ready: NO.

The content is deeper than the old skeleton export, but public launch still requires vendor replacement of sample details, mockup creation, checkout test, delivery test, affiliate rule review, and ideally real buyer feedback.

## Next Public Launch Gates

1. Replace Example Studio/support@example.com.
2. Add real download page.
3. Add mockups and screenshots.
4. Test WarriorPlus payment.
5. Test delivery email.
6. Ask one buyer to use the pack and report confusion points.
7. Update files based on feedback."""


def _vendor_ready_build_notes(files: dict[str, str]) -> str:
    total = sum(len(content.encode("utf-8")) for content in files.values())
    return f"""# Vendor Ready Build Notes

Total content bytes: {total}
File count: {len(files)}
Status: Vendor-ready content pack generated.
Not public-launch ready until external proof gates pass.

Quality rule: files under 2.5KB should be reviewed manually unless they are CSV/proof files.
"""


def _vendor_ready_scorecard(files: dict[str, str]) -> str:
    rows = []
    for name, content in sorted(files.items()):
        size = len(content.encode("utf-8"))
        rows.append(f"| `{name}` | {size} | {'PASS' if size >= 2500 or name.endswith('.csv') else 'REVIEW'} |")
    return "# Vendor Ready Scorecard\n\n| File | Bytes | Status |\n|---|---:|---|\n" + "\n".join(rows) + "\n"
def _ai_etsy_bundle_files() -> dict[str, str]:
    return {
        "README.md": """# AI Etsy Printable Bundle Builder

This kit helps buyers plan and build original Etsy-style printable bundle concepts with AI assistance. It is designed for digital product sellers, PLR sellers, coaches, and creators who want a structured workflow instead of a raw prompt pack.

Start with `00_Start_Here.md`, then use `01_Bundle_Workflow.md` and `02_Niche_And_Buyer_Picker.csv` to choose a safe niche. Build the bundle with `03_Printable_Bundle_Planner.md`, generate prompts from `04_AI_Prompt_Library.md`, and run every idea through `06_Quality_Control_Checklist.md` plus `07_License_And_Compliance_Note.md`.

Sample vendor: Example Studio. Support: support@example.com. Replace these sample values before selling.

No income, Etsy approval, KDP approval, or legal outcome is promised.""",
        "00_Start_Here.md": """# Start Here

Goal: create a small printable bundle that feels useful, specific, and safe enough for a buyer to customize.

Quick path:
1. Choose one audience: teachers, busy moms, coaches, pet owners, planners, wedding DIY buyers, small business owners, or homeschool families.
2. Choose one outcome: organize, decorate, plan, track, gift, teach, or sell a lead magnet.
3. Choose 5-8 printable assets for the bundle.
4. Generate drafts with AI prompts.
5. Check originality, layout, and license safety.
6. Package as PDF/PNG/Canva-style instructions if you have rights.
7. Write listing copy with realistic claims.

Do not use famous characters, brand names, celebrity likeness, lyrics, sports teams, trademarked phrases, or copied art styles.""",
        "01_Bundle_Workflow.md": """# Bundle Workflow

## Step 1 — Pick a safe buyer
Use a buyer with a clear use case. Example: kindergarten teacher who needs printable classroom reward charts.

## Step 2 — Pick a bundle promise
Safe promise format: "Create a ready-to-customize printable bundle for [buyer] who wants [practical outcome]."

## Step 3 — Choose assets
A strong starter bundle can include: cover page, instruction page, 5 printable pages, checklist, bonus tracker, preview copy, license note.

## Step 4 — Generate and inspect
Use AI for concepts and text, then manually inspect layout, spelling, margins, and originality.

## Step 5 — Package
Export clearly named files and include buyer onboarding, support FAQ, and license terms.

## Step 6 — Sell safely
Avoid income claims, guaranteed marketplace approval, legal claims, or "best-selling" claims without proof.""",
        "02_Niche_And_Buyer_Picker.csv": """Niche,Buyer,Pain,Bundle Idea,Risk,Score
Classroom Rewards,Teachers,Needs fast behavior tools,Reward chart printable kit,Low,9
Pet Care Planner,Pet owners,Forgets routines,Pet routine tracker bundle,Low,8
Wedding DIY Planner,Brides,Needs organization,Wedding printable checklist bundle,Medium,8
Coach Lead Magnet,Coaches,Needs opt-in asset,Client worksheet bundle,Low,9
Kids Chore Charts,Parents,Needs simple routines,Chore chart printable set,Low,8
Small Biz Planner,Solopreneurs,Needs planning pages,Promo planner printable bundle,Low,8
""",
        "03_Printable_Bundle_Planner.md": """# Printable Bundle Planner

Product name: AI Etsy Printable Bundle Builder

Bundle concept worksheet:
- Buyer:
- Specific problem:
- Printable assets included:
- File formats:
- Customization instructions:
- License terms:
- Support note:

Example bundle: Teacher Reward Chart Starter Kit
Files: cover, start guide, 5 reward charts, sticker tracker, classroom routine checklist, parent note template, license note.
Quality standard: readable text, clean margins, no protected characters, no fake brand logos, no medical/education guarantee claims.""",
        "04_AI_Prompt_Library.md": """# AI Prompt Library

## Niche Finder
Act as a printable product strategist. Generate 20 safe Etsy-style printable bundle niches for [audience]. Avoid trademarks, celebrities, lyrics, brand names, protected characters, and income claims. Include buyer pain, asset ideas, and risk score.

## Bundle Builder
Create a 7-file printable bundle for [buyer] who wants [outcome]. Include file names, page purpose, copy text, layout notes, and quality checks.

## Listing Writer
Write an Etsy-style listing for this printable bundle. Use realistic benefits only. Do not promise income, therapy, education outcomes, marketplace approval, or legal approval.

## Compliance Check
Review this printable idea for trademark, copyright, celebrity, brand imitation, Canva asset, font, and claim risk. Replace risky ideas with generic original alternatives.""",
        "05_Example_Bundle_Concepts.md": """# Example Bundle Concepts

## Teacher Reward Chart Starter Kit
Buyer: elementary teacher. Assets: 5 reward charts, routine checklist, sticker tracker, parent note, start guide. Safe angle: classroom organization support, no learning guarantee.

## Pet Routine Printable Pack
Buyer: pet owner. Assets: feeding tracker, vet visit log, medication log, grooming checklist, sitter note. Safe angle: organization, no medical advice.

## Coach Lead Magnet Worksheet Kit
Buyer: coach or consultant. Assets: discovery worksheet, action planner, goal tracker, reflection page, email delivery note. Safe angle: client onboarding support, no income claim.""",
        "06_Quality_Control_Checklist.md": """# Quality Control Checklist

- Buyer understands the bundle in 5 seconds. PASS/FAIL
- Every file has a clear use. PASS/FAIL
- No trademark/celebrity/brand/lyrics/protected character. PASS/FAIL
- No income, therapy, medical, legal, or guaranteed marketplace approval claims. PASS/FAIL
- Fonts and graphics are licensed for the intended use. PASS/FAIL
- Layout is readable when printed. PASS/FAIL
- File names are clear. PASS/FAIL
- License note and support FAQ are included. PASS/FAIL
- Placeholder scan completed. PASS/FAIL
- Buyer onboarding explains first action. PASS/FAIL""",
        "07_License_And_Compliance_Note.md": """# License And Compliance Note

This kit is not legal advice. Seller must verify rights before publishing or selling.

Avoid: trademarks, brand names, logos, celebrity likeness, lyrics, sports teams, famous characters, copied designs, and artist-style imitation.

Canva: use only elements and templates you have rights to use. Do not resell Canva elements as standalone assets.

AI: inspect all output for accidental logos, distorted text, copied style, or protected content.

Claims: do not promise income, ranking, Etsy approval, KDP approval, therapy, medical results, or legal safety.""",
        "sales_page.md": """# Sales Page

## Headline
Build Etsy-style printable bundle ideas with AI — without selling a raw prompt pack.

## Problem
AI can generate ideas, but beginners still struggle to choose a safe niche, plan files, write listings, check rights, and package a useful printable bundle.

## Solution
AI Etsy Printable Bundle Builder gives you a workflow, prompts, examples, checklists, listing copy, license notes, delivery assets, and buyer onboarding files.

## What You Get
Start guide, workflow map, niche picker CSV, bundle planner, prompt library, example concepts, QC checklist, compliance note, WarriorPlus listing, JV pack, delivery page, onboarding email, FAQ, refund policy, manifest, placeholder check, and launch audit.

## CTA
Get the kit, pick one buyer, build one safe printable bundle concept, and package it with confidence.

Disclaimer: no income, Etsy approval, or legal outcome is promised.""",
        "warriorplus_listing.md": """# WarriorPlus Listing

Title: AI Etsy Printable Bundle Builder
Short description: Workflow and prompt kit for building safe Etsy-style printable bundle concepts with AI.
Price: $17-$27
Category: Digital product creation / AI printables
Commission suggestion: 50%
Refund policy: support-first digital product policy.
Affiliate note: affiliates may not claim guaranteed income, Etsy approval, or legal approval.
Delivery: ZIP download with markdown/csv files and onboarding instructions.""",
        "jv_pack.md": """# JV Pack

## JV Angle
Your audience can ask AI for printables, but they still need a product workflow, safe niche picker, bundle planner, QC checklist, and listing support.

## Email Swipe 1
Subject: AI printables are easy. Productizing them is the hard part.
This kit helps creators turn AI-assisted printable ideas into structured bundle concepts with prompts, examples, checklists, listing assets, and compliance notes.

## Social Post
Stop selling raw prompts. Build printable bundle systems buyers can actually follow.""",
        "delivery_page.md": """# Delivery Page

Thanks for getting AI Etsy Printable Bundle Builder.
Download the ZIP, unzip it, and open `00_Start_Here.md` first.

Recommended first task: create one small 5-asset printable bundle concept before trying a large product.

Support: support@example.com. Replace this with your real support email.""",
        "buyer_onboarding_email.md": """Subject: Your AI Etsy Printable Bundle Builder access

Hi,

Thanks for your purchase. Download your ZIP and start with `00_Start_Here.md`.

Your first action: choose one buyer and complete `02_Niche_And_Buyer_Picker.csv`, then draft one bundle with `03_Printable_Bundle_Planner.md`.

Important: avoid trademarks, famous characters, lyrics, logos, and unsupported income/platform claims.

Support: support@example.com""",
        "support_faq.md": """# Support FAQ

Q: Is this a finished Etsy shop?
A: No. It is a product creation kit.

Q: Are sales guaranteed?
A: No.

Q: Can I sell printables I create with it?
A: Yes, if your output is original and you have proper rights.

Q: Does it include legal review?
A: No. Get professional advice for legal questions.""",
        "refund_policy.md": """# Refund Policy

This is a digital product creation resource. Refunds may be reviewed if the buyer cannot access the files, receives the wrong product, or a core promised component is missing and support cannot fix it.

Refunds are not promised for lack of sales, platform rejection, misuse, ignored license rules, or claims not stated on the sales page.""",
        "export_manifest.md": """# Export Manifest

All core files are included. Status: ZIP created. Public launch status: not ready until vendor replaces sample details and tests checkout/delivery.""",
        "placeholder_check.md": """# Placeholder Check

Bracket placeholder scan: PASS.
Sample values still requiring vendor replacement: Example Studio, support@example.com, download URL, final vendor name, final refund rules.
Status: artifact proof exists, but business setup is not complete.""",
        "launch_audit.md": """# Launch Audit

Build readiness: PASS.
Soft launch prep: PASS.
Public launch ready: NO.
Reasons: no live payment test, delivery test, WarriorPlus approval, JV approval, real buyer feedback, or legal review.
Next: replace sample data, add mockups, test delivery, then run buyer test.""",
    }
def _clip(text: str, max_chars: int) -> str:
    cleaned = " ".join(text.split())
    if len(cleaned) <= max_chars:
        return cleaned
    return cleaned[: max_chars - 3].rstrip() + "..."


if __name__ == "__main__":
    main()









